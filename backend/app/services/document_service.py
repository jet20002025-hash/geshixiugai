from __future__ import annotations

import base64
import io
import json
import os
import re
import shutil
import sys
import uuid
import xml.sax.saxutils
from datetime import datetime
from pathlib import Path
from typing import Dict, Tuple, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.text.paragraph import Paragraph
from fastapi import UploadFile

from .utils import docx_format_utils
from .storage_factory import get_storage
from .thesis_format_standard import (
    FONT_STANDARDS,
    STYLE_MAPPING_RULES,
    DEFAULT_STYLE,
    PAGE_SETTINGS,
    HEADER_SETTINGS,
    PAGE_NUMBER_FORMAT,
    REFERENCE_REQUIREMENTS,
)


class DocumentService:
    def __init__(self, document_dir: Path, template_dir: Path) -> None:
        self.document_dir = document_dir
        self.template_dir = template_dir
        self.document_dir.mkdir(parents=True, exist_ok=True)
        # 获取存储实例（如果可用）
        self.storage = get_storage()
        self.use_storage = self.storage is not None
    
    def _log_to_file(self, message: str) -> None:
        """将日志消息同时输出到 stderr 和日志文件（双重保险）"""
        print(message, file=sys.stderr, flush=True)
        try:
            with open("/var/log/geshixiugai/error.log", "a") as f:
                f.write(f"{message}\n")
        except Exception:
            pass

    async def process_document(
        self, 
        template_id: Optional[str] = None, 
        university_id: Optional[str] = None,
        upload: Optional[UploadFile] = None
    ) -> Tuple[str, Dict]:
        if not upload or not upload.filename or not upload.filename.lower().endswith(".docx"):
            raise ValueError("仅支持 docx 文档")
        
        # 验证参数：template_id 和 university_id 必须二选一
        if not template_id and not university_id:
            raise ValueError("必须提供 template_id 或 university_id 之一")
        if template_id and university_id:
            raise ValueError("不能同时提供 template_id 和 university_id")
        
        # 加载模板元数据
        if university_id:
            template_metadata = self._load_university_template(university_id)
        else:
            template_metadata = self._load_template(template_id)
        document_id = uuid.uuid4().hex
        # 生成唯一的下载 token，用于验证用户身份
        download_token = uuid.uuid4().hex
        task_dir = self.document_dir / document_id
        task_dir.mkdir(parents=True, exist_ok=True)

        # 保存原始文件名
        original_filename = upload.filename
        
        original_path = task_dir / "original.docx"
        original_path.write_bytes(await upload.read())

        # 加载文档
        document = Document(original_path)
        
        # 诊断1：检查原始文档中诚信承诺和摘要的分页情况
        self._log_to_file(f"[诊断] ========== 开始诊断：原始文档 ==========")
        original_diagnosis = self._diagnose_integrity_abstract_separation(document)
        self._log_to_file(f"[诊断] 原始文档诊断结果: {original_diagnosis['issue'] if original_diagnosis['issue'] else '有分页符'}")
        self._log_to_file(f"[诊断] 分页符位置: {len(original_diagnosis['page_break_locations'])} 个")
        
        # 应用页面设置（优先使用标准）
        self._apply_page_settings(document)
        
        # 检测页眉（不修改，只检测）
        # 不再自动应用页眉，只检测是否存在
        
        # 合并模板规则和标准规则（标准优先）
        # 如果是预设模板，使用 parameters；如果是自定义模板，使用 styles
        if template_metadata.get("university_id"):
            # 预设模板：从 parameters 中提取格式规则
            university_params = template_metadata.get("parameters", {})
            template_rules = self._convert_university_params_to_rules(university_params)
        else:
            # 自定义模板：使用 styles
            template_rules = template_metadata.get("styles", {})
        
        merged_rules = self._merge_rules_with_standard(template_rules)
        
        final_doc, stats = self._apply_rules(
            document=document,
            rules=merged_rules,
            default_style=template_metadata.get("default_style") or DEFAULT_STYLE,
        )
        
        # 检测图片并检查图题
        figure_issues = self._check_figure_captions(final_doc)
        if figure_issues:
            stats["figure_issues"] = figure_issues
        
        # 检测参考文献引用标注
        reference_issues = self._check_reference_citations(final_doc)
        if reference_issues:
            stats["reference_issues"] = reference_issues
        
        # 修复前先诊断一次，记录初始状态
        self._log_to_file(f"[检测] ========== 修复前检测：诚信承诺和摘要分页结果 ==========")
        pre_fix_diagnosis = self._diagnose_integrity_abstract_separation(final_doc)
        if pre_fix_diagnosis["has_page_break_between"]:
            self._log_to_file(f"[检测] ✅ 修复前已有分页符，无需修复")
        else:
            self._log_to_file(f"[检测] ❌ 修复前没有分页符，需要修复")
        
        # 确保诚信承诺和摘要分开在不同页（在空白行删除之前）
        self._log_to_file(f"[修复] ========== 开始修复：确保诚信承诺和摘要分开在不同页 ==========")
        separation_fixed = self._ensure_integrity_abstract_separation(final_doc)
        if separation_fixed:
            self._log_to_file(f"[修复] ✅ 已确保诚信承诺和摘要分开在不同页")
            stats["integrity_abstract_separation_fixed"] = True
        else:
            self._log_to_file(f"[修复] ⚠️ 未进行修复（可能已有分页符或未找到诚信承诺/摘要）")
        
        # 修复后再次检测，确认分页结果
        self._log_to_file(f"[检测] ========== 修复后检测：诚信承诺和摘要分页结果 ==========")
        post_fix_diagnosis = self._diagnose_integrity_abstract_separation(final_doc)
        if post_fix_diagnosis["has_page_break_between"]:
            self._log_to_file(f"[检测] ✅ 修复成功：诚信承诺和摘要已分开在不同页")
            self._log_to_file(f"[检测] 分页符位置: {len(post_fix_diagnosis['page_break_locations'])} 个")
            for loc in post_fix_diagnosis['page_break_locations']:
                self._log_to_file(f"[检测]   - 段落 {loc['index']}: {loc['type']}")
            stats["post_fix_separation_status"] = "已分开"
        else:
            self._log_to_file(f"[检测] ❌ 修复失败：诚信承诺和摘要仍然没有分页符")
            self._log_to_file(f"[检测] 问题: {post_fix_diagnosis.get('issue', '未知')}")
            stats["post_fix_separation_status"] = "未分开"
        stats["post_fix_diagnosis"] = post_fix_diagnosis
        
        # 确保中文摘要和英文摘要分开在不同页
        self._log_to_file(f"[修复] ========== 开始修复：确保中文摘要和英文摘要分开在不同页 ==========")
        abstract_separation_fixed = self._ensure_abstract_separation(final_doc)
        if abstract_separation_fixed:
            self._log_to_file(f"[修复] ✅ 已确保中文摘要和英文摘要分开在不同页")
            stats["abstract_separation_fixed"] = True
        
        # 检测大段空白
        blank_issues = self._check_excessive_blanks(final_doc)
        if blank_issues:
            stats["blank_issues"] = blank_issues
        
        # 检测并删除整页空白页（不允许整页空白）
        blank_page_issues = self._check_and_remove_blank_pages(final_doc)
        if blank_page_issues:
            stats["blank_page_issues"] = blank_page_issues
        
        # 检测页眉
        header_issues = self._check_header(final_doc)
        if header_issues:
            stats["header_issues"] = header_issues

        final_path = task_dir / "final.docx"
        final_doc.save(final_path)

        # 诊断2：检查格式修改后的文档中诚信承诺和摘要的分页情况
        self._log_to_file(f"[诊断] ========== 开始诊断：格式修改后的文档 ==========")
        final_diagnosis = self._diagnose_integrity_abstract_separation(final_doc)
        self._log_to_file(f"[诊断] 格式修改后诊断结果: {final_diagnosis['issue'] if final_diagnosis['issue'] else '有分页符'}")
        self._log_to_file(f"[诊断] 分页符位置: {len(final_diagnosis['page_break_locations'])} 个")
        
        # 对比诊断结果
        if original_diagnosis["has_page_break_between"] and not final_diagnosis["has_page_break_between"]:
            self._log_to_file(f"[诊断] ⚠️ 警告：格式修改过程中丢失了分页符！")
            stats["diagnosis_warning"] = "格式修改过程中丢失了诚信承诺和摘要之间的分页符"
        elif not original_diagnosis["has_page_break_between"]:
            self._log_to_file(f"[诊断] ⚠️ 警告：原始文档中就没有分页符！")
            stats["diagnosis_warning"] = "原始文档中诚信承诺和摘要之间没有分页符"
        
        stats["original_diagnosis"] = original_diagnosis
        stats["final_diagnosis"] = final_diagnosis

        # 验证格式修改是否正确：对比原始文档和修改后的文档
        print(f"[格式验证] 开始验证格式修改是否正确...")
        format_verification = self._verify_format_changes(original_path, final_path, merged_rules)
        stats["format_verification"] = format_verification
        
        # 输出格式验证结果
        if format_verification.get("errors"):
            print(f"[格式验证] ⚠️ 发现 {len(format_verification['errors'])} 个格式问题：")
            for error in format_verification["errors"][:10]:  # 只显示前10个
                print(f"[格式验证]   - {error}")
        else:
            print(f"[格式验证] ✅ 格式验证通过，所有格式修改正确")
        
        if format_verification.get("summary"):
            print(f"[格式验证] 格式修改摘要：")
            for key, value in format_verification["summary"].items():
                print(f"[格式验证]   - {key}: {value}")

        preview_path = task_dir / "preview.docx"
        self._generate_watermarked_preview(final_path, preview_path)
        
        # 诊断3：检查预览文档（带水印的Word文档）中诚信承诺和摘要的分页情况
        self._log_to_file(f"[诊断] ========== 开始诊断：预览文档（带水印） ==========")
        preview_doc = Document(preview_path)
        preview_diagnosis = self._diagnose_integrity_abstract_separation(preview_doc)
        self._log_to_file(f"[诊断] 预览文档诊断结果: {preview_diagnosis['issue'] if preview_diagnosis['issue'] else '有分页符'}")
        self._log_to_file(f"[诊断] 分页符位置: {len(preview_diagnosis['page_break_locations'])} 个")
        
        # 对比诊断结果
        if final_diagnosis["has_page_break_between"] and not preview_diagnosis["has_page_break_between"]:
            self._log_to_file(f"[诊断] ⚠️ 警告：生成预览文档过程中丢失了分页符！")
            stats["diagnosis_warning"] = "生成预览文档过程中丢失了诚信承诺和摘要之间的分页符"
        stats["preview_diagnosis"] = preview_diagnosis
        
        # 生成PDF预览（优先使用LibreOffice直接转换，保持格式完全一致）
        # 预览PDF = 修改后的文档 + 水印，格式与最终文档完全一致，仅格式为PDF
        pdf_path = preview_path.with_suffix('.pdf')
        print(f"[预览] 开始生成PDF预览: {pdf_path}")
        print(f"[预览] 预览策略：直接从Word转PDF（使用LibreOffice），保持格式完全一致，仅添加水印")
        # 初始化 html_path，确保在所有情况下都有定义
        html_path = preview_path.with_suffix('.html')
        
        # 优先使用LibreOffice直接从Word转PDF（格式完美，只加水印，与最终文档完全一致）
        # 先生成临时PDF，然后添加水印
        temp_pdf_path = pdf_path.with_suffix('.temp.pdf')
        pdf_success = self._try_libreoffice_pdf_conversion(preview_path, temp_pdf_path)
        
        # 检测5：PDF生成后，检查PDF中诚信承诺和摘要的分页结果
        if pdf_success and temp_pdf_path.exists():
            try:
                from pypdf import PdfReader
                pdf_reader = PdfReader(str(temp_pdf_path))
                pdf_page_count = len(pdf_reader.pages)
                self._log_to_file(f"[检测] ========== PDF生成后检测：诚信承诺和摘要分页结果 ==========")
                self._log_to_file(f"[检测] PDF总页数: {pdf_page_count}")
                
                # 检查每一页，找到诚信承诺和摘要所在的页面
                integrity_page = None
                abstract_page = None
                
                # 检查前10页（通常诚信承诺和摘要在前几页，但可能因为格式问题延后）
                check_pages = min(10, pdf_page_count)
                self._log_to_file(f"[检测] 检查前 {check_pages} 页以查找诚信承诺和摘要")
                
                for page_num in range(check_pages):
                    try:
                        page_text = pdf_reader.pages[page_num].extract_text()
                        # 调试：输出每页的前100个字符（用于排查）
                        if page_num < 5:  # 只输出前5页的调试信息
                            preview_text = page_text[:100].replace('\n', ' ').strip()
                            self._log_to_file(f"[检测] 第 {page_num + 1} 页文本预览: {preview_text}...")
                        
                        # 检查是否包含诚信承诺（支持多种变体）
                        has_integrity = (
                            '诚信承诺' in page_text or 
                            ('诚' in page_text and '信' in page_text and '承' in page_text and '诺' in page_text) or
                            '诚信' in page_text and '承诺' in page_text
                        )
                        
                        # 检查是否包含摘要（支持多种变体）
                        has_abstract = (
                            '摘要' in page_text or
                            'ABSTRACT' in page_text.upper() or
                            ('摘' in page_text and '要' in page_text)
                        )
                        
                        if has_integrity and integrity_page is None:
                            integrity_page = page_num + 1
                            self._log_to_file(f"[检测] ✅ 第 {page_num + 1} 页包含诚信承诺")
                        if has_abstract and abstract_page is None:
                            abstract_page = page_num + 1
                            self._log_to_file(f"[检测] ✅ 第 {page_num + 1} 页包含摘要")
                    except Exception as e:
                        self._log_to_file(f"[检测] ❌ 无法提取第 {page_num + 1} 页文本: {e}")
                
                # 判断结果并输出
                self._log_to_file(f"[检测] ========== PDF分页结果 ==========")
                # 始终输出页码信息（如果找到）
                if integrity_page is not None:
                    self._log_to_file(f"[检测] 📄 诚信承诺所在页码: 第 {integrity_page} 页")
                    stats["pdf_integrity_page"] = integrity_page
                if abstract_page is not None:
                    self._log_to_file(f"[检测] 📄 摘要所在页码: 第 {abstract_page} 页")
                    stats["pdf_abstract_page"] = abstract_page
                
                # 判断分页情况
                if integrity_page is not None and abstract_page is not None:
                    if integrity_page == abstract_page:
                        self._log_to_file(f"[检测] ❌ PDF中诚信承诺和摘要在同一页（第 {integrity_page} 页）")
                        self._log_to_file(f"[检测] ⚠️ 警告：Word转PDF过程中分页符可能失效")
                        stats["pdf_separation_status"] = "合并在同一页"
                        stats["pdf_separation_warning"] = f"PDF中诚信承诺和摘要在同一页（第 {integrity_page} 页），Word转PDF过程中分页符可能失效"
                    else:
                        self._log_to_file(f"[检测] ✅ PDF中诚信承诺和摘要分开在不同页")
                        self._log_to_file(f"[检测] 📊 页码对比: 诚信承诺(第 {integrity_page} 页) vs 摘要(第 {abstract_page} 页)")
                        stats["pdf_separation_status"] = "已分开"
                        stats["pdf_separation_pages"] = {"integrity": integrity_page, "abstract": abstract_page}
                elif integrity_page is not None:
                    self._log_to_file(f"[检测] ⚠️ 找到诚信承诺（第 {integrity_page} 页），但未找到摘要")
                    stats["pdf_separation_status"] = "未找到摘要"
                elif abstract_page is not None:
                    self._log_to_file(f"[检测] ⚠️ 找到摘要（第 {abstract_page} 页），但未找到诚信承诺")
                    stats["pdf_separation_status"] = "未找到诚信承诺"
                else:
                    self._log_to_file(f"[检测] ⚠️ 未找到诚信承诺和摘要")
                    stats["pdf_separation_status"] = "未找到"
                self._log_to_file(f"[检测] ========================================")
                    
            except Exception as e:
                self._log_to_file(f"[检测] ❌ 无法读取PDF: {e}")
                stats["pdf_separation_status"] = "检测失败"
        
        # 如果LibreOffice转换失败，回退到HTML转PDF（不推荐，格式会有变化）
        if not pdf_success:
            print(f"[预览] ⚠️ LibreOffice转换失败，回退到HTML转PDF（格式可能有变化，不推荐）")
            pdf_success = self._generate_pdf_preview(preview_path, temp_pdf_path, stats)
        
        if pdf_success and temp_pdf_path.exists():
            # 为预览PDF添加水印（确保免费用户只能看到带水印的PDF）
            print(f"[预览] 为预览PDF添加水印...")
            watermark_success = self._add_pdf_watermarks(
                pdf_path=temp_pdf_path,
                output_path=pdf_path,
                watermark_text="www.geshixiugai.cn",
                watermarks_per_page=10
            )
            
            if watermark_success and pdf_path.exists():
                pdf_size = pdf_path.stat().st_size
                print(f"[预览] ✅ PDF预览生成成功（已添加水印）: {pdf_path}, 大小: {pdf_size / 1024:.2f} KB")
                # 删除临时PDF文件
                if temp_pdf_path.exists():
                    temp_pdf_path.unlink()
            else:
                print(f"[预览] ⚠️ 水印添加失败，使用原始PDF")
                if temp_pdf_path.exists():
                    # 如果水印添加失败，使用原始PDF（但应该确保原始PDF也有水印）
                    temp_pdf_path.rename(pdf_path)
                pdf_success = True
        elif pdf_success:
            print(f"[预览] ⚠️ PDF生成返回成功但文件不存在: {temp_pdf_path}")
            pdf_success = False
        
        if not pdf_success:
            # 回退到HTML预览
            print(f"[预览] PDF生成失败，回退到HTML预览")
            self._generate_html_preview(preview_path, html_path, stats)
            if html_path.exists():
                html_size = html_path.stat().st_size
                print(f"[预览] HTML预览生成成功: {html_path}, 大小: {html_size / 1024:.2f} KB")
            else:
                print(f"[预览] ⚠️ HTML预览生成失败，文件不存在: {html_path}")

        report_data = {
            "document_id": document_id,
            "template_id": template_id,
            "summary": stats,
        }

        report_path = task_dir / "report.json"
        report_path.write_text(json.dumps(report_data, ensure_ascii=False, indent=2), encoding="utf-8")

        # 如果使用云存储，将文件上传到云存储
        if self.use_storage:
            files_to_save = {
                "original": original_path,
                "final": final_path,
                "preview": preview_path,
                "report": report_path,
            }
            # 添加PDF或HTML预览文件
            # 注意：pdf_path 和 html_path 已经在上面定义过了
            if pdf_path.exists():
                pdf_size = pdf_path.stat().st_size
                print(f"[存储] 准备上传PDF预览文件: {pdf_path}, 大小: {pdf_size / 1024:.2f} KB")
                files_to_save["pdf"] = pdf_path
            else:
                print(f"[存储] PDF文件不存在，检查HTML文件")
                # html_path 已经在上面定义过了，直接使用
                if html_path.exists():
                    html_size = html_path.stat().st_size
                    print(f"[存储] 准备上传HTML预览文件: {html_path}, 大小: {html_size / 1024:.2f} KB")
                    files_to_save["html"] = html_path
                else:
                    print(f"[存储] ⚠️ 警告: PDF和HTML预览文件都不存在！")
            
            self._save_to_storage(document_id, files_to_save)

        # 确保 template_id 不为 None（如果使用 university_id，则使用 university_id 作为标识）
        final_template_id = template_id if template_id else (f"university_{university_id}" if university_id else "unknown")
        
        metadata = {
            "document_id": document_id,
            "template_id": final_template_id,
            "status": "completed",
            "paid": False,
            "download_token": download_token,  # 下载验证 token
            "original_filename": original_filename,  # 保存原始文件名
            "summary": stats,
            "report_path": str(report_path),
            "preview_path": str(preview_path),
            "preview_html_path": str(html_path),
            "final_path": str(final_path),
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat(),
        }

        metadata_path = task_dir / "metadata.json"
        metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
        
        # 如果使用云存储，也上传 metadata
        if self.use_storage:
            self._save_file_to_storage(f"documents/{document_id}/metadata.json", metadata_path.read_bytes())
        
        return document_id, stats

    def get_document_metadata(self, document_id: str) -> Dict:
        # 优先从云存储读取
        if self.use_storage:
            metadata_key = f"documents/{document_id}/metadata.json"
            if self.storage.file_exists(metadata_key):
                content = self.storage.download_file(metadata_key)
                if content:
                    return json.loads(content.decode("utf-8"))
        
        # 回退到本地文件系统
        metadata_path = self.document_dir / document_id / "metadata.json"
        if not metadata_path.exists():
            return {}
        return json.loads(metadata_path.read_text(encoding="utf-8"))

    def update_metadata(self, document_id: str, **kwargs) -> Dict:
        # 先加载 metadata（优先从存储）
        data = self.get_document_metadata(document_id)
        if not data:
            raise FileNotFoundError("metadata not found")
        
        # 更新数据
        data.update(kwargs)
        data["updated_at"] = datetime.utcnow().isoformat()
        
        # 保存到本地和存储
        task_dir = self.document_dir / document_id
        task_dir.mkdir(parents=True, exist_ok=True)
        metadata_path = task_dir / "metadata.json"
        metadata_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        
        # 如果使用云存储，也更新存储中的 metadata
        if self.use_storage:
            metadata_key = f"documents/{document_id}/metadata.json"
            content = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
            self._save_file_to_storage(metadata_key, content)
        
        return data

    def _load_template(self, template_id: str) -> Dict:
        """加载用户上传的自定义模板"""
        metadata_path = self.template_dir / template_id / "metadata.json"
        if not metadata_path.exists():
            raise FileNotFoundError("template not found")
        return json.loads(metadata_path.read_text(encoding="utf-8"))
    
    def _load_university_template(self, university_id: str) -> Dict:
        """加载预设大学模板"""
        from .university_template_service import UniversityTemplateService
        
        service = UniversityTemplateService()
        template = service.get_university_template(university_id)
        if not template:
            raise FileNotFoundError(f"未找到大学模板: {university_id}")
        
        # 将预设模板转换为与自定义模板相同的格式
        parameters = template.get("parameters", {})
        
        # 构建模板元数据格式
        metadata = {
            "template_id": f"university_{university_id}",
            "name": template.get("display_name", template.get("name")),
            "university_id": university_id,
            "styles": {},  # 预设模板不使用 styles，而是使用 parameters
            "parameters": parameters,  # 预设模板的参数
            "default_style": "body_text",
        }
        
        return metadata

    def _paragraph_has_image_or_equation(self, paragraph) -> bool:
        """判断段落是否包含图片或公式"""
        # 检查是否包含图片
        has_image = False
        try:
            # 方法1: 检查段落中的runs是否包含图片
            for run in paragraph.runs:
                if not hasattr(run, 'element'):
                    continue
                run_xml = str(run.element.xml)
                # 排除VML形状的水印
                if 'v:shape' in run_xml.lower() and 'textpath' in run_xml.lower():
                    continue
                # 检查是否包含真正的图片元素
                if ('pic:pic' in run_xml or 'a:blip' in run_xml) and ('r:embed' in run_xml or 'r:link' in run_xml or 'a:blip' in run_xml):
                    has_image = True
                    break
        except:
            pass
        
        # 方法2: 检查段落元素中是否包含图片
        if not has_image:
            try:
                para_xml = str(paragraph._element.xml)
                if 'v:shape' in para_xml.lower() and 'textpath' in para_xml.lower():
                    pass  # 这是水印，跳过
                elif ('pic:pic' in para_xml or 'a:blip' in para_xml) and ('r:embed' in para_xml or 'r:link' in para_xml or 'a:blip' in para_xml):
                    has_image = True
            except:
                pass
        
        # 方法3: 使用xpath查找drawing元素
        if not has_image:
            try:
                from docx.oxml.ns import qn
                # 使用findall配合qn，而不是xpath with namespaces
                drawings = paragraph._element.findall('.//' + qn('w:drawing'))
                if drawings:
                    for drawing in drawings:
                        drawing_xml = str(drawing.xml)
                        if 'v:shape' in drawing_xml.lower() and 'textpath' in drawing_xml.lower():
                            continue
                        if ('pic:pic' in drawing_xml or 'a:blip' in drawing_xml) and ('r:embed' in drawing_xml or 'r:link' in drawing_xml or 'a:blip' in drawing_xml):
                            has_image = True
                            break
            except:
                pass
        
        # 检查是否包含公式（Office Math 或 MathType）
        has_equation = False
        try:
            para_xml = str(paragraph._element.xml)
            # 检查Office Math (oMath)
            if 'm:oMath' in para_xml or 'm:oMathPara' in para_xml:
                has_equation = True
            # 检查MathType公式（通常包含object标签）
            elif 'object' in para_xml.lower() and ('mathtype' in para_xml.lower() or 'equation' in para_xml.lower()):
                has_equation = True
            # 检查段落中的runs是否包含公式
            if not has_equation:
                for run in paragraph.runs:
                    if not hasattr(run, 'element'):
                        continue
                    run_xml = str(run.element.xml)
                    if 'm:oMath' in run_xml or 'm:oMathPara' in run_xml:
                        has_equation = True
                        break
        except:
            pass
        
        return has_image or has_equation
    
    def _paragraph_has_flowchart(self, paragraph) -> bool:
        """判断段落是否包含流程图（由多个形状组成的流程图）"""
        try:
            para_xml = str(paragraph._element.xml)
            
            # 方法1: 检测 Word Processing Shapes (wps:wsp) - 现代 Word 文档中的形状
            # 流程图通常包含多个形状，如果段落中有多个 wps:wsp 元素，可能是流程图
            if 'wps:wsp' in para_xml:
                # 计算形状数量
                shape_count = para_xml.count('wps:wsp')
                # 如果包含多个形状（至少2个），可能是流程图
                if shape_count >= 2:
                    return True
            
            # 方法2: 检测 VML Shapes (v:shape) - 旧版 Word 文档中的形状
            # 排除水印（包含 textpath 的 v:shape 通常是水印）
            vml_shapes = []
            if 'v:shape' in para_xml.lower():
                # 检查是否有多个非水印的形状
                # 简单判断：如果包含 v:shape 但不包含 textpath，可能是流程图的一部分
                vml_count = para_xml.lower().count('v:shape')
                textpath_count = para_xml.lower().count('textpath')
                # 如果有形状且不是水印，可能是流程图
                if vml_count > textpath_count and vml_count >= 2:
                    return True
            
            # 方法3: 检测 SmartArt 流程图
            # SmartArt 在 XML 中通常包含 'smartart' 或特定的命名空间
            if 'smartart' in para_xml.lower() or 'dgm:' in para_xml:
                return True
            
            # 方法4: 检测 drawing 元素中的多个形状
            # 使用 findall 查找 drawing 元素，检查是否包含多个形状
            try:
                from docx.oxml.ns import qn
                drawings = paragraph._element.findall('.//' + qn('w:drawing'))
                if drawings:
                    # 检查每个 drawing 元素中是否包含多个形状
                    for drawing in drawings:
                        drawing_xml = str(drawing.xml)
                        # 计算形状数量
                        wps_count = drawing_xml.count('wps:wsp')
                        vml_count = drawing_xml.lower().count('v:shape') - drawing_xml.lower().count('textpath')
                        # 如果包含多个形状，可能是流程图
                        if wps_count >= 2 or (vml_count >= 2 and vml_count > 0):
                            return True
            except:
                pass
            
            # 方法5: 检测段落中的 runs 是否包含形状
            try:
                shape_count = 0
                for run in paragraph.runs:
                    if not hasattr(run, 'element'):
                        continue
                    run_xml = str(run.element.xml)
                    # 排除水印
                    if 'v:shape' in run_xml.lower() and 'textpath' in run_xml.lower():
                        continue
                    # 检查是否包含形状
                    if 'wps:wsp' in run_xml or ('v:shape' in run_xml.lower() and 'textpath' not in run_xml.lower()):
                        shape_count += 1
                # 如果包含多个形状，可能是流程图
                if shape_count >= 2:
                    return True
            except:
                pass
            
        except:
            pass
        
        return False

    def _apply_page_settings(self, document: Document) -> None:
        """应用页面设置（页边距等），但不修改任何页边距，保持文档原始页边距"""
        # 不修改任何 section 的页边距，保持文档原始页边距
        # 封面页和后续页面的页边距都不修改
        pass
    
    def _check_header(self, document: Document) -> list:
        """检测页眉是否存在，如果不存在则提示添加"""
        issues = []
        
        # 检查所有节的页眉
        has_header = False
        for section in document.sections:
            header = section.header
            # 检查页眉是否有内容
            for para in header.paragraphs:
                if para.text and para.text.strip():
                    has_header = True
                    break
            if has_header:
                break
        
        # 如果没有任何页眉，添加提示
        if not has_header:
            issues.append({
                "type": "missing_header",
                "severity": "warning",
                "message": "文档缺少页眉",
                "suggestion": f"请在文档中添加页眉，建议内容：{HEADER_SETTINGS['text']}"
            })
        
        return issues
    
    def _convert_university_params_to_rules(self, university_params: Dict) -> Dict[str, Dict]:
        """
        将预设大学模板的参数转换为格式规则
        
        Args:
            university_params: 大学模板参数字典，包含 body_text, page_settings 等
            
        Returns:
            格式规则字典，格式与 FONT_STANDARDS 相同
        """
        rules = {}
        
        # 复制标准规则作为基础
        for style_name, style_config in FONT_STANDARDS.items():
            rules[style_name] = style_config.copy()
        
        # 应用预设模板的参数覆盖
        # 主要覆盖 body_text 的行距等参数
        if "body_text" in university_params:
            body_params = university_params["body_text"]
            if "body_text" in rules:
                # 覆盖 body_text 的参数
                for key, value in body_params.items():
                    rules["body_text"][key] = value
        
        # 如果预设模板有其他样式参数，也可以覆盖
        for style_name, style_params in university_params.items():
            if style_name != "body_text" and style_name != "page_settings":
                if style_name in rules:
                    for key, value in style_params.items():
                        rules[style_name][key] = value
        
        return rules
    
    def _merge_rules_with_standard(self, template_rules: Dict[str, Dict]) -> Dict[str, Dict]:
        """
        合并模板规则和标准规则
        优先级：标准规则 > 模板规则
        如果模板规则中有标准规则没有的样式，保留模板规则
        """
        merged = {}
        
        # 首先添加标准规则
        for style_name, style_config in FONT_STANDARDS.items():
            merged[style_name] = style_config.copy()
        
        # 然后添加模板规则（如果模板规则中的样式名不在标准中，则添加）
        for style_name, style_config in template_rules.items():
            # 如果模板样式名不在标准中，保留模板样式
            if style_name not in merged:
                merged[style_name] = style_config.copy()
            else:
                # 如果模板样式在标准中，但标准中没有某些字段，则补充模板的字段
                standard_style = merged[style_name]
                for key, value in style_config.items():
                    if key not in standard_style or standard_style[key] is None:
                        standard_style[key] = value
        
        return merged
    
    def _detect_paragraph_style(self, paragraph: Paragraph) -> str:
        """
        根据段落内容自动检测应该应用的样式
        返回样式名称（对应FONT_STANDARDS中的key）
        """
        text = paragraph.text.strip() if paragraph.text else ""
        if not text:
            return DEFAULT_STYLE
        
        # 优先检测特殊标题：摘要、ABSTRACT、目录、绪论、概述
        # 这些标题需要设置为黑体、三号字、加粗、居中
        if text == "摘要" or text.startswith("摘要"):
            return "abstract_title"
        if text == "ABSTRACT" or text.startswith("ABSTRACT"):
            return "abstract_title_en"
        if text == "目录" or text.startswith("目录") or text == "Contents" or text.startswith("Contents"):
            return "toc_title"
        if text == "绪论" or text == "概述" or text.startswith("1 绪论") or text.startswith("1 概述"):
            # 如果是独立的"绪论"或"概述"，且段落较短，则认为是标题
            if len(text) < 50:
                return "title_level_1"
        
        # 根据样式映射规则检测
        for rule in STYLE_MAPPING_RULES:
            if re.match(rule["pattern"], text, re.IGNORECASE):
                return rule["style"]
        
        # 检查是否是标题
        style_name = paragraph.style.name if paragraph.style else None
        if style_name:
            style_lower = style_name.lower()
            if "标题" in style_name or "heading" in style_lower:
                # 根据标题级别判断
                if "1" in style_name or "一" in style_name or "heading 1" in style_lower:
                    return "title_level_1"
                elif "2" in style_name or "二" in style_name or "heading 2" in style_lower:
                    return "title_level_2"
                elif "3" in style_name or "三" in style_name or "heading 3" in style_lower:
                    return "title_level_3"
        
        # 检查段落内容特征
        if text.startswith("图") and len(text) < 100:
            return "figure_caption"
        if text.startswith("表") and len(text) < 100:
            return "table_caption"
        
        # 章节标题检测：必须是独立的、较短的段落
        # 避免将正文中的"第二章的方案"等误识别为标题
        # 标题一般不会超过一行，字数不会超过30个
        chapter_match = re.match(r"^(第[一二三四五六七八九十\d]+章|第\d+章|Chapter\s+\d+)([，,。.：:；;]?)$", text)
        if chapter_match:
            # 如果匹配到章节标题，且段落较短（标题通常是独立的短段落，不超过30个字符）
            # 或者后面只有标点符号，则认为是标题
            # 换行以后就是新的内容了，标题一般不会超过一行
            remaining_text = text[len(chapter_match.group(0)):].strip()
            if len(text) <= 30 and (len(remaining_text) == 0 or remaining_text in ["，", "。", "：", "；", ",", ".", ":", ";"]):
                return "title_level_1"
        
        # 二级标题检测：格式为 数字.数字 或 数字.数字 后跟文字内容
        # 例如：2.1、3.1、4.1 或 2.1 系统设计、3.1 需求分析 等
        # 标题一般不会超过一行，字数不会超过50个
        section_match = re.match(r"^(\d+\.\d+)(\s*[，,。.：:；;]?\s*)(.*)$", text)
        if section_match:
            # 匹配到 数字.数字 格式
            # 如果后面有文字内容（如"2.1 系统设计"），也是标题
            # 如果后面只有标点符号或为空，也是标题
            remaining_text = section_match.group(3).strip() if section_match.group(3) else ""
            # 标题通常不会超过50个字符，且段落较短
            if len(text) <= 50:
                return "title_level_2"
        
        # 也支持"第X节"格式的二级标题
        section_chinese_match = re.match(r"^(第[一二三四五六七八九十\d]+节)([，,。.：:；;]?)$", text)
        if section_chinese_match:
            remaining_text = text[len(section_chinese_match.group(0)):].strip()
            if len(text) <= 30 and (len(remaining_text) == 0 or remaining_text in ["，", "。", "：", "；", ",", ".", ":", ";"]):
                return "title_level_2"
        
        # 三级标题检测：必须是独立的、较短的段落
        # 标题格式：数字.数字.数字 或 数字.数字.数字 后跟标点符号，且后面没有其他文字内容
        # 标题一般不会超过一行，字数不会超过30个
        subsection_match = re.match(r"^(\d+\.\d+\.\d+)([，,。.：:；;]?)$", text)
        if subsection_match:
            remaining_text = text[len(subsection_match.group(0)):].strip()
            # 只有当剩余文本为空或只有标点符号时，且总长度不超过30个字符，才认为是标题
            # 如果后面还有文字内容（如"3.2.4 12864 液晶显示屏"），则不是标题，是正文
            # 换行以后就是新的内容了，标题一般不会超过一行
            if len(text) <= 30 and (len(remaining_text) == 0 or remaining_text in ["，", "。", "：", "；", ",", ".", ":", ";"]):
                return "title_level_3"
        
        # 默认返回正文样式
        return DEFAULT_STYLE

    def _find_cover_end_index(self, document: Document) -> int:
        """找到封面结束的段落索引，跳过封面部分"""
        # 封面的结束标志：通常是"摘要"、"目录"、"引言"、"第一章"等
        cover_end_keywords = [
            "摘要", "ABSTRACT", "目录", "Contents", 
            "引言", "绪论", "前言", "第一章", "第1章", "Chapter 1",
            "1 引言", "1 绪论", "1 概述"
        ]
        
        # 从前往后查找，找到第一个封面结束标志
        for idx, paragraph in enumerate(document.paragraphs):
            para_text = paragraph.text.strip() if paragraph.text else ""
            if not para_text:
                continue
            
            # 检查是否是封面结束标志
            for keyword in cover_end_keywords:
                if para_text.startswith(keyword) or keyword in para_text:
                    # 确保不是封面中的文字（封面通常较短）
                    if len(para_text) < 200:  # 封面中的文字通常较短
                        return idx
        
        # 如果找不到，跳过前20个段落（通常是封面）
        return min(20, len(document.paragraphs) - 1)
    
    def _find_section_ranges(self, document: Document) -> Dict[str, Tuple[int, int]]:
        """
        识别文档各个部分的段落范围
        返回: {
            "cover": (0, cover_end),  # 封面（第一页）
            "integrity": (start, end),  # 诚信承诺（第二页）
            "abstract_zh": (start, end),  # 中文摘要（第三页）
            "abstract_en": (start, end),  # 英文摘要
            "toc": (start, end),  # 目录
            "body": (start, end),  # 正文
        }
        """
        ranges = {}
        cover_end = self._find_cover_end_index(document)
        ranges["cover"] = (0, cover_end)
        
        integrity_start = None
        integrity_end = None
        abstract_zh_start = None
        abstract_zh_end = None
        abstract_en_start = None
        abstract_en_end = None
        toc_start = None
        toc_end = None
        body_start = None
        
        # 查找诚信承诺（通常在封面之后，摘要之前，第二页）
        # 支持"诚信承诺"中间有空格的情况，如"诚信 承诺"、"诚 信 承 诺"等
        # 注意：诚信承诺可能在封面范围内，所以从段落0开始查找
        integrity_pattern = re.compile(r'诚\s*信\s*承\s*诺', re.IGNORECASE)
        integrity_keywords = ["学术诚信", "原创性声明", "原创声明"]
        
        self._log_to_file(f"[修复] 开始查找诚信承诺，从段落 0 开始（cover_end={cover_end}）")
        # 先在前50个段落中查找（通常诚信承诺在前几页）
        search_range = min(50, len(document.paragraphs))
        
        # 改进的查找逻辑：四个字可以分散在不同段落，中间可以有任意空格
        # 先找"诚"，再找"信"，再找"承"，再找"诺"，按顺序出现即可
        cheng_idx = None  # 诚
        xin_idx = None    # 信
        cheng2_idx = None # 承
        nuo_idx = None    # 诺
        
        for idx in range(0, search_range):
            para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
            if not para_text:
                continue
            
            # 调试：输出前20个段落的文本（用于排查）
            if idx < 20 and integrity_start is None:
                self._log_to_file(f"[修复] 段落 {idx} 文本预览: {para_text[:80]}")
            
            # 第一步：找"诚"
            if cheng_idx is None and '诚' in para_text:
                cheng_idx = idx
                self._log_to_file(f"[修复] 步骤1: 找到'诚'，段落索引: {idx}, 文本: {para_text[:80]}")
                # 如果四个字都在同一个段落，一次性检查
                if '信' in para_text and '承' in para_text and '诺' in para_text:
                    xin_idx = idx
                    cheng2_idx = idx
                    nuo_idx = idx
                    integrity_start = idx
                    self._log_to_file(f"[修复] ✅ 在同一段落找到完整的'诚信承诺'，段落索引: {idx}")
                    break
                # 否则继续查找其他字
                continue
            
            # 第二步：找到"诚"后，找"信"（可以在同一段落或之后）
            if cheng_idx is not None and xin_idx is None and idx >= cheng_idx and '信' in para_text:
                xin_idx = idx
                self._log_to_file(f"[修复] 步骤2: 找到'信'，段落索引: {idx}, 文本: {para_text[:80]}")
                # 如果"承"和"诺"也在同一段落，一次性检查
                if '承' in para_text and '诺' in para_text:
                    cheng2_idx = idx
                    nuo_idx = idx
                    integrity_start = cheng_idx
                    self._log_to_file(f"[修复] ✅ 找到完整的'诚信承诺'，起始段落索引: {integrity_start}")
                    break
                continue
            
            # 第三步：找到"信"后，找"承"（可以在同一段落或之后）
            if xin_idx is not None and cheng2_idx is None and idx >= xin_idx and '承' in para_text:
                cheng2_idx = idx
                self._log_to_file(f"[修复] 步骤3: 找到'承'，段落索引: {idx}, 文本: {para_text[:80]}")
                # 如果"诺"也在同一段落，一次性检查
                if '诺' in para_text:
                    nuo_idx = idx
                    integrity_start = cheng_idx
                    self._log_to_file(f"[修复] ✅ 找到完整的'诚信承诺'，起始段落索引: {integrity_start}")
                    break
                continue
            
            # 第四步：找到"承"后，找"诺"（可以在同一段落或之后）
            if cheng2_idx is not None and nuo_idx is None and idx >= cheng2_idx and '诺' in para_text:
                nuo_idx = idx
                self._log_to_file(f"[修复] 步骤4: 找到'诺'，段落索引: {idx}, 文本: {para_text[:80]}")
                # 找到所有四个字，设置诚信承诺的起始位置为"诚"所在的段落
                integrity_start = cheng_idx
                self._log_to_file(f"[修复] ✅ 找到完整的'诚信承诺'，起始段落索引: {integrity_start}")
                break
            
            # 也检查其他诚信承诺相关关键词（作为备选方案）
            if integrity_start is None:
                found_keyword = False
                for keyword in integrity_keywords:
                    if keyword in para_text:
                        integrity_start = idx
                        self._log_to_file(f"[修复] ✅ 找到诚信承诺（关键词匹配：{keyword}），段落索引: {idx}, 文本: {para_text[:80]}")
                        found_keyword = True
                        break
                if found_keyword:
                    break
        
        # 如果通过四个字分别查找的方式找到了，但还没有设置 integrity_start
        if cheng_idx is not None and xin_idx is not None and cheng2_idx is not None and nuo_idx is not None and integrity_start is None:
            integrity_start = cheng_idx
            self._log_to_file(f"[修复] ✅ 通过分步查找找到完整的'诚信承诺'，起始段落索引: {integrity_start}")
        
        # 诚信承诺的结束标志：遇到"摘要"或"ABSTRACT"时结束
        # 注意：诚信承诺应该在独立的一页，所以遇到"摘要"就应该结束
        if integrity_start is not None:
            # 从诚信承诺开始位置之后查找摘要
            abstract_pattern = re.compile(r'^摘\s*要', re.IGNORECASE)
            for idx in range(integrity_start + 1, min(integrity_start + 30, len(document.paragraphs))):
                para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
                if not para_text:
                    continue
                
                # 检查是否是摘要开始（支持"摘要"中间有空格，或"ABSTRACT"大小写不敏感）
                abstract_en_pattern = re.compile(r'^abstract', re.IGNORECASE)
                if abstract_pattern.match(para_text) or abstract_en_pattern.match(para_text):
                    # 检查摘要前是否有分页符，如果有，说明诚信承诺和摘要已经分开
                    # 如果没有分页符，但摘要标题前有分页符，也认为已经分开
                    if idx > 0:
                        prev_para = document.paragraphs[idx - 1]
                        if prev_para.paragraph_format.page_break_before:
                            integrity_end = idx
                            break
                        # 检查前一个段落是否有分页符
                        for run in prev_para.runs:
                            if hasattr(run, 'element'):
                                run_xml = str(run.element.xml)
                                if 'w:br' in run_xml and 'type="page"' in run_xml:
                                    integrity_end = idx
                                    break
                        if integrity_end is not None:
                            break
                    # 如果摘要标题本身有分页符，也认为已经分开
                    abstract_para = document.paragraphs[idx]
                    if abstract_para.paragraph_format.page_break_before:
                        integrity_end = idx
                        break
                    # 检查摘要标题的runs中是否有分页符
                    for run in abstract_para.runs:
                        if hasattr(run, 'element'):
                            run_xml = str(run.element.xml)
                            if 'w:br' in run_xml and 'type="page"' in run_xml:
                                integrity_end = idx
                                break
                    if integrity_end is not None:
                        break
                    # 如果没有分页符，但已经找到摘要标题，也结束诚信承诺（避免合并）
                    integrity_end = idx
                    break
        
        # 如果找到了诚信承诺，但没找到结束标志，假设到摘要之前
        if integrity_start is not None and integrity_end is None:
            abstract_pattern = re.compile(r'^摘\s*要', re.IGNORECASE)
            for idx in range(integrity_start + 1, len(document.paragraphs)):
                para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
                abstract_en_pattern = re.compile(r'^abstract', re.IGNORECASE)
                if abstract_pattern.match(para_text) or abstract_en_pattern.match(para_text):
                    integrity_end = idx
                    break
        
        # 确定查找后续部分的起始位置
        search_start = integrity_end if integrity_end is not None else cover_end
        
        # 查找中文摘要（支持"摘要"中间有空格）
        abstract_pattern = re.compile(r'^摘\s*要', re.IGNORECASE)
        self._log_to_file(f"[修复] 开始查找中文摘要，从段落 {search_start} 开始")
        for idx in range(search_start, len(document.paragraphs)):
            para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
            if abstract_pattern.match(para_text) and abstract_zh_start is None:
                abstract_zh_start = idx
                self._log_to_file(f"[修复] ✅ 找到中文摘要，段落索引: {idx}, 文本: {para_text[:50]}")
            elif abstract_zh_start is not None:
                # 检查是否是关键词、ABSTRACT（大小写不敏感）或目录
                abstract_en_pattern = re.compile(r'^abstract', re.IGNORECASE)
                if para_text.startswith("关键词") or abstract_en_pattern.match(para_text) or para_text.startswith("目录"):
                    abstract_zh_end = idx
                    break
        
        # 如果没找到结束标志，假设摘要到"ABSTRACT"（大小写不敏感）或"目录"之前
        if abstract_zh_start is not None and abstract_zh_end is None:
            abstract_en_pattern = re.compile(r'^abstract', re.IGNORECASE)
            for idx in range(abstract_zh_start + 1, len(document.paragraphs)):
                para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
                if abstract_en_pattern.match(para_text) or para_text.startswith("目录"):
                    abstract_zh_end = idx
                    break
        
        # 查找英文摘要（支持大小写不敏感，如 "Abstract", "ABSTRACT", "abstract"）
        self._log_to_file(f"[修复] 开始查找英文摘要，从段落 {search_start} 开始")
        abstract_en_pattern = re.compile(r'^abstract', re.IGNORECASE)
        for idx in range(search_start, len(document.paragraphs)):
            para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
            # 检查是否是英文摘要标题（大小写不敏感）
            if abstract_en_pattern.match(para_text) and abstract_en_start is None:
                abstract_en_start = idx
                self._log_to_file(f"[修复] ✅ 找到英文摘要，段落索引: {idx}, 文本: {para_text[:50]}")
            elif abstract_en_start is not None:
                # 检查是否是英文摘要结束标志：Keywords/Key words/目录/Contents/第一章等
                # 支持 "Keywords"、"Key words"、"Key words:" 等多种格式
                is_end_marker = (
                    para_text.startswith("Keywords") or 
                    para_text.startswith("Key words") or 
                    para_text.startswith("Key Words") or
                    para_text.startswith("目录") or 
                    para_text.startswith("Contents") or 
                    para_text.startswith("第一章") or 
                    para_text.startswith("第1章")
                )
                if is_end_marker:
                    # 找到结束标志，但需要找到"Key words"或"Keywords"之后的内容结束位置
                    # 继续查找，直到找到目录或正文开始
                    abstract_en_end = idx
                    self._log_to_file(f"[修复] 找到英文摘要结束标志，段落索引: {idx}, 文本: {para_text[:50]}")
                    # 继续查找，找到"Key words"或"Keywords"之后的内容结束位置
                    # 如果后面是目录或正文，则英文摘要结束
                    for next_idx in range(idx + 1, min(idx + 10, len(document.paragraphs))):
                        next_para_text = document.paragraphs[next_idx].text.strip() if document.paragraphs[next_idx].text else ""
                        if next_para_text.startswith("目录") or next_para_text.startswith("Contents") or next_para_text.startswith("第一章") or next_para_text.startswith("第1章"):
                            abstract_en_end = next_idx
                            self._log_to_file(f"[修复] 英文摘要结束位置: {next_idx}, 文本: {next_para_text[:50]}")
                            break
                    break
        
        # 查找目录
        for idx in range(search_start, len(document.paragraphs)):
            para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
            if (para_text.startswith("目录") or para_text.startswith("Contents")) and toc_start is None:
                toc_start = idx
            elif toc_start is not None and (para_text.startswith("第一章") or para_text.startswith("第1章") or para_text.startswith("Chapter 1") or para_text.startswith("1 引言") or para_text.startswith("1 绪论")):
                toc_end = idx
                break
        
        # 查找正文开始（从"绪论"或"概述"开始）
        for idx in range(search_start, len(document.paragraphs)):
            para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
            if (para_text.startswith("第一章") or para_text.startswith("第1章") or para_text.startswith("Chapter 1") or 
                para_text.startswith("1 引言") or para_text.startswith("1 绪论") or para_text.startswith("1 概述") or
                para_text == "绪论" or para_text == "概述" or para_text.startswith("绪论") or para_text.startswith("概述")):
                body_start = idx
                break
        
        if integrity_start is not None:
            ranges["integrity"] = (integrity_start, integrity_end if integrity_end else (abstract_zh_start if abstract_zh_start else len(document.paragraphs)))
            self._log_to_file(f"[修复] 设置 integrity 范围: {ranges['integrity']}")
        else:
            self._log_to_file(f"[修复] ⚠️ 未找到诚信承诺（integrity_start is None）")
        if abstract_zh_start is not None:
            ranges["abstract_zh"] = (abstract_zh_start, abstract_zh_end if abstract_zh_end else (abstract_en_start if abstract_en_start else (toc_start if toc_start else len(document.paragraphs))))
            self._log_to_file(f"[修复] 设置 abstract_zh 范围: {ranges['abstract_zh']}")
        else:
            self._log_to_file(f"[修复] ⚠️ 未找到中文摘要（abstract_zh_start is None）")
        if abstract_en_start is not None:
            # 如果找到了英文摘要但没找到结束位置，尝试查找"Key words"或"Keywords"之后的内容
            if abstract_en_end is None:
                # 从英文摘要开始位置之后查找"Key words"或"Keywords"
                for idx in range(abstract_en_start + 1, len(document.paragraphs)):
                    para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
                    if (para_text.startswith("Keywords") or para_text.startswith("Key words") or 
                        para_text.startswith("Key Words") or para_text.startswith("目录") or 
                        para_text.startswith("Contents") or para_text.startswith("第一章") or 
                        para_text.startswith("第1章")):
                        # 找到结束标志，继续查找后面的内容结束位置
                        abstract_en_end = idx
                        # 继续查找，直到找到目录或正文
                        for next_idx in range(idx + 1, min(idx + 20, len(document.paragraphs))):
                            next_para_text = document.paragraphs[next_idx].text.strip() if document.paragraphs[next_idx].text else ""
                            if next_para_text.startswith("目录") or next_para_text.startswith("Contents") or next_para_text.startswith("第一章") or next_para_text.startswith("第1章"):
                                abstract_en_end = next_idx
                                break
                        break
            ranges["abstract_en"] = (abstract_en_start, abstract_en_end if abstract_en_end else (toc_start if toc_start else (body_start if body_start else len(document.paragraphs))))
            self._log_to_file(f"[修复] 设置 abstract_en 范围: {ranges['abstract_en']}")
        if toc_start is not None:
            ranges["toc"] = (toc_start, toc_end if toc_end else (body_start if body_start else len(document.paragraphs)))
        if body_start is not None:
            ranges["body"] = (body_start, len(document.paragraphs))
        
        return ranges

    def _apply_rules(
        self,
        document: Document,
        rules: Dict[str, Dict],
        default_style: str | None,
    ) -> Tuple[Document, Dict]:
        total_paragraphs = len(document.paragraphs)
        adjusted_paragraphs = 0
        used_styles: set[str] = set()
        changes_log = []  # 记录详细修改日志

        default_rule = rules.get(default_style) if default_style else None
        
        # 找到封面结束位置，跳过封面部分
        cover_end_idx = self._find_cover_end_index(document)
        
        # 识别各个部分的段落范围
        section_ranges = self._find_section_ranges(document)

        for idx, paragraph in enumerate(document.paragraphs):
            # 跳过封面部分，不修改封面内容
            if idx < cover_end_idx:
                continue
            
            # 判断当前段落属于哪个部分
            current_section = None
            if "integrity" in section_ranges:
                start, end = section_ranges["integrity"]
                if start <= idx < end:
                    current_section = "integrity"
            if current_section is None and "abstract_zh" in section_ranges:
                start, end = section_ranges["abstract_zh"]
                if start <= idx < end:
                    current_section = "abstract_zh"
            if current_section is None and "abstract_en" in section_ranges:
                start, end = section_ranges["abstract_en"]
                if start <= idx < end:
                    current_section = "abstract_en"
            if current_section is None and "toc" in section_ranges:
                start, end = section_ranges["toc"]
                if start <= idx < end:
                    current_section = "toc"
            if current_section is None and "body" in section_ranges:
                start, end = section_ranges["body"]
                if start <= idx < end:
                    current_section = "body"
            
            # 跳过诚信承诺部分，不修改任何内容（只检查有无即可）
            if current_section == "integrity":
                continue
            
            style_name = paragraph.style.name if paragraph.style else None
            rule = None
            applied_rule_name = None
            paragraph_text = paragraph.text.strip() if paragraph.text else ""
            
            # 根据当前部分应用特定格式规则
            # 处理中文摘要部分
            if current_section == "abstract_zh":
                # 摘要标题（支持"摘"和"要"中间有空格，如"摘 要"、"摘  要"等）
                abstract_pattern = re.compile(r'^摘\s*要', re.IGNORECASE)
                if abstract_pattern.match(paragraph_text):
                    if "abstract_title" in rules:
                        rule = rules["abstract_title"].copy()
                        applied_rule_name = "abstract_title"
                    else:
                        rule = FONT_STANDARDS.get("abstract_title", {}).copy()
                        applied_rule_name = "abstract_title"
                    # 强制确保摘要标题：黑体三号（16pt）、加粗、居中
                    rule["font_name"] = "黑体"
                    rule["font_size"] = 16  # 三号字
                    rule["bold"] = True
                    rule["alignment"] = "center"
                # 关键词标签
                elif paragraph_text.startswith("关键词"):
                    if "keywords_label" in rules:
                        rule = rules["keywords_label"].copy()
                        applied_rule_name = "keywords_label"
                    else:
                        rule = FONT_STANDARDS.get("keywords_label", {}).copy()
                        applied_rule_name = "keywords_label"
                # 摘要正文内容
                else:
                    if "abstract_content" in rules:
                        rule = rules["abstract_content"].copy()
                        applied_rule_name = "abstract_content"
                    else:
                        rule = FONT_STANDARDS.get("abstract_content", {}).copy()
                        applied_rule_name = "abstract_content"
                    # 确保摘要正文：宋体小四（12pt），行距20磅
                    rule["font_name"] = "宋体"
                    rule["font_size"] = 12
                    rule["line_spacing"] = 20
            
            # 处理英文摘要部分
            elif current_section == "abstract_en":
                # 英文摘要标题（支持大小写不敏感，如"Abstract"、"ABSTRACT"、"abstract"）
                abstract_en_pattern = re.compile(r'^abstract', re.IGNORECASE)
                if abstract_en_pattern.match(paragraph_text):
                    if "abstract_title_en" in rules:
                        rule = rules["abstract_title_en"].copy()
                        applied_rule_name = "abstract_title_en"
                    else:
                        rule = FONT_STANDARDS.get("abstract_title_en", {}).copy()
                        applied_rule_name = "abstract_title_en"
                    # 强制确保ABSTRACT标题：黑体三号（16pt）、加粗、居中
                    rule["font_name"] = "黑体"
                    rule["font_size"] = 16  # 三号字
                    rule["bold"] = True
                    rule["alignment"] = "center"
                # 关键词标签
                elif paragraph_text.startswith("Keywords") or paragraph_text.startswith("Key words"):
                    if "keywords_label_en" in rules:
                        rule = rules["keywords_label_en"].copy()
                        applied_rule_name = "keywords_label_en"
                    else:
                        rule = FONT_STANDARDS.get("keywords_label_en", {}).copy()
                        applied_rule_name = "keywords_label_en"
                # 英文摘要正文内容
                else:
                    if "abstract_content_en" in rules:
                        rule = rules["abstract_content_en"].copy()
                        applied_rule_name = "abstract_content_en"
                    else:
                        rule = FONT_STANDARDS.get("abstract_content_en", {}).copy()
                        applied_rule_name = "abstract_content_en"
                    # 确保英文摘要正文：Times New Roman小四（12pt）
                    rule["font_name"] = "Times New Roman"
                    rule["font_size"] = 12
            
            # 处理目录部分
            elif current_section == "toc":
                # 目录标题（支持中间最多5个空格的变体，如"目 录"、"目  录"、"目    录"等）
                # 检查是否包含"目"和"录"（允许中间最多5个空格）
                is_toc_title_para = False
                if "目" in paragraph_text and "录" in paragraph_text:
                    # 去除空格和标点后检查是否等于"目录"
                    cleaned_toc_text = re.sub(r'[\s\u3000：:，,。.；;！!？?、]', '', paragraph_text)
                    if cleaned_toc_text == "目录":
                        # 检查"目"和"录"之间的字符是否只有空格（最多5个）
                        mu_pos = paragraph_text.find("目")
                        lu_pos = paragraph_text.find("录")
                        if mu_pos >= 0 and lu_pos > mu_pos:
                            between_text = paragraph_text[mu_pos + 1:lu_pos]
                            # 如果中间只有空格（最多5个），或者是空字符串，认为是目录标题
                            if len(between_text) <= 5 and all(c in ' \t\u3000' for c in between_text):
                                is_toc_title_para = True
                            elif len(between_text) == 0:
                                is_toc_title_para = True
                elif paragraph_text.startswith("Contents") or paragraph_text.startswith("contents"):
                    cleaned_toc_text = re.sub(r'[\s\u3000：:，,。.；;！!？?、]', '', paragraph_text).upper()
                    if cleaned_toc_text == "CONTENTS":
                        is_toc_title_para = True
                
                if is_toc_title_para:
                    if "toc_title" in rules:
                        rule = rules["toc_title"].copy()
                        applied_rule_name = "toc_title"
                    else:
                        rule = FONT_STANDARDS.get("toc_title", {}).copy()
                        applied_rule_name = "toc_title"
                    # 强制确保目录标题：黑体三号（16pt）、加粗、居中
                    rule["font_name"] = "黑体"
                    rule["font_size"] = 16  # 三号字
                    rule["bold"] = True
                    rule["alignment"] = "center"
                # 目录内容
                else:
                    if "toc_content" in rules:
                        rule = rules["toc_content"].copy()
                        applied_rule_name = "toc_content"
                    else:
                        rule = FONT_STANDARDS.get("toc_content", {}).copy()
                        applied_rule_name = "toc_content"
                    # 确保目录内容：宋体小四（12pt），行距20磅
                    rule["font_name"] = "宋体"
                    rule["font_size"] = 12
                    rule["line_spacing"] = 20
            
            # 处理正文部分（使用原有逻辑）
            else:
                # 优先使用标准格式检测
                detected_style = self._detect_paragraph_style(paragraph)
                if detected_style in rules:
                    rule = rules[detected_style].copy()
                    applied_rule_name = detected_style
                # 如果标准格式中没有，尝试使用模板中的样式名
                elif style_name and style_name in rules:
                    rule = rules[style_name].copy()
                    applied_rule_name = style_name
                # 如果都没有，使用默认规则
                elif default_rule:
                    rule = default_rule.copy()
                    applied_rule_name = default_style or "默认样式"
                
                # 如果仍然没有规则，使用标准默认样式
                if not rule:
                    if DEFAULT_STYLE in rules:
                        rule = rules[DEFAULT_STYLE].copy()
                        applied_rule_name = DEFAULT_STYLE
                    elif default_rule:
                        rule = default_rule.copy()
                        applied_rule_name = default_style or "默认样式"
            
            # 强制统一正文段落格式：毕业论文正文固定为小四（12pt）宋体，固定行距20磅
            if rule:
                paragraph_text = paragraph.text.strip() if paragraph.text else ""
                # 判断是否是标题（使用更严格的判断，避免把正文误判为标题）
                is_heading = False
                if applied_rule_name:
                    # 如果应用的规则是标题样式，则认为是标题
                    if applied_rule_name in ["title_level_1", "title_level_2", "title_level_3", "abstract_title", "toc_title", "reference_title", "acknowledgment_title", "abstract_title_en"]:
                        is_heading = True
                        if idx < 10:  # 只记录前10个段落的详细信息
                            print(f"[格式应用] 段落 {idx} 被识别为标题（规则: {applied_rule_name}）")
                    # 或者检查样式名称（但更严格）
                    elif style_name and ("标题" in style_name.lower() or "heading" in style_name.lower()):
                        # 只有当段落很短（<=30字符）且居中对齐时，才认为是标题
                        if len(paragraph_text) <= 30 and paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                            is_heading = True
                            if idx < 10:
                                print(f"[格式应用] 段落 {idx} 被识别为标题（样式: {style_name}，内容: {paragraph_text[:20]}）")
                    # 或者检查段落内容特征（居中对齐的短文本，或"绪论"、"概述"等）
                    elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and len(paragraph_text) < 30:
                        # 更严格：只有非常短的文本（<=20字符）且居中对齐才认为是标题
                        if len(paragraph_text) <= 20:
                            is_heading = True
                            if idx < 10:
                                print(f"[格式应用] 段落 {idx} 被识别为标题（居中短文本: {paragraph_text[:20]}）")
                    # 或者检查是否是"绪论"、"概述"等标题
                    elif paragraph_text == "绪论" or paragraph_text == "概述" or paragraph_text.startswith("1 绪论") or paragraph_text.startswith("1 概述"):
                        if len(paragraph_text) <= 20:  # 更严格：只有很短的文本才认为是标题
                            is_heading = True
                            if idx < 10:
                                print(f"[格式应用] 段落 {idx} 被识别为标题（绪论/概述: {paragraph_text}）")
                    # 或者检查是否以数字开头且较短（标题一般不会超过一行，字数不会超过50个）
                    elif paragraph_text and paragraph_text[0].isdigit() and len(paragraph_text) <= 50:
                        # 二级标题格式：数字.数字 或 数字.数字 后跟文字（如"2.1 系统设计"）
                        if re.match(r'^(\d+\.\d+)(\s*[，,。.：:；;]?\s*)(.*)$', paragraph_text):
                            is_heading = True
                            if idx < 10:
                                print(f"[格式应用] 段落 {idx} 被识别为二级标题（数字编号: {paragraph_text}）")
                        # 三级标题格式：数字.数字.数字
                        elif re.match(r'^(\d+\.\d+\.\d+)([，,。.：:；;]?)$', paragraph_text):
                            is_heading = True
                            if idx < 10:
                                print(f"[格式应用] 段落 {idx} 被识别为标题（数字编号: {paragraph_text}）")
                # 如果没有应用规则名称，使用备用判断逻辑（更严格）
                if not is_heading:
                    is_heading = (
                        (style_name and ("标题" in style_name.lower() or "heading" in style_name.lower()) and len(paragraph_text) <= 20 and paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER) or
                        (paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and len(paragraph_text) <= 15) or  # 更严格：<=15字符
                        # 更严格的判断：只有纯数字编号格式才认为是标题（标题一般不会超过一行，字数不会超过15个）
                        (paragraph_text and paragraph_text[0].isdigit() and len(paragraph_text) <= 15 and 
                         re.match(r'^(\d+\.\d+\.\d+|\d+\.\d+|\d+)([，,。.：:；;]?)$', paragraph_text)) or
                        ((paragraph_text == "绪论" or paragraph_text == "概述" or paragraph_text.startswith("1 绪论") or paragraph_text.startswith("1 概述")) and len(paragraph_text) <= 20)
                    )
                
                # 判断是否包含图片、公式或流程图
                has_image_or_equation = self._paragraph_has_image_or_equation(paragraph)
                has_flowchart = self._paragraph_has_flowchart(paragraph)
                
                # 判断是否是图题（图片说明）
                is_figure_caption = False
                if paragraph_text and len(paragraph_text) < 100:
                    # 检查是否以"图"开头，且包含数字（如"图1-1"、"图2.1"等）
                    if (paragraph_text.startswith("图") and 
                        (re.search(r'图\s*\d+[\.\-]\d+', paragraph_text) or re.search(r'图\s*\d+', paragraph_text))):
                        is_figure_caption = True
                    # 检查是否是流程图标题（流程图X-X、流程图X.X等）
                    elif (paragraph_text.startswith("流程图") and 
                          (re.search(r'流程图\s*\d+[\.\-]\d+', paragraph_text) or re.search(r'流程图\s*\d+', paragraph_text))):
                        is_figure_caption = True
                
                # 对于标题，移除行距设置，保持标题的原始行距
                if is_heading:
                    rule.pop("line_spacing", None)
                
                # 对于包含图片、公式或流程图的段落，移除行距设置，避免被压缩看不见，并强制居中
                if has_image_or_equation or has_flowchart:
                    # 移除行距设置，保持图片/流程图段落的原始行距
                    rule.pop("line_spacing", None)
                    # 也移除首行缩进，图片/流程图段落通常不需要缩进
                    rule.pop("first_line_indent", None)
                    # 强制设置图片段落居中对齐
                    rule["alignment"] = "center"
                
                # 对于图题（图片说明），强制居中并应用图题格式
                if is_figure_caption:
                    # 使用图题格式标准
                    if "figure_caption" in rules:
                        rule = rules["figure_caption"].copy()
                        applied_rule_name = "figure_caption"
                    else:
                        rule = FONT_STANDARDS.get("figure_caption", {}).copy()
                        applied_rule_name = "figure_caption"
                    # 强制确保图题居中对齐
                    rule["alignment"] = "center"
                
                # 对于正文段落（非标题、非图片、非公式、非流程图），保留原有字体，不强制统一
                if not is_heading and not has_image_or_equation and not has_flowchart:
                    # 检查段落中所有 runs 的字体，如果段落内字体不一致，保留各自的字体
                    run_fonts = []
                    for run in paragraph.runs:
                        if run.text.strip():  # 只检查有文本的 run
                            # 从 run 中提取字体
                            run_font = None
                            if run.font and run.font.name:
                                run_font = run.font.name
                            else:
                                # 尝试从 XML 中提取
                                try:
                                    r_pr = run._element.get_or_add_rPr()
                                    r_fonts = r_pr.rFonts
                                    if r_fonts is not None:
                                        run_font = r_fonts.get(qn("w:eastAsia")) or r_fonts.get(qn("w:ascii"))
                                except:
                                    pass
                            if run_font:
                                run_fonts.append(run_font)
                    
                    # 如果段落中有多种字体，不设置 rule["font_name"]，保留原有字体
                    unique_fonts = set(run_fonts)
                    if len(unique_fonts) > 1:
                        # 段落中有多种字体，保留各自的字体，只统一字号和行距
                        print(f"[格式应用] 段落 {idx} 检测到多种字体: {unique_fonts}，保留各自字体")
                        # 移除字体设置，并标记为保留字体
                        rule.pop("font_name", None)
                        rule["_preserve_fonts"] = True  # 标记为保留字体
                        # 只设置字号和行距（如果规则中有）
                        if DEFAULT_STYLE in FONT_STANDARDS:
                            standard_body = FONT_STANDARDS[DEFAULT_STYLE]
                            if "font_size" not in rule:
                                rule["font_size"] = standard_body.get("font_size", 12)
                            if "line_spacing" not in rule:
                                rule["line_spacing"] = standard_body.get("line_spacing", 20)
                            if "bold" not in rule:
                                rule["bold"] = standard_body.get("bold", False)
                            if "first_line_indent" not in rule:
                                rule["first_line_indent"] = standard_body.get("first_line_indent", 24)
                    elif len(unique_fonts) == 1:
                        # 段落中只有一种字体，检查是否需要保留
                        extracted_font = list(unique_fonts)[0]
                        font_lower = extracted_font.lower()
                        # 如果是支持的字体（楷体、宋体、Times New Roman、黑体），保留
                        if ("楷" in extracted_font or "kaiti" in font_lower or "kai" in font_lower or
                            "宋" in extracted_font or "simsun" in font_lower or "song" in font_lower or
                            "times" in font_lower or "new roman" in font_lower or "tnr" in font_lower or
                            "黑" in extracted_font or "simhei" in font_lower or "hei" in font_lower):
                            # 保留原有字体
                            rule["font_name"] = extracted_font
                            print(f"[格式应用] 段落 {idx} 保留字体：{extracted_font}")
                            # 只设置字号和行距（如果规则中有）
                            if DEFAULT_STYLE in FONT_STANDARDS:
                                standard_body = FONT_STANDARDS[DEFAULT_STYLE]
                                if "font_size" not in rule:
                                    rule["font_size"] = standard_body.get("font_size", 12)
                                if "line_spacing" not in rule:
                                    rule["line_spacing"] = standard_body.get("line_spacing", 20)
                                if "bold" not in rule:
                                    rule["bold"] = standard_body.get("bold", False)
                                if "first_line_indent" not in rule:
                                    rule["first_line_indent"] = standard_body.get("first_line_indent", 24)
                        else:
                            # 不支持的字体，使用默认宋体
                            if DEFAULT_STYLE in FONT_STANDARDS:
                                standard_body = FONT_STANDARDS[DEFAULT_STYLE]
                                rule["font_name"] = standard_body.get("font_name", "宋体")
                                rule["font_size"] = standard_body.get("font_size", 12)
                                rule["line_spacing"] = standard_body.get("line_spacing", 20)
                                rule["bold"] = standard_body.get("bold", False)
                                rule["first_line_indent"] = standard_body.get("first_line_indent", 24)
                            print(f"[格式应用] 段落 {idx} 使用默认字体：宋体、12pt、行距20磅")
                    else:
                        # 没有提取到字体，使用默认宋体
                        if DEFAULT_STYLE in FONT_STANDARDS:
                            standard_body = FONT_STANDARDS[DEFAULT_STYLE]
                            rule["font_name"] = standard_body.get("font_name", "宋体")
                            rule["font_size"] = standard_body.get("font_size", 12)
                            rule["bold"] = standard_body.get("bold", False)
                            rule["line_spacing"] = standard_body.get("line_spacing", 20)
                            rule["first_line_indent"] = standard_body.get("first_line_indent", 24)
                        print(f"[格式应用] 段落 {idx} 使用默认字体：宋体、12pt、行距20磅")
                # 对于标题，根据级别设置字体
                elif is_heading:
                    # 根据标题级别设置字体
                    if applied_rule_name in ["title_level_1", "title_level_2", "title_level_3"]:
                        if applied_rule_name in FONT_STANDARDS:
                            title_style = FONT_STANDARDS[applied_rule_name]
                            # 一级标题：黑体；二级标题：宋体；三级标题：黑体
                            rule["font_name"] = title_style.get("font_name", "黑体" if applied_rule_name != "title_level_2" else "宋体")
                            rule["font_size"] = title_style.get("font_size", 16 if applied_rule_name == "title_level_1" else 12)
                            rule["bold"] = title_style.get("bold", True)
                            print(f"[格式应用] 段落 {idx} 应用标题格式：{applied_rule_name}，字体：{rule['font_name']}，字号：{rule['font_size']}pt")
                    else:
                        # 其他标题（如摘要、目录等）使用黑体
                        if rule.get("font_name") is None or "黑" not in str(rule.get("font_name", "")):
                            rule["font_name"] = "黑体"
                            print(f"[格式应用] 段落 {idx} 强制设置为标题格式：黑体")

            if rule:
                # 记录修改前的格式
                before_format = docx_format_utils.extract_paragraph_format(paragraph)
                paragraph_text = paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text
                
                # 再次确认：如果段落包含流程图，确保行距不被修改
                # 流程图视为图片，不修改行距
                if has_flowchart:
                    # 确保规则中不包含行距设置
                    rule.pop("line_spacing", None)
                    rule.pop("first_line_indent", None)
                
                # 应用规则
                docx_format_utils.apply_paragraph_rule(paragraph, rule)
                
                # 最终检查：确保"摘要"、"ABSTRACT"和"目录"标题始终居中（防止被其他逻辑覆盖）
                para_text_check = paragraph.text.strip() if paragraph.text else ""
                if para_text_check:
                    # 去除所有空格、标点符号和空白字符，只保留字母和汉字
                    cleaned_text_check = re.sub(r'[\s\u3000：:，,。.；;！!？?、]', '', para_text_check)
                    cleaned_text_check_upper = cleaned_text_check.upper()
                    
                    is_abstract_title_check = False
                    is_toc_title_check = False
                    # 检查去除空格和标点后是否等于"摘要"、"ABSTRACT"或"目录"
                    if cleaned_text_check == "摘要" or cleaned_text_check_upper == "ABSTRACT":
                        is_abstract_title_check = True
                    elif cleaned_text_check == "目录" or cleaned_text_check_upper == "CONTENTS":
                        is_toc_title_check = True
                    # 如果去除空格后的文本较短，也检查是否包含这些关键词
                    elif len(cleaned_text_check) <= 15:  # 基于去除空格后的长度
                        # 检查是否包含"摘"和"要"（允许中间有空格或其他字符）
                        if "摘" in para_text_check and "要" in para_text_check:
                            if cleaned_text_check == "摘要":
                                is_abstract_title_check = True
                        # 检查是否包含"目"和"录"（允许中间最多5个空格）
                        elif "目" in para_text_check and "录" in para_text_check:
                            # 检查"目"和"录"之间的字符是否只有空格（最多5个）
                            mu_pos = para_text_check.find("目")
                            lu_pos = para_text_check.find("录")
                            if mu_pos >= 0 and lu_pos > mu_pos:
                                between_text = para_text_check[mu_pos + 1:lu_pos]
                                # 如果中间只有空格（最多5个），或者是空字符串，认为是目录标题
                                if len(between_text) <= 5 and all(c in ' \t\u3000' for c in between_text):
                                    if cleaned_text_check == "目录":
                                        is_toc_title_check = True
                                elif len(between_text) == 0:
                                    if cleaned_text_check == "目录":
                                        is_toc_title_check = True
                        # 检查是否包含"ABSTRACT"（不区分大小写）
                        elif "ABSTRACT" in cleaned_text_check_upper or "abstract" in para_text_check.lower():
                            if cleaned_text_check_upper == "ABSTRACT":
                                is_abstract_title_check = True
                        # 检查是否包含"Contents"（不区分大小写）
                        elif "CONTENTS" in cleaned_text_check_upper or "contents" in para_text_check.lower():
                            if cleaned_text_check_upper == "CONTENTS":
                                is_toc_title_check = True
                    
                    if is_abstract_title_check or is_toc_title_check:
                        # 确保标题格式：黑体三号（16pt）、加粗、居中
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # 设置字体和字号
                        for run in paragraph.runs:
                            run.font.name = "黑体"
                            run.font.size = Pt(16)  # 三号字
                            run.font.bold = True
                        # 如果段落没有runs，创建一个run并设置格式
                        if not paragraph.runs:
                            run = paragraph.add_run()
                            run.font.name = "黑体"
                            run.font.size = Pt(16)
                            run.font.bold = True
                
                # 记录修改后的格式
                after_format = docx_format_utils.extract_paragraph_format(paragraph)
                
                # 找出实际修改的字段
                changed_fields = []
                for key in before_format:
                    before_val = before_format.get(key)
                    after_val = after_format.get(key)
                    if before_val != after_val:
                        changed_fields.append({
                            "field": key,
                            "before": before_val,
                            "after": after_val
                        })
                
                if changed_fields:
                    adjusted_paragraphs += 1
                    changes_log.append({
                        "paragraph_index": idx,
                        "paragraph_preview": paragraph_text.strip() or "(空段落)",
                        "style_name": style_name,
                        "applied_rule": applied_rule_name,
                        "changes": changed_fields
                    })
                    if style_name:
                        used_styles.add(style_name)

        # 统计修改类型
        change_summary = {}
        for change in changes_log:
            for field_change in change["changes"]:
                field_name = field_change["field"]
                if field_name not in change_summary:
                    change_summary[field_name] = 0
                change_summary[field_name] += 1

        stats = {
            "paragraphs_total": total_paragraphs,
            "paragraphs_adjusted": adjusted_paragraphs,
            "styles_applied": sorted(list(used_styles)),
            "changes_summary": change_summary,
            "changes_detail": changes_log[:50],  # 只保留前50条详细记录，避免报告过大
        }

        return document, stats

    def _find_body_start_index(self, document: Document) -> int:
        """找到正文开始的段落索引，跳过封面、目录等前置部分"""
        # 正文开始的标志关键词（按优先级排序）
        # 高优先级：明确的章节标题
        chapter_keywords = [
            "第一章", "第二章", "第三章", "第四章", "第五章", "第六章", "第七章", "第八章", "第九章", "第十章",
            "第1章", "第2章", "第3章", "第4章", "第5章", "第6章", "第7章", "第8章", "第9章", "第10章",
        ]
        
        # 中优先级：章节关键词
        section_keywords = [
            "引言", "绪论", "前言", "概述", "正文", "正文部分",
        ]
        
        # 低优先级：带编号的章节（需要更严格的匹配）
        numbered_sections = [
            "1 引言", "1 绪论", "1 概述", "1 前言",
            "1.1", "1.2", "2.1", "2.2",  # 小节编号
        ]
        
        # 方法1: 查找明确的章节标题（最高优先级）
        for idx, paragraph in enumerate(document.paragraphs):
            paragraph_text = paragraph.text.strip() if paragraph.text else ""
            if not paragraph_text:
                continue
            
            # 检查是否是明确的章节标题
            for keyword in chapter_keywords:
                if keyword in paragraph_text:
                    # 确保不是目录中的引用（目录通常较短且包含"目录"字样）
                    if "目录" not in paragraph_text:
                        # 章节标题通常较短，或者段落开头就是章节标题
                        if len(paragraph_text) < 100 or paragraph_text.startswith(keyword):
                            return idx
        
        # 方法2: 查找章节关键词（中优先级）
        for idx, paragraph in enumerate(document.paragraphs):
            paragraph_text = paragraph.text.strip() if paragraph.text else ""
            if not paragraph_text:
                continue
            
            # 检查是否是章节关键词，且段落开头包含关键词（避免匹配到正文中的引用）
            for keyword in section_keywords:
                if paragraph_text.startswith(keyword) or (keyword in paragraph_text and len(paragraph_text) > 50):
                    # 确保不是目录中的引用
                    if "目录" not in paragraph_text and len(paragraph_text) > 20:
                        return idx
        
        # 方法3: 查找带编号的章节（需要更严格的匹配）
        for idx, paragraph in enumerate(document.paragraphs):
            paragraph_text = paragraph.text.strip() if paragraph.text else ""
            if not paragraph_text:
                continue
            
            # 检查是否是带编号的章节（段落开头必须是编号）
            for keyword in numbered_sections:
                if paragraph_text.startswith(keyword):
                    # 确保不是目录中的引用
                    if "目录" not in paragraph_text and len(paragraph_text) > 20:
                        return idx
        
        # 方法4: 如果找不到关键词，跳过前N个段落（通常是封面和目录）
        # 跳过前20个段落，或者文档总段落数的10%（取较大值）
        skip_count = max(20, len(document.paragraphs) // 10)
        return min(skip_count, len(document.paragraphs) - 1)

    def _check_figure_captions(self, document: Document) -> list:
        """检测文档中的图片，检查是否有图题，返回缺失图题的图片列表
        注意：只从正文开始检测，跳过封面、目录等前置部分
        注意：不在文档中插入标记，只记录问题到issues中，保持文档干净"""
        issues = []
        missing_caption_indices = []  # 记录缺少图题的图片段落索引
        
        # 找到正文开始的段落索引
        body_start_idx = self._find_body_start_index(document)
        
        # 只从正文开始检测图片
        for idx, paragraph in enumerate(document.paragraphs):
            # 跳过正文之前的段落
            if idx < body_start_idx:
                continue
            # 检查段落中是否包含图片
            has_image = False
            paragraph_text = paragraph.text.strip() if paragraph.text else ""
            
            # 跳过明显不是图片的段落（比如纯文本段落、标题等）
            # 如果段落有大量文字且没有drawing相关标签，不太可能是图片段落
            # 但不要完全跳过，因为图片段落可能包含一些文字说明
            # 先检查是否有drawing相关标签，如果没有且文字很多，才跳过
            if len(paragraph_text) > 200:
                para_xml_preview = str(paragraph._element.xml)[:500] if hasattr(paragraph, '_element') else ""
                if 'drawing' not in para_xml_preview.lower() and 'pic:pic' not in para_xml_preview and 'a:blip' not in para_xml_preview:
                    continue
            
            # 方法1: 检查段落中的runs是否包含真正的图片（必须包含pic:pic或a:blip）
            try:
                for run in paragraph.runs:
                    if not hasattr(run, 'element'):
                        continue
                    run_xml = str(run.element.xml)
                    # 排除明显是VML形状的水印（通过检查是否有textpath等特征）
                    if 'v:shape' in run_xml.lower() and 'textpath' in run_xml.lower():
                        continue  # 这是水印，跳过
                    # 必须包含pic:pic或a:blip，这些才是真正的图片元素
                    # 同时需要验证有图片引用（r:embed或r:link）
                    if ('pic:pic' in run_xml or 'a:blip' in run_xml) and ('r:embed' in run_xml or 'r:link' in run_xml or 'a:blip' in run_xml):
                        has_image = True
                        break
            except:
                pass
            
            # 方法2: 检查段落元素中是否包含真正的图片
            if not has_image:
                try:
                    para_xml = str(paragraph._element.xml)
                    # 排除VML形状的水印
                    if 'v:shape' in para_xml.lower() and 'textpath' in para_xml.lower():
                        pass  # 这是水印，跳过
                    # 必须包含pic:pic或a:blip，且需要验证有图片引用
                    elif ('pic:pic' in para_xml or 'a:blip' in para_xml) and ('r:embed' in para_xml or 'r:link' in para_xml or 'a:blip' in para_xml):
                        has_image = True
                except:
                    pass
            
            # 方法3: 使用findall查找drawing元素，并验证包含真正的图片
            if not has_image:
                try:
                    from docx.oxml.ns import qn
                    # 查找drawing元素（使用findall配合qn，而不是xpath with namespaces）
                    drawings = paragraph._element.findall('.//' + qn('w:drawing'))
                    if drawings:
                        # 检查drawing中是否包含真正的图片（pic:pic或a:blip）
                        for drawing in drawings:
                            drawing_xml = str(drawing.xml)
                            # 排除VML形状的水印
                            if 'v:shape' in drawing_xml.lower() and 'textpath' in drawing_xml.lower():
                                continue
                            # 必须包含pic:pic或a:blip，且需要验证有图片引用
                            if ('pic:pic' in drawing_xml or 'a:blip' in drawing_xml) and ('r:embed' in drawing_xml or 'r:link' in drawing_xml or 'a:blip' in drawing_xml):
                                has_image = True
                                break
                except:
                    pass
            
            # 如果找到图片，还需要验证段落确实包含图片（不能只是文字说明）
            # 如果段落只有文字且没有图片元素，跳过
            if has_image:
                # 再次验证：如果段落只有文字说明（如"如下图所示"），但没有实际图片，则跳过
                # 检查段落中是否有实际的图片元素，而不仅仅是文字
                has_actual_image_element = False
                try:
                    para_xml_full = str(paragraph._element.xml)
                    # 必须包含pic:pic元素（这是真正的图片元素）
                    if 'pic:pic' in para_xml_full:
                        # 进一步验证：pic:pic中应该包含blip（图片数据）
                        # 或者包含embed/link引用
                        if 'a:blip' in para_xml_full or 'r:embed' in para_xml_full or 'r:link' in para_xml_full:
                            has_actual_image_element = True
                    # 或者直接包含a:blip且有引用
                    elif 'a:blip' in para_xml_full and ('r:embed' in para_xml_full or 'r:link' in para_xml_full):
                        has_actual_image_element = True
                except:
                    pass
                
                # 如果没有实际的图片元素，只是误判，跳过
                if not has_actual_image_element:
                    has_image = False
            
            # 如果找到图片，强制设置段落对齐为居中
            if has_image:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 如果找到图片，检查后面几个段落是否有图题
            if has_image:
                # 检查当前段落及后面最多5个段落是否有图题
                is_caption = False
                caption_paragraph_idx = None
                
                # 检查范围：当前段落 + 后面5个段落
                check_range = min(6, len(document.paragraphs) - idx)
                for offset in range(check_range):
                    check_idx = idx + offset
                    if check_idx >= len(document.paragraphs):
                        break
                    check_para = document.paragraphs[check_idx]
                    check_text = check_para.text.strip() if check_para.text else ""
                    
                    # 判断是否是图题：以"图"开头，且包含数字（如"图1-1"、"图2.1"等）
                    # 或者以"流程图"开头（如"流程图1-1"、"流程图2.1"等）
                    if check_text and len(check_text) < 100:
                        # 检查是否包含图号格式（图X-X、图X.X等）
                        if (check_text.startswith("图") and (re.search(r'图\s*\d+[\.\-]\d+', check_text) or re.search(r'图\s*\d+', check_text))):
                            is_caption = True
                            caption_paragraph_idx = check_idx
                            break
                        # 检查是否是流程图标题（流程图X-X、流程图X.X等）
                        elif (check_text.startswith("流程图") and (re.search(r'流程图\s*\d+[\.\-]\d+', check_text) or re.search(r'流程图\s*\d+', check_text))):
                            is_caption = True
                            caption_paragraph_idx = check_idx
                            break
                    
                    # 如果检查的段落已经有大量文字，说明图题不太可能在更后面了
                    if offset > 0 and len(check_text) > 50 and not check_text.startswith("图"):
                        break
                
                # 如果找到图题，强制设置图题段落居中对齐
                if is_caption and caption_paragraph_idx is not None:
                    caption_para = document.paragraphs[caption_paragraph_idx]
                    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # 如果没有找到图题，记录问题
                if not is_caption:
                    # 获取图片所在位置的上下文（前后各一段）
                    context_before = ""
                    context_after = ""
                    if idx > 0:
                        context_before = document.paragraphs[idx - 1].text.strip()[:50]
                    if idx + 1 < len(document.paragraphs):
                        context_after = document.paragraphs[idx + 1].text.strip()[:50]
                    
                    issues.append({
                        "paragraph_index": idx,
                        "type": "missing_figure_caption",
                        "message": "图片后缺少图题说明",
                        "context_before": context_before,
                        "context_after": context_after,
                        "suggestion": "请在图片后添加图题，格式如：图X-X 图片说明"
                    })
                    missing_caption_indices.append(idx)
        
        # 不再在文档中插入标记，只记录问题到issues中
        # 最终文档应该看起来像标准文档，不显示修改痕迹
        
        return issues

    def _check_reference_citations(self, document: Document) -> list:
        """检测参考文献引用标注，检查正文中是否有引用标注，返回缺失引用的问题列表
        注意：不在文档中插入标记，只记录问题到issues中，保持文档干净
        """
        issues = []
        
        # 1. 找到参考文献部分的起始位置（从后往前查找，找到最后一个"参考文献"标题）
        reference_start_idx = None
        reference_section_text = ""
        
        # 从后往前查找，找到最后一个"参考文献"标题（避免匹配到目录中的"参考文献"）
        for idx in range(len(document.paragraphs) - 1, -1, -1):
            paragraph = document.paragraphs[idx]
            para_text = paragraph.text.strip() if paragraph.text else ""
            # 检测参考文献标题（可能包含"参考文献"、"References"、"参考书目"等）
            if re.search(r'参考(文献|书目)', para_text) or para_text.lower().startswith('references') or para_text.lower().startswith('bibliography'):
                # 确保是标题格式（通常较短，且可能是居中或单独一行）
                if len(para_text) < 50 or para_text in ["参考文献", "References", "参考书目", "Bibliography"]:
                    reference_start_idx = idx
                    # 收集参考文献部分的内容（最多收集100个段落）
                    ref_paragraphs = []
                    for i in range(idx, min(idx + 100, len(document.paragraphs))):
                        ref_paragraphs.append(document.paragraphs[i].text.strip() if document.paragraphs[i].text else "")
                    reference_section_text = "\n".join(ref_paragraphs)
                    break
        
        # 如果没有找到参考文献部分，提示用户
        if reference_start_idx is None:
            issues.append({
                "type": "no_reference_section",
                "message": "未找到参考文献部分",
                "suggestion": "请在文档末尾添加参考文献部分，标题为'参考文献'"
            })
            return issues
        
        # 2. 提取参考文献列表（通常以数字编号开头，如 [1]、1. 等）
        reference_items = []
        reference_patterns = [
            r'^\[\d+\]',  # [1] 格式
            r'^\d+\.',    # 1. 格式
            r'^\(\d+\)',  # (1) 格式
        ]
        
        for idx in range(reference_start_idx + 1, min(reference_start_idx + 100, len(document.paragraphs))):
            para = document.paragraphs[idx]
            # 获取原始文本，不strip，以便检查开头格式
            para_text_raw = para.text if para.text else ""
            para_text = para_text_raw.strip()
            
            # 如果遇到新的章节标题，停止收集
            if len(para_text) < 50 and (para_text.startswith("第") or para_text.startswith("Chapter") or 
                                         para_text.startswith("附录") or para_text.startswith("Appendix")):
                break
            
            # 排除章节标题（如"1.2"、"1.2.1"、"第一章"等）
            is_section_title = False
            # 检查是否是章节标题格式
            if re.match(r'^\d+\.\d+', para_text) or re.match(r'^\d+\.\d+\.\d+', para_text):  # 1.2 或 1.2.1 格式
                # 如果段落较短（通常是标题），且不包含参考文献特征，则不是参考文献
                if len(para_text) < 100:
                    is_section_title = True
            # 检查是否是章节标题（如"第一章"、"第1章"等）
            if re.match(r'^第[一二三四五六七八九十\d]+章|^第\d+章|^Chapter\s+\d+', para_text):
                is_section_title = True
            # 检查是否是标题样式
            if para.style and ("标题" in para.style.name or "heading" in para.style.name.lower()):
                is_section_title = True
            
            # 如果确定是章节标题，跳过
            if is_section_title:
                continue
            
            # 检查是否符合参考文献格式
            is_reference = False
            ref_number = None
            
            # 首先尝试从段落开头提取编号（更准确）
            # 检查常见的参考文献编号格式
            # 注意：参考文献格式可能是 [1]  作者名...（[1]后面有多个空格）
            # 改进：使用 search 而不是 match，允许前面有少量空格
            number_match = None
            
            # 检查 [数字] 格式（优先检查，因为这是最常见的格式）
            # 只支持半角方括号 [数字]（参考文献标注一定带英文版的方括号）
            # 使用 search 查找，但检查是否在段落开头（允许前面有少量空格）
            bracket_match = re.search(r'\[(\d+)\]', para_text)
            
            if bracket_match:
                # 检查 [数字] 是否在段落开头（允许前面有少量空格）
                bracket_pos = para_text.find(bracket_match.group(0))
                # 改进：允许 [数字] 在段落开头10个字符内，提高容错性
                if bracket_pos <= 10:  # [数字] 在段落开头10个字符内
                    # 进一步验证：确保 [数字] 后面有内容（不是单独的 [2]）
                    bracket_end = bracket_pos + len(bracket_match.group(0))
                    remaining_after_bracket = para_text[bracket_end:].strip()
                    # 如果 [数字] 后面有内容（至少5个字符），认为是参考文献
                    if len(remaining_after_bracket) >= 5:
                        is_reference = True
                        ref_number = int(bracket_match.group(1))
                        print(f"[DocumentService] 通过半角方括号 [数字] 格式识别参考文献: {ref_number} (位置: {bracket_pos}, 后续文本长度: {len(remaining_after_bracket)})")
            
            # 如果还没有识别，继续检查其他格式
            if not is_reference:
                if re.match(r'^\d+\.', para_text):  # 1. 格式
                    number_match = re.search(r'^\d+', para_text)
                    if number_match:
                        is_reference = True
                        ref_number = int(number_match.group())
                        print(f"[DocumentService] 通过 数字. 格式识别参考文献: {ref_number}")
                elif re.match(r'^\(\d+\)', para_text):  # (1) 格式
                    number_match = re.search(r'\d+', para_text)
                    if number_match:
                        is_reference = True
                        ref_number = int(number_match.group())
                        print(f"[DocumentService] 通过 (数字) 格式识别参考文献: {ref_number}")
                else:
                    # 尝试其他格式：可能是空格分隔的编号，如 "1 作者名..."
                    number_match = re.match(r'^(\d+)\s+', para_text)
                    if number_match:
                        # 检查后面是否有参考文献特征
                        remaining_text = para_text[len(number_match.group(0)):].strip()
                        # 如果后面有作者名、年份等特征，可能是参考文献
                        has_year = re.search(r'\d{4}', remaining_text)
                        has_author = re.search(r'[，,]\s*\d{4}|[A-Z][a-z]+\s+[A-Z]', remaining_text)
                        if has_year or (has_author and len(remaining_text) > 20):
                            is_reference = True
                            ref_number = int(number_match.group(1))
                            print(f"[DocumentService] 通过 数字空格 格式识别参考文献: {ref_number}")
            
            # 如果还没有识别为参考文献，但段落较长且包含作者、年份等信息，也可能是参考文献
            # 但必须排除章节标题
            if not is_reference and len(para_text) > 20:
                # 检查是否包含常见的参考文献特征（作者名、年份、期刊名等）
                # 参考文献通常包含：作者、年份、期刊名、出版社等
                # 改进：支持更多年份格式（中文和英文）
                has_author_pattern_cn = re.search(r'[，,]\s*\d{4}[，,]', para_text)  # 中文格式：年份前后有逗号
                has_author_pattern_en = re.search(r'[A-Z][a-z]+\.[A-Z]', para_text)  # 英文格式：作者名（如 A. I.）
                has_journal_pattern = re.search(r'\[[JC]\]|期刊|学报|Journal|Conference', para_text, re.IGNORECASE)  # 期刊标识 [J] 或 [C]
                has_publisher_pattern = re.search(r'出版社|Press|Publishing', para_text, re.IGNORECASE)  # 出版社
                has_year = re.search(r'\d{4}', para_text)  # 年份（4位数字）
                
                # 改进识别逻辑：支持英文参考文献格式
                # 参考文献必须同时满足：有年份，且（有作者模式或期刊标识或出版社），且段落较长
                # 或者：有 [数字] 格式在开头，且有年份和期刊标识
                # 只支持半角方括号（参考文献标注一定带英文版的方括号）
                has_bracket_at_start = False
                bracket_match_at_start = re.search(r'\[(\d+)\]', para_text)
                if bracket_match_at_start:
                    bracket_pos = para_text.find(bracket_match_at_start.group(0))
                    if bracket_pos <= 10:  # [数字] 在段落开头10个字符内
                        has_bracket_at_start = True
                
                # 如果段落开头有 [数字] 格式，且有年份和期刊标识，认为是参考文献
                if has_bracket_at_start and has_year and has_journal_pattern:
                    is_reference = True
                    ref_number = int(bracket_match_at_start.group(1))
                    print(f"[DocumentService] 通过 [数字]+年份+期刊标识 识别参考文献: {ref_number}")
                # 或者满足传统的识别条件
                elif has_year and (has_author_pattern_cn or has_author_pattern_en or has_journal_pattern or has_publisher_pattern) and len(para_text) > 30:
                    is_reference = True
                    # 尝试从段落开头提取编号（更宽松的匹配）
                    # 可能格式：数字开头，后面跟空格或标点，或者 [数字] 格式
                    # 只尝试 [数字] 格式（只支持半角方括号）
                    bracket_match = re.search(r'\[(\d+)\]', para_text)
                    if bracket_match:
                        bracket_pos = para_text.find(bracket_match.group(0))
                        if bracket_pos <= 10:  # 在段落开头10个字符内
                            ref_number = int(bracket_match.group(1))
                    else:
                        # 尝试数字开头格式
                        number_match = re.search(r'^(\d+)', para_text)
                        if number_match:
                            ref_number = int(number_match.group(1))
                        else:
                            # 如果没有找到编号，使用序号（但这种情况应该很少）
                            ref_number = len(reference_items) + 1
                    print(f"[DocumentService] 通过内容特征识别参考文献: {ref_number} (年份: {has_year is not None}, 期刊: {has_journal_pattern is not None}, 作者: {has_author_pattern_cn is not None or has_author_pattern_en is not None})")
            
            if is_reference:
                # 如果还是没有编号，尝试从段落开头提取（更宽松的匹配）
                if ref_number is None:
                    # 尝试匹配：数字开头，后面跟空格、点、方括号、圆括号等
                    number_match = re.match(r'^(\d+)', para_text)
                    if number_match:
                        ref_number = int(number_match.group(1))
                    else:
                        # 如果还是没有，使用序号（确保每个参考文献都有编号）
                        ref_number = len(reference_items) + 1
                        print(f"[DocumentService] 警告：参考文献没有明确编号，使用序号 {ref_number}: {para_text[:50]}")
                
                reference_items.append({
                    "index": ref_number,
                    "number": ref_number,  # 确保编号一致
                    "text": para_text[:100],  # 只保存前100个字符
                    "paragraph_index": idx,
                    "paragraph": para,  # 保存段落对象，用于后续修改
                })
                print(f"[DocumentService] 识别参考文献 {ref_number}: {para_text[:50]}")
        
        # 如果没有找到参考文献条目，提示
        if not reference_items:
            issues.append({
                "type": "no_reference_items",
                "message": "参考文献部分为空或格式不正确",
                "suggestion": "请确保参考文献部分包含编号的参考文献条目"
            })
            return issues
        
        # 2.5. 检查参考文献数量是否满足要求（至少10篇）
        reference_count = len(reference_items)
        min_required = REFERENCE_REQUIREMENTS.get("min_total", 10)
        if reference_count < min_required:
            issues.append({
                "type": "insufficient_references",
                "message": f"参考文献数量不足：当前 {reference_count} 篇，至少需要 {min_required} 篇",
                "suggestion": f"请添加更多参考文献，至少需要 {min_required} 篇",
                "current_count": reference_count,
                "required_count": min_required,
                "missing_count": min_required - reference_count
            })
        
        # 3. 检查正文中是否有引用标注，并找出被引用的参考文献编号
        # 正文部分：从封面结束到参考文献部分之前
        body_start_idx = self._find_body_start_index(document)
        print(f"[DocumentService] 正文开始位置: {body_start_idx}, 参考文献开始位置: {reference_start_idx}")
        
        body_text = ""
        body_paragraphs = []
        # 记录每个引用所在的段落索引（用于计算页码）
        citation_locations = {}  # {ref_number: [paragraph_index1, paragraph_index2, ...]}
        
        # 从正文开始到参考文献之前的所有段落（包括短段落，因为引用可能在图片说明等短段落中）
        # 改进：确保能正确提取段落文本，包括所有 runs 的文本
        for idx in range(body_start_idx, reference_start_idx):
            para = document.paragraphs[idx]
            # 方法1：使用 para.text（这是最可靠的方法，会自动合并所有 runs）
            para_text = para.text.strip() if para.text else ""
            
            # 方法2：如果 para.text 为空，尝试手动合并所有 runs 的文本
            if not para_text:
                para_text = "".join([run.text for run in para.runs if run.text]).strip()
            
            # 方法3：如果还是为空，尝试从 XML 中提取文本（最后的手段）
            if not para_text:
                try:
                    para_xml = str(para._element.xml)
                    # 提取所有文本节点
                    import xml.etree.ElementTree as ET
                    root = ET.fromstring(para_xml)
                    texts = []
                    for elem in root.iter():
                        if elem.text:
                            texts.append(elem.text)
                    para_text = "".join(texts).strip()
                except:
                    pass
            
            # 检查所有段落（包括短段落），因为引用可能在图片说明、表格说明等短段落中
            if len(para_text) > 0:  # 只要有内容就检查
                body_text += para_text + " "
                body_paragraphs.append((idx, para_text))
                
                # 注意：不在遍历段落时检测普通文本中的引用
                # 只检测上标格式的引用（通过检查runs的格式）
                
                # 调试：如果段落包含 [4] 或 [5]，打印详细信息
                if '[4]' in para_text or '[5]' in para_text:
                    print(f"[DocumentService] 调试：段落 {idx} 包含引用")
                    print(f"[DocumentService] 段落文本: {para_text}")
                    print(f"[DocumentService] para.text: {para.text if para.text else 'None'}")
                    print(f"[DocumentService] runs数量: {len(para.runs)}")
                    for run_idx, run in enumerate(para.runs):
                        print(f"[DocumentService]   run {run_idx}: '{run.text}' (上标: {run.font.superscript if run.font else 'N/A'})")
        
        # 调试：检查 body_text 中是否包含 [4] 和 [5]
        print(f"[DocumentService] 正文文本总长度: {len(body_text)} 字符")
        if '[4]' in body_text:
            # 找到所有 [4] 的位置
            positions = []
            start = 0
            while True:
                pos = body_text.find('[4]', start)
                if pos == -1:
                    break
                positions.append(pos)
                start = pos + 1
            print(f"[DocumentService] 在正文中找到 {len(positions)} 个 [4]，位置: {positions}")
            for pos in positions[:3]:  # 只显示前3个
                context = body_text[max(0, pos-30):min(len(body_text), pos+30)]
                print(f"[DocumentService]   [4] 上下文: ...{context}...")
        else:
            print(f"[DocumentService] 警告：正文文本中未找到 [4]")
        
        if '[5]' in body_text:
            # 找到所有 [5] 的位置
            positions = []
            start = 0
            while True:
                pos = body_text.find('[5]', start)
                if pos == -1:
                    break
                positions.append(pos)
                start = pos + 1
            print(f"[DocumentService] 在正文中找到 {len(positions)} 个 [5]，位置: {positions}")
            for pos in positions[:3]:  # 只显示前3个
                context = body_text[max(0, pos-30):min(len(body_text), pos+30)]
                print(f"[DocumentService]   [5] 上下文: ...{context}...")
        else:
            print(f"[DocumentService] 警告：正文文本中未找到 [5]")
        
        # 检测引用标注的常见格式，并提取被引用的参考文献编号
        # 改进：支持更多格式，包括多个编号的完整提取
        citation_patterns = [
            (r'\[(\d+)\]', 'single'),                    # [1] 格式
            (r'\[(\d+)[,\s]+(\d+)\]', 'range_comma'),   # [1,2,3] 格式（逗号分隔，但只匹配两个数字）
            (r'\[(\d+)[\-\s]+(\d+)\]', 'range_dash'),   # [1-5] 或 [1 5] 格式（连字符或空格分隔）
            (r'\((\d+)\)', 'paren_single'),              # (1) 格式（圆括号）
            (r'（(\d+)）', 'paren_single_cn'),          # （1）格式（中文圆括号）
            (r'\((\d+)[,\s]+(\d+)\)', 'paren_range'),  # (1,2,3) 格式
            (r'（(\d+)[,\s]+(\d+)）', 'paren_range_cn'), # （1,2,3）格式
            # 注意：年份格式 (2020) 不提取为参考文献编号，因为可能是作者-年份引用格式
        ]
        
        cited_reference_numbers = set()  # 被引用的参考文献编号集合
        
        # 根据用户要求：只有上标格式的 [数字] 才算文献引用，别的都不算
        # 不再检测普通文本中的引用格式，只检测上标格式的引用（通过检查runs的格式）
        print(f"[DocumentService] 开始检测引用，正文文本长度: {len(body_text)}")
        print(f"[DocumentService] 注意：只检测上标格式的引用，普通文本中的引用不算")
        
        # 只检测上标格式的引用（通过检查runs的格式）
        # 根据用户要求：只有上标格式的 [数字] 才算文献引用，别的都不算
        # 毕业论文中，引用通常是在文字上方加入 [1], [2] 这种格式，通常是上标格式
        for idx in range(body_start_idx, reference_start_idx):
            para = document.paragraphs[idx]
            for run in para.runs:
                run_text = run.text.strip() if run.text else ""
                if not run_text:
                    continue
                
                # 只检查上标格式（这是唯一算作引用的格式）
                if run.font.superscript:
                    # 上标格式的引用可能是：
                    # 1. 纯数字：1, 2, 3
                    # 2. 方括号数字：[1], [2], [3]
                    # 3. 多个数字：[1,2,3] 或 [1-5]
                    
                    # 检查方括号格式的上标引用 [1], [2] 等（只支持半角方括号）
                    # 先检测半角方括号
                    bracket_matches = re.finditer(r'\[(\d+)\]', run_text)
                    for match in bracket_matches:
                        try:
                            num = int(match.group(1))
                            if 1 <= num <= 1000:
                                cited_reference_numbers.add(num)
                                # 记录引用位置（避免重复记录）
                                if num not in citation_locations:
                                    citation_locations[num] = []
                                # 只有当这个段落索引还没有记录时才添加，避免重复
                                if idx not in citation_locations[num]:
                                    citation_locations[num].append(idx)
                                print(f"[DocumentService] 检测到上标格式引用 [{num}]")
                        except ValueError:
                            pass
                    
                    # 检查多个编号的上标引用 [1,2,3,4,5] 或 [1-5]（改进：支持任意数量的编号，只支持半角方括号）
                    # 先检测多个编号格式 [1,2,3,4,5]（半角）
                    multi_bracket_pattern = r'\[(\d+(?:[,\s]+\d+)+)\]'
                    multi_matches = re.finditer(multi_bracket_pattern, run_text)
                    for match in multi_matches:
                        try:
                            numbers_str = match.group(1)  # 提取括号内的内容
                            # 提取所有数字
                            numbers = re.findall(r'\d+', numbers_str)
                            for num_str in numbers:
                                num = int(num_str.strip())
                                if 1 <= num <= 1000:
                                    cited_reference_numbers.add(num)
                                    # 记录引用位置（避免重复记录）
                                    if num not in citation_locations:
                                        citation_locations[num] = []
                                    # 只有当这个段落索引还没有记录时才添加，避免重复
                                    if idx not in citation_locations[num]:
                                        citation_locations[num].append(idx)
                                    print(f"[DocumentService] 检测到上标格式多个编号引用 [{num}]")
                        except ValueError:
                            pass
                    
                    # 再检测范围格式 [1-5] 或两个编号 [1,2]（只支持半角方括号）
                    range_matches = re.finditer(r'\[(\d+)[,\-\s]+(\d+)\]', run_text)
                    for match in range_matches:
                        try:
                            # 提取所有数字
                            numbers_str = match.group(0).strip('[]')
                            if ',' in numbers_str:
                                # 逗号分隔：[1,2] 或 [1,2,3]
                                for num_str in numbers_str.split(','):
                                    num = int(num_str.strip())
                                    if 1 <= num <= 1000:
                                        cited_reference_numbers.add(num)
                                        # 记录引用位置（避免重复记录）
                                        if num not in citation_locations:
                                            citation_locations[num] = []
                                        # 只有当这个段落索引还没有记录时才添加，避免重复
                                        if idx not in citation_locations[num]:
                                            citation_locations[num].append(idx)
                                        print(f"[DocumentService] 检测到上标格式逗号分隔引用 [{num}]")
                            elif '-' in numbers_str:
                                # 连字符分隔：[1-5]
                                parts = numbers_str.split('-')
                                if len(parts) == 2:
                                    start = int(parts[0].strip())
                                    end = int(parts[1].strip())
                                    if 1 <= start <= end <= 1000:
                                        for num in range(start, end + 1):
                                            cited_reference_numbers.add(num)
                                            # 记录引用位置（避免重复记录）
                                            if num not in citation_locations:
                                                citation_locations[num] = []
                                            # 只有当这个段落索引还没有记录时才添加，避免重复
                                            if idx not in citation_locations[num]:
                                                citation_locations[num].append(idx)
                                        print(f"[DocumentService] 检测到上标格式范围引用 [{start}-{end}]")
                        except ValueError:
                            pass
                    
                    # 注意：根据用户要求，只有上标格式的 [数字] 才算引用
                    # 纯数字的上标（没有方括号的）不算引用，所以不再检测
                
                # 注意：不再检测普通文本中的引用格式
                # 只检测上标格式的引用（已在上面的 if run.font.superscript 中处理）
        
        # 4. 找出未被引用的参考文献
        # 调试信息：打印检测到的参考文献编号和引用编号
        print(f"[DocumentService] 检测到 {len(reference_items)} 条参考文献")
        print(f"[DocumentService] 参考文献编号: {[ref['number'] for ref in reference_items]}")
        print(f"[DocumentService] 正文中引用的编号: {sorted(cited_reference_numbers)}")
        print(f"[DocumentService] 正文文本长度: {len(body_text)} 字符")
        print(f"[DocumentService] 正文段落数量: {len(body_paragraphs)}")
        
        # 额外调试：检查正文中是否包含 [4] 和 [5]
        if '[4]' in body_text:
            print(f"[DocumentService] 调试：正文中包含 [4]")
        if '[5]' in body_text:
            print(f"[DocumentService] 调试：正文中包含 [5]")
        if '[4,5]' in body_text or '[4, 5]' in body_text:
            print(f"[DocumentService] 调试：正文中包含 [4,5] 格式")
        
        uncited_references = []
        for ref_item in reference_items:
            ref_num = ref_item["number"]
            # 检查是否真正被引用：必须在 cited_reference_numbers 中，并且有位置记录
            # 如果不在引用集合中，或者没有位置记录，都统一标记为未引用
            locations = citation_locations.get(ref_num, [])
            if ref_num not in cited_reference_numbers or not locations:
                uncited_references.append(ref_item)
                print(f"[DocumentService] 未引用的参考文献: {ref_num} - {ref_item['text'][:50]}")
        
        print(f"[DocumentService] 未引用的参考文献数量: {len(uncited_references)}")
        print(f"[DocumentService] 引用位置记录: {citation_locations}")
        
        # 5. 在参考文献段落中标记引用信息
        # 逻辑：只标记未找到引用的参考文献
        # 1. 找到引用的：不添加任何标记
        # 2. 未找到引用的：显示"未找到标注页"（红色标记）
        for ref_item in reference_items:
            ref_num = ref_item["number"]
            para = ref_item["paragraph"]
            
            # 获取引用位置（段落索引列表）
            locations = citation_locations.get(ref_num, [])
            
            # 检查是否找到了引用：必须在 cited_reference_numbers 中，并且有位置记录
            # 不再在文档中标记，只记录问题到issues中
            # 最终文档应该看起来像标准文档，不显示修改痕迹
            if ref_num in cited_reference_numbers and locations:
                # 找到了引用，不添加任何标记
                print(f"[DocumentService] 参考文献 {ref_num} 已找到引用")
            else:
                # 未找到标注页码，只记录到日志，不修改文档
                print(f"[DocumentService] 参考文献 {ref_num} 未找到标注页（仅记录，不修改文档）")
        
        # 6. 生成问题报告
        # 统计未找到标注页的参考文献数量
        uncited_refs = [ref for ref in reference_items 
                       if ref["number"] not in cited_reference_numbers 
                       or not citation_locations.get(ref["number"], [])]
        if uncited_refs:
            issues.append({
                "type": "uncited_references",
                "message": f"发现 {len(uncited_refs)} 条参考文献未找到标注页",
                "suggestion": "请在正文中添加引用标注，或删除未被引用的参考文献",
                "uncited_count": len(uncited_refs),
                "uncited_references": [
                    {
                        "number": ref["number"],
                        "text_preview": ref["text"][:80] + "..."
                    }
                    for ref in uncited_refs[:10]  # 只显示前10个
                ]
            })
        
        # 如果没有找到引用标注，提示用户
        if not cited_reference_numbers and len(reference_items) > 0:
            # 找到正文段落中可能缺少引用的位置
            missing_citation_paragraphs = []
            for para_idx, para_text in body_paragraphs:
                # 如果段落较长（可能是正文），但没有引用标注，记录
                if len(para_text) > 100:
                    # 检查段落是否包含可能引用的内容（如"研究"、"文献"、"表明"等学术词汇）
                    academic_keywords = ['研究', '文献', '表明', '发现', '提出', '分析', '方法', '理论', '模型']
                    if any(keyword in para_text for keyword in academic_keywords):
                        missing_citation_paragraphs.append({
                            "paragraph_index": para_idx,
                            "text_preview": para_text[:80] + "..."
                        })
            
            if missing_citation_paragraphs:
                issues.append({
                    "type": "missing_citations",
                    "message": f"正文中缺少参考文献引用标注（发现 {len(reference_items)} 条参考文献，但正文中未找到引用标注）",
                    "suggestion": "请在正文中添加引用标注，格式如：[1] 或 [1,2,3] 或 (作者, 年份)",
                    "reference_count": len(reference_items),
                    "missing_citation_paragraphs": missing_citation_paragraphs[:10]  # 只显示前10个
                })
            else:
                issues.append({
                    "type": "missing_citations",
                    "message": f"正文中缺少参考文献引用标注（发现 {len(reference_items)} 条参考文献，但正文中未找到引用标注）",
                    "suggestion": "请在正文中添加引用标注，格式如：[1] 或 [1,2,3] 或 (作者, 年份)",
                    "reference_count": len(reference_items)
                })
        
        return issues

    def _check_excessive_blanks(self, document: Document) -> list:
        """
        检测文档中的大段空白
        
        规则：
        - 只在正文部分检测空白行
        - 封面、诚信承诺、摘要、目录等部分不检测空白行
        - 先识别大章节（1、2、3或一、二、三开头，字体三号约16磅）
        - 只在大章节内部检测空白行（连续2个以上空白段落）
        - 章节之间不需要检测空白行
        
        Returns:
            问题列表
        """
        issues = []
        
        # 1. 使用 _find_section_ranges 获取正文范围
        # 明确排除封面、诚信承诺、摘要、Abstract、目录等部分，这些部分完全不检测空白行
        section_ranges = self._find_section_ranges(document)
        body_start_idx = None
        body_end_idx = len(document.paragraphs)
        
        # 获取正文范围
        if "body" in section_ranges:
            body_start_idx, body_end_idx = section_ranges["body"]
        
        # 如果没有找到正文范围，使用原来的方法查找
        if body_start_idx is None:
            body_start_idx = self._find_body_start_index(document)
            body_end_idx = len(document.paragraphs)
        
        # 2. 找到参考文献开始位置（作为检测结束位置）
        reference_start_idx = None
        for idx, paragraph in enumerate(document.paragraphs):
            para_text = paragraph.text.strip() if paragraph.text else ""
            if re.search(r'参考(文献|书目)', para_text) or para_text.lower().startswith('references') or para_text.lower().startswith('bibliography'):
                reference_start_idx = idx
                break
        
        # 3. 找到致谢部分（如果存在，也要排除）
        acknowledgement_start_idx = None
        for idx, paragraph in enumerate(document.paragraphs):
            para_text = paragraph.text.strip() if paragraph.text else ""
            if re.search(r'^(致谢|Acknowledgement)', para_text, re.IGNORECASE):
                acknowledgement_start_idx = idx
                break
        
        # 确定检测范围：只在正文范围内检测空白行
        check_start_idx = body_start_idx
        if check_start_idx is None:
            return issues
        
        # 检测结束位置：正文结束位置、参考文献开始或致谢开始，取较早的
        check_end_idx = body_end_idx
        if reference_start_idx is not None and reference_start_idx < check_end_idx:
            check_end_idx = reference_start_idx
        if acknowledgement_start_idx is not None and acknowledgement_start_idx < check_end_idx:
            check_end_idx = acknowledgement_start_idx
        
        if check_start_idx >= check_end_idx:
            return issues
        
        # 获取诚信承诺和摘要的范围，确保不删除它们之间的内容
        integrity_start = None
        integrity_end = None
        abstract_zh_start = None
        if "integrity" in section_ranges:
            integrity_start, integrity_end = section_ranges["integrity"]
        if "abstract_zh" in section_ranges:
            abstract_zh_start, _ = section_ranges["abstract_zh"]
        
        # 2. 识别大章节标题
        # 大章节特征：数字（1、2、3、4、5、6、7、8等）或中文一、二、三开头，字体三号（约16磅）
        def is_major_chapter_title(paragraph) -> bool:
            para_text = paragraph.text.strip() if paragraph.text else ""
            if not para_text:
                return False
            
            # 检查是否以数字（1-9）或中文一、二、三开头
            # 支持格式：1 绪论、1. 绪论、4. 剔除粗大误差、第一章、第1章、第4章、一 绪论等
            major_chapter_patterns = [
                r'^\d+\s+',  # 1 、2 、3、4、5、6、7、8 等开头（数字+空格）
                r'^\d+\.',  # 1.、2.、3.、4.、5. 等开头（数字+点）
                r'^[一二三四五六七八九十]\s+',  # 一 、二 、三 开头（中文数字+空格）
                r'^第[一二三四五六七八九十]章',  # 第一章、第二章等
                r'^第\d+章',  # 第1章、第2章、第3章、第4章等
                r'^\d+\s+[^\d]',  # 1 绪论、2 概述、4 剔除粗大误差等（数字+空格+非数字）
            ]
            
            # 先检查文本模式
            pattern_matched = False
            for pattern in major_chapter_patterns:
                if re.match(pattern, para_text):
                    pattern_matched = True
                    break
            
            if not pattern_matched:
                return False
            
            # 检查字体大小是否为三号（约16磅，允许14-18磅的范围，因为可能有些偏差）
            # 大章节必须是三号字体，不能仅凭文本模式判断
            has_three_size_font = False
            for run in paragraph.runs:
                if run.text.strip():
                    font_size = run.font.size.pt if run.font.size else None
                    if font_size is not None:
                        # 三号字通常是16磅，允许14-18磅的范围
                        if 14 <= font_size <= 18:
                            has_three_size_font = True
                            break
            
            # 大章节必须同时满足：文本模式匹配（数字或中文数字开头）AND 三号字体
            # 如果只有文本模式匹配但没有三号字体，不是大章节（可能是小节标题或其他格式）
            return has_three_size_font
        
        # 3. 识别所有大章节的边界（只在正文范围内识别）
        major_chapters = []  # [(start_idx, end_idx), ...]
        current_chapter_start = None
        
        # 确保检测范围从正文开始，不包括摘要、Abstract、目录等
        excluded_keywords = ['摘要', 'Abstract', '目录', 'Contents', '关键词', 'Key words', 'KeyWords']
        
        for idx in range(check_start_idx, check_end_idx):
            paragraph = document.paragraphs[idx]
            para_text = paragraph.text.strip() if paragraph.text else ""
            
            # 再次检查是否在排除部分内（双重保险）
            is_excluded = False
            for keyword in excluded_keywords:
                if re.search(rf'^{re.escape(keyword)}', para_text, re.IGNORECASE):
                    is_excluded = True
                    break
            
            if is_excluded:
                continue
            
            if is_major_chapter_title(paragraph):
                # 如果之前有章节，结束之前的章节
                if current_chapter_start is not None:
                    major_chapters.append((current_chapter_start, idx))
                # 开始新的大章节
                current_chapter_start = idx
        
        # 处理最后一个章节
        if current_chapter_start is not None:
            major_chapters.append((current_chapter_start, check_end_idx))
        
        # 如果没有找到大章节，将整个检测范围作为一个章节处理
        # 这样可以确保即使没有识别到大章节，也能检测空白行
        if not major_chapters:
            major_chapters = [(check_start_idx, check_end_idx)]
        
        # 调试信息：打印识别到的大章节
        # print(f"[空白行检测] 识别到 {len(major_chapters)} 个大章节: {major_chapters}")
        
        # 4. 只在大章节内部检测空白行（确保不在摘要、Abstract、目录等部分）
        def is_blank_paragraph(paragraph) -> bool:
            para_text = paragraph.text.strip() if paragraph.text else ""
            return len(para_text) == 0
        
        def has_page_break(paragraph) -> bool:
            """检查段落是否包含分页符"""
            # 检查段落格式中的分页符
            if paragraph.paragraph_format.page_break_before:
                return True
            # 检查runs中的分页符
            for run in paragraph.runs:
                if hasattr(run, 'element'):
                    run_xml = str(run.element.xml)
                    if 'w:br' in run_xml and 'type="page"' in run_xml:
                        return True
            return False
        
        # 排除关键词列表（摘要、Abstract、目录等部分完全不检测空白行）
        excluded_keywords = ['摘要', 'Abstract', '目录', 'Contents', '关键词', 'Key words', 'KeyWords']
        
        # 如果大章节范围为空，直接在整个检测范围内检测空白行
        if not major_chapters:
            # 在整个检测范围内检测空白行
            consecutive_blanks = 0
            blank_start_idx = None
            
            for idx in range(check_start_idx, check_end_idx):
                paragraph = document.paragraphs[idx]
                para_text = paragraph.text.strip() if paragraph.text else ""
                
                # 检查是否在排除部分内
                is_excluded = False
                for keyword in excluded_keywords:
                    if re.search(rf'^{re.escape(keyword)}', para_text, re.IGNORECASE):
                        is_excluded = True
                        break
                
                if is_excluded:
                    consecutive_blanks = 0
                    blank_start_idx = None
                    continue
                
                # 检查是否在诚信承诺和摘要之间，如果是，则不删除任何内容
                is_between_integrity_and_abstract = False
                if integrity_end is not None and abstract_zh_start is not None:
                    if blank_start_idx is not None:
                        if integrity_end <= blank_start_idx < abstract_zh_start:
                            is_between_integrity_and_abstract = True
                    elif integrity_end <= idx < abstract_zh_start:
                        is_between_integrity_and_abstract = True
                
                if is_between_integrity_and_abstract:
                    # 在诚信承诺和摘要之间，不删除任何内容，重置计数
                    consecutive_blanks = 0
                    blank_start_idx = None
                    continue
                
                if is_blank_paragraph(paragraph):
                    if consecutive_blanks == 0:
                        blank_start_idx = idx
                    consecutive_blanks += 1
                else:
                    # 遇到非空白段落
                    if consecutive_blanks >= 2 and blank_start_idx is not None:
                        # 直接删除连续空白段落
                        deleted_count = 0
                        for delete_idx in range(blank_start_idx + consecutive_blanks - 1, blank_start_idx - 1, -1):
                            if delete_idx < len(document.paragraphs):
                                para_to_delete = document.paragraphs[delete_idx]
                                if is_blank_paragraph(para_to_delete):
                                    # 检查是否包含分页符，如果包含则不删除（避免导致空白页）
                                    if has_page_break(para_to_delete):
                                        continue
                                    para_to_delete._element.getparent().remove(para_to_delete._element)
                                    deleted_count += 1
                        
                        if deleted_count > 0:
                            issues.append({
                                "type": "excessive_blanks_in_chapter",
                                "message": f"已删除第 {blank_start_idx + 1} 段到第 {blank_start_idx + consecutive_blanks} 段之间的 {deleted_count} 个连续空白段落",
                                "suggestion": "已自动删除章节内的多余空白",
                                "blank_start": blank_start_idx,
                                "blank_count": deleted_count,
                                "paragraph_indices": list(range(blank_start_idx, blank_start_idx + consecutive_blanks))
                            })
                    
                    consecutive_blanks = 0
                    blank_start_idx = None
            
            # 处理末尾的连续空白
            if consecutive_blanks >= 2 and blank_start_idx is not None:
                deleted_count = 0
                for delete_idx in range(blank_start_idx + consecutive_blanks - 1, blank_start_idx - 1, -1):
                    if delete_idx < len(document.paragraphs):
                        para_to_delete = document.paragraphs[delete_idx]
                        if is_blank_paragraph(para_to_delete):
                            # 检查：确保不删除包含字段代码的段落（如TOC字段）
                            para_xml = para_to_delete._element.xml if hasattr(para_to_delete._element, 'xml') else ""
                            if 'TOC' in para_xml or 'w:fldChar' in para_xml or 'w:instrText' in para_xml:
                                # 包含字段代码，不删除
                                continue
                            # 检查是否包含分页符，如果包含则不删除（避免导致空白页）
                            if has_page_break(para_to_delete):
                                continue
                            para_to_delete._element.getparent().remove(para_to_delete._element)
                            deleted_count += 1
                
                if deleted_count > 0:
                    issues.append({
                        "type": "excessive_blanks_in_chapter",
                        "message": f"已删除第 {blank_start_idx + 1} 段到第 {blank_start_idx + consecutive_blanks} 段之间的 {deleted_count} 个连续空白段落",
                        "suggestion": "已自动删除章节内的多余空白",
                        "blank_start": blank_start_idx,
                        "blank_count": deleted_count,
                        "paragraph_indices": list(range(blank_start_idx, blank_start_idx + consecutive_blanks))
                    })
        
        for chapter_start, chapter_end in major_chapters:
            consecutive_blanks = 0
            blank_start_idx = None
            
            # 在大章节内部遍历
            for idx in range(chapter_start + 1, chapter_end):  # +1 跳过章节标题本身
                # 确保索引在检测范围内（从正文开始，不包括摘要、目录等）
                if idx < check_start_idx:
                    continue
                
                paragraph = document.paragraphs[idx]
                para_text = paragraph.text.strip() if paragraph.text else ""
                
                # 检查是否在排除部分内（摘要、Abstract、目录等部分完全不检测）
                is_excluded = False
                for keyword in excluded_keywords:
                    if re.search(rf'^{re.escape(keyword)}', para_text, re.IGNORECASE):
                        is_excluded = True
                        break
                
                # 检查段落是否包含目录字段代码（TOC字段），如果包含则不删除
                if not is_excluded:
                    # 检查段落XML中是否包含TOC字段
                    para_xml = paragraph._element.xml if hasattr(paragraph._element, 'xml') else ""
                    if 'TOC' in para_xml or 'w:fldChar' in para_xml or 'w:instrText' in para_xml:
                        is_excluded = True
                
                if is_excluded:
                    # 如果在排除部分内，重置空白计数，不检测
                    consecutive_blanks = 0
                    blank_start_idx = None
                    continue
                
                # 检查是否在诚信承诺和摘要之间，如果是，则不删除任何内容
                is_between_integrity_and_abstract = False
                if integrity_end is not None and abstract_zh_start is not None:
                    if blank_start_idx is not None:
                        if integrity_end <= blank_start_idx < abstract_zh_start:
                            is_between_integrity_and_abstract = True
                    elif integrity_end <= idx < abstract_zh_start:
                        is_between_integrity_and_abstract = True
                
                if is_between_integrity_and_abstract:
                    # 在诚信承诺和摘要之间，不删除任何内容，重置计数
                    consecutive_blanks = 0
                    blank_start_idx = None
                    continue
                
                if is_blank_paragraph(paragraph):
                    if consecutive_blanks == 0:
                        blank_start_idx = idx
                    consecutive_blanks += 1
                else:
                    # 遇到非空白段落
                    if consecutive_blanks >= 2 and blank_start_idx is not None:
                        # 检查空白行是否在章节边界处（目录和正文之间、大章节之间）
                        # 如果空白行紧邻大章节标题，则不删除（这是章节间的空白，应该保留）
                        is_at_chapter_boundary = False
                        
                        # 检查空白行之后是否有大章节标题（如果空白行后面是大章节标题，这是章节间的空白，不删除）
                        for next_idx in range(blank_start_idx + consecutive_blanks, min(blank_start_idx + consecutive_blanks + 3, len(document.paragraphs))):
                            if next_idx < len(document.paragraphs):
                                next_para = document.paragraphs[next_idx]
                                if is_major_chapter_title(next_para):
                                    is_at_chapter_boundary = True
                                    break
                        
                        # 检查空白行之前是否有大章节标题（如果空白行前面是大章节标题，这也是章节间的空白，不删除）
                        if not is_at_chapter_boundary:
                            for prev_idx in range(max(0, blank_start_idx - 2), blank_start_idx):
                                if prev_idx < len(document.paragraphs):
                                    prev_para = document.paragraphs[prev_idx]
                                    if is_major_chapter_title(prev_para):
                                        is_at_chapter_boundary = True
                                        break
                        
                        # 检查空白行是否在目录和正文之间
                        # 如果空白行之前有"目录"关键词，且空白行之后有正文开始标记（如"1 称重技术和衡器的发展"），则不删除
                        if not is_at_chapter_boundary:
                            has_toc_before = False
                            has_body_after = False
                            
                            # 检查空白行之前是否有"目录"关键词（扩大检查范围到20个段落）
                            for prev_idx in range(max(0, blank_start_idx - 20), blank_start_idx):
                                if prev_idx < len(document.paragraphs):
                                    prev_text = document.paragraphs[prev_idx].text.strip() if document.paragraphs[prev_idx].text else ""
                                    if re.search(r'^(目录|Contents)', prev_text, re.IGNORECASE):
                                        has_toc_before = True
                                        break
                            
                            # 检查空白行之后是否有正文开始标记（如"1 称重技术和衡器的发展"、"第一章"等）
                            if has_toc_before:
                                for next_idx in range(blank_start_idx + consecutive_blanks, min(blank_start_idx + consecutive_blanks + 10, len(document.paragraphs))):
                                    if next_idx < len(document.paragraphs):
                                        next_text = document.paragraphs[next_idx].text.strip() if document.paragraphs[next_idx].text else ""
                                        # 检查是否是正文开始标记
                                        if (re.match(r'^[1-9]\s+', next_text) or  # 1 称重技术和衡器的发展
                                            re.match(r'^[1-9]\.', next_text) or  # 1. 绪论
                                            re.match(r'^第一章', next_text) or  # 第一章
                                            re.match(r'^第1章', next_text) or  # 第1章
                                            re.match(r'^第[一二三四五六七八九十]章', next_text) or  # 第一章、第二章等
                                            next_text == "绪论" or next_text == "概述"):  # 绪论、概述
                                            has_body_after = True
                                            break
                            
                            # 如果空白行在目录和正文之间，不删除
                            if has_toc_before and has_body_after:
                                is_at_chapter_boundary = True
                        
                        # 检查空白行之前是否有小节标题（如"4. 剔除粗大误差"），如果是小节标题后的空白，应该删除
                        # 小节标题格式：数字. 文字（如"4. 剔除粗大误差"）
                        if is_at_chapter_boundary:
                            # 如果空白行之前是小节标题（不是大章节），则应该删除
                            for prev_idx in range(max(0, blank_start_idx - 3), blank_start_idx):
                                if prev_idx < len(document.paragraphs):
                                    prev_para = document.paragraphs[prev_idx]
                                    prev_text = prev_para.text.strip() if prev_para.text else ""
                                    # 检查是否是小节标题格式（数字. 文字，但不是大章节）
                                    if prev_text and re.match(r'^\d+\.\s+', prev_text):
                                        # 进一步确认不是大章节（大章节必须是1、2、3开头且三号字体）
                                        if not is_major_chapter_title(prev_para):
                                            # 是小节标题，不是大章节，应该删除空白行
                                            is_at_chapter_boundary = False
                                            break
                        
                        # 只有不在章节边界处的空白行才删除
                        if not is_at_chapter_boundary:
                            # 在大章节内部发现连续空白，直接删除这些空白段落
                            # 从后往前删除，避免索引变化
                            deleted_count = 0
                            for delete_idx in range(blank_start_idx + consecutive_blanks - 1, blank_start_idx - 1, -1):
                                if delete_idx < len(document.paragraphs):
                                    para_to_delete = document.paragraphs[delete_idx]
                                    # 确认是空白段落再删除
                                    if is_blank_paragraph(para_to_delete):
                                        # 再次检查：确保不删除包含字段代码的段落（如TOC字段）
                                        para_xml = para_to_delete._element.xml if hasattr(para_to_delete._element, 'xml') else ""
                                        if 'TOC' in para_xml or 'w:fldChar' in para_xml or 'w:instrText' in para_xml:
                                            # 包含字段代码，不删除
                                            continue
                                        # 检查是否包含分页符，如果包含则不删除（避免导致空白页）
                                        if has_page_break(para_to_delete):
                                            continue
                                        # 删除段落
                                        para_to_delete._element.getparent().remove(para_to_delete._element)
                                        deleted_count += 1
                            
                            # 记录删除的空白段落信息（用于报告）
                            if deleted_count > 0:
                                issues.append({
                                    "type": "excessive_blanks_in_chapter",
                                    "message": f"已删除第 {blank_start_idx + 1} 段到第 {blank_start_idx + consecutive_blanks} 段之间的 {deleted_count} 个连续空白段落（大章节内）",
                                    "suggestion": "已自动删除章节内的多余空白",
                                    "blank_start": blank_start_idx,
                                    "blank_count": deleted_count,
                                    "paragraph_indices": list(range(blank_start_idx, blank_start_idx + consecutive_blanks))
                                })
                    
                    consecutive_blanks = 0
                    blank_start_idx = None
            
            # 处理章节末尾的连续空白
            if consecutive_blanks >= 2 and blank_start_idx is not None:
                # 再次确认不在排除部分内
                is_excluded = False
                if blank_start_idx < len(document.paragraphs) and blank_start_idx >= check_start_idx:
                    para_text = document.paragraphs[blank_start_idx].text.strip() if document.paragraphs[blank_start_idx].text else ""
                    for keyword in excluded_keywords:
                        if re.search(rf'^{re.escape(keyword)}', para_text, re.IGNORECASE):
                            is_excluded = True
                            break
                
                if not is_excluded and blank_start_idx >= check_start_idx:
                    # 检查空白行是否在章节边界处（如果空白行后面是大章节标题，不删除）
                    is_at_chapter_boundary = False
                    
                    # 检查空白行之后是否有大章节标题（虽然已经到章节末尾，但也要检查）
                    for next_idx in range(blank_start_idx + consecutive_blanks, min(blank_start_idx + consecutive_blanks + 3, len(document.paragraphs))):
                        if next_idx < len(document.paragraphs):
                            next_para = document.paragraphs[next_idx]
                            if is_major_chapter_title(next_para):
                                is_at_chapter_boundary = True
                                break
                    
                    # 检查空白行之前是否有大章节标题
                    if not is_at_chapter_boundary:
                        for prev_idx in range(max(0, blank_start_idx - 2), blank_start_idx):
                            if prev_idx < len(document.paragraphs):
                                prev_para = document.paragraphs[prev_idx]
                                if is_major_chapter_title(prev_para):
                                    is_at_chapter_boundary = True
                                    break
                    
                    # 检查空白行是否在目录和正文之间（处理章节末尾的连续空白时也要检查）
                    if not is_at_chapter_boundary:
                        has_toc_before = False
                        has_body_after = False
                        
                        # 检查空白行之前是否有"目录"关键词（扩大检查范围到20个段落）
                        for prev_idx in range(max(0, blank_start_idx - 20), blank_start_idx):
                            if prev_idx < len(document.paragraphs):
                                prev_text = document.paragraphs[prev_idx].text.strip() if document.paragraphs[prev_idx].text else ""
                                if re.search(r'^(目录|Contents)', prev_text, re.IGNORECASE):
                                    has_toc_before = True
                                    break
                        
                        # 检查空白行之后是否有正文开始标记（如"1 称重技术和衡器的发展"、"第一章"等）
                        if has_toc_before:
                            for next_idx in range(blank_start_idx + consecutive_blanks, min(blank_start_idx + consecutive_blanks + 10, len(document.paragraphs))):
                                if next_idx < len(document.paragraphs):
                                    next_text = document.paragraphs[next_idx].text.strip() if document.paragraphs[next_idx].text else ""
                                    # 检查是否是正文开始标记
                                    if (re.match(r'^[1-9]\s+', next_text) or  # 1 称重技术和衡器的发展
                                        re.match(r'^[1-9]\.', next_text) or  # 1. 绪论
                                        re.match(r'^第一章', next_text) or  # 第一章
                                        re.match(r'^第1章', next_text) or  # 第1章
                                        re.match(r'^第[一二三四五六七八九十]章', next_text) or  # 第一章、第二章等
                                        next_text == "绪论" or next_text == "概述"):  # 绪论、概述
                                        has_body_after = True
                                        break
                        
                        # 如果空白行在目录和正文之间，不删除
                        if has_toc_before and has_body_after:
                            is_at_chapter_boundary = True
                    
                    # 只有不在章节边界处的空白行才删除
                    if not is_at_chapter_boundary:
                        # 直接删除章节末尾的连续空白段落
                        # 从后往前删除，避免索引变化
                        deleted_count = 0
                        for delete_idx in range(blank_start_idx + consecutive_blanks - 1, blank_start_idx - 1, -1):
                            if delete_idx < len(document.paragraphs):
                                para_to_delete = document.paragraphs[delete_idx]
                                # 确认是空白段落再删除
                                if is_blank_paragraph(para_to_delete):
                                    # 再次检查：确保不删除包含字段代码的段落（如TOC字段）
                                    para_xml = para_to_delete._element.xml if hasattr(para_to_delete._element, 'xml') else ""
                                    if 'TOC' in para_xml or 'w:fldChar' in para_xml or 'w:instrText' in para_xml:
                                        # 包含字段代码，不删除
                                        continue
                                    # 检查是否包含分页符，如果包含则不删除（避免导致空白页）
                                    if has_page_break(para_to_delete):
                                        continue
                                    # 删除段落
                                    para_to_delete._element.getparent().remove(para_to_delete._element)
                                    deleted_count += 1
                        
                        # 记录删除的空白段落信息（用于报告）
                        if deleted_count > 0:
                            issues.append({
                                "type": "excessive_blanks_in_chapter",
                                "message": f"已删除第 {blank_start_idx + 1} 段到第 {blank_start_idx + consecutive_blanks} 段之间的 {deleted_count} 个连续空白段落（大章节内）",
                                "suggestion": "已自动删除章节内的多余空白",
                                "blank_start": blank_start_idx,
                                "blank_count": deleted_count,
                                "paragraph_indices": list(range(blank_start_idx, blank_start_idx + consecutive_blanks))
                            })
        
        # 空白段落已直接删除，不需要标记
        return issues

    def _diagnose_integrity_abstract_separation(self, document: Document) -> Dict:
        """
        诊断诚信承诺和摘要之间的分页情况
        
        Returns:
            诊断信息字典
        """
        diagnosis = {
            "integrity_found": False,
            "abstract_found": False,
            "integrity_start_idx": None,
            "abstract_start_idx": None,
            "has_page_break_between": False,
            "page_break_locations": [],
            "paragraphs_between": [],
            "issue": None
        }
        
        # 查找诚信承诺
        integrity_pattern = re.compile(r'诚\s*信\s*承\s*诺', re.IGNORECASE)
        for idx, paragraph in enumerate(document.paragraphs):
            para_text = paragraph.text.strip() if paragraph.text else ""
            if integrity_pattern.search(para_text) and not diagnosis["integrity_found"]:
                diagnosis["integrity_found"] = True
                diagnosis["integrity_start_idx"] = idx
                self._log_to_file(f"[诊断] 找到诚信承诺，段落索引: {idx}, 文本: {para_text[:50]}")
                break
        
        # 查找摘要
        abstract_pattern = re.compile(r'^摘\s*要', re.IGNORECASE)
        for idx, paragraph in enumerate(document.paragraphs):
            para_text = paragraph.text.strip() if paragraph.text else ""
            if abstract_pattern.match(para_text) and not diagnosis["abstract_found"]:
                diagnosis["abstract_found"] = True
                diagnosis["abstract_start_idx"] = idx
                self._log_to_file(f"[诊断] 找到摘要，段落索引: {idx}, 文本: {para_text[:50]}")
                break
        
        if not diagnosis["integrity_found"] or not diagnosis["abstract_found"]:
            diagnosis["issue"] = "未找到诚信承诺或摘要"
            return diagnosis
        
        if diagnosis["abstract_start_idx"] <= diagnosis["integrity_start_idx"]:
            diagnosis["issue"] = "摘要在诚信承诺之前，顺序异常"
            return diagnosis
        
        # 检查诚信承诺和摘要之间的段落
        start_idx = diagnosis["integrity_start_idx"]
        end_idx = diagnosis["abstract_start_idx"]
        
        self._log_to_file(f"[诊断] 诚信承诺和摘要之间的段落索引: {start_idx} 到 {end_idx}")
        
        # 检查每个段落是否有分页符
        for idx in range(start_idx, end_idx):
            paragraph = document.paragraphs[idx]
            para_text = paragraph.text.strip() if paragraph.text else ""
            
            # 检查段落格式中的分页符
            has_page_break = False
            if paragraph.paragraph_format.page_break_before:
                has_page_break = True
                diagnosis["page_break_locations"].append({
                    "index": idx,
                    "type": "paragraph_format.page_break_before",
                    "text": para_text[:50]
                })
                self._log_to_file(f"[诊断] 段落 {idx} 有分页符 (paragraph_format.page_break_before): {para_text[:50]}")
            
            # 检查runs中的分页符
            for run_idx, run in enumerate(paragraph.runs):
                if hasattr(run, 'element'):
                    run_xml = str(run.element.xml)
                    if 'w:br' in run_xml and 'type="page"' in run_xml:
                        has_page_break = True
                        diagnosis["page_break_locations"].append({
                            "index": idx,
                            "type": f"run_{run_idx}_page_break",
                            "text": para_text[:50]
                        })
                        self._log_to_file(f"[诊断] 段落 {idx}, Run {run_idx} 有分页符: {para_text[:50]}")
            
            # 记录段落信息
            diagnosis["paragraphs_between"].append({
                "index": idx,
                "text": para_text[:100],
                "has_page_break": has_page_break,
                "is_blank": len(para_text) == 0
            })
        
        # 检查摘要标题本身是否有分页符
        abstract_para = document.paragraphs[diagnosis["abstract_start_idx"]]
        if abstract_para.paragraph_format.page_break_before:
            diagnosis["has_page_break_between"] = True
            diagnosis["page_break_locations"].append({
                "index": diagnosis["abstract_start_idx"],
                "type": "abstract_title_page_break_before",
                "text": abstract_para.text.strip()[:50]
            })
            self._log_to_file(f"[诊断] 摘要标题本身有分页符 (page_break_before)")
        
        # 检查摘要标题的runs中是否有分页符
        for run_idx, run in enumerate(abstract_para.runs):
            if hasattr(run, 'element'):
                run_xml = str(run.element.xml)
                if 'w:br' in run_xml and 'type="page"' in run_xml:
                    diagnosis["has_page_break_between"] = True
                    diagnosis["page_break_locations"].append({
                        "index": diagnosis["abstract_start_idx"],
                        "type": f"abstract_title_run_{run_idx}_page_break",
                        "text": abstract_para.text.strip()[:50]
                    })
                    self._log_to_file(f"[诊断] 摘要标题的Run {run_idx} 有分页符")
        
        # 检查前一个段落是否有分页符
        if diagnosis["abstract_start_idx"] > 0:
            prev_para = document.paragraphs[diagnosis["abstract_start_idx"] - 1]
            if prev_para.paragraph_format.page_break_before:
                diagnosis["has_page_break_between"] = True
                diagnosis["page_break_locations"].append({
                    "index": diagnosis["abstract_start_idx"] - 1,
                    "type": "prev_paragraph_page_break_before",
                    "text": prev_para.text.strip()[:50]
                })
                self._log_to_file(f"[诊断] 摘要前一个段落有分页符 (page_break_before)")
            
            for run_idx, run in enumerate(prev_para.runs):
                if hasattr(run, 'element'):
                    run_xml = str(run.element.xml)
                    if 'w:br' in run_xml and 'type="page"' in run_xml:
                        diagnosis["has_page_break_between"] = True
                        diagnosis["page_break_locations"].append({
                            "index": diagnosis["abstract_start_idx"] - 1,
                            "type": f"prev_paragraph_run_{run_idx}_page_break",
                            "text": prev_para.text.strip()[:50]
                        })
                        self._log_to_file(f"[诊断] 摘要前一个段落的Run {run_idx} 有分页符")
        
        if not diagnosis["has_page_break_between"]:
            diagnosis["issue"] = "诚信承诺和摘要之间没有分页符"
        else:
            diagnosis["issue"] = None
        
        return diagnosis

    def _ensure_integrity_abstract_separation(self, document: Document) -> bool:
        """
        确保诚信承诺和摘要分开在不同页
        
        如果它们之间没有分页符，在摘要标题前添加分页符
        
        Returns:
            bool: 是否进行了修复
        """
        # 1. 查找诚信承诺和摘要的位置
        section_ranges = self._find_section_ranges(document)
        
        self._log_to_file(f"[修复] 查找结果: {list(section_ranges.keys())}")
        
        if "integrity" not in section_ranges or "abstract_zh" not in section_ranges:
            self._log_to_file(f"[修复] ❌ 未找到诚信承诺或摘要，跳过分页修复")
            if "integrity" not in section_ranges:
                self._log_to_file(f"[修复]   缺少: integrity")
            if "abstract_zh" not in section_ranges:
                self._log_to_file(f"[修复]   缺少: abstract_zh")
            return False
        
        integrity_start, integrity_end = section_ranges["integrity"]
        abstract_zh_start, _ = section_ranges["abstract_zh"]
        
        self._log_to_file(f"[修复] 诚信承诺范围: {integrity_start} 到 {integrity_end}")
        self._log_to_file(f"[修复] 摘要起始位置: {abstract_zh_start}")
        
        # 2. 检查摘要标题前是否有分页符
        if abstract_zh_start >= len(document.paragraphs):
            self._log_to_file(f"[修复] ⚠️ 摘要位置超出文档范围")
            return False
        
        abstract_para = document.paragraphs[abstract_zh_start]
        
        # 检查摘要标题本身是否有分页符
        if abstract_para.paragraph_format.page_break_before:
            self._log_to_file(f"[修复] ✅ 摘要标题已有分页符 (page_break_before)")
            return False  # 已经有分页符
        
        # 检查摘要标题的runs中是否有分页符
        for run in abstract_para.runs:
            if hasattr(run, 'element'):
                run_xml = str(run.element.xml)
                if 'w:br' in run_xml and 'type="page"' in run_xml:
                    self._log_to_file(f"[修复] ✅ 摘要标题的runs中已有分页符")
                    return False  # 已经有分页符
        
        # 检查前一个段落是否有分页符
        if abstract_zh_start > 0:
            prev_para = document.paragraphs[abstract_zh_start - 1]
            if prev_para.paragraph_format.page_break_before:
                self._log_to_file(f"[修复] ✅ 摘要前一个段落已有分页符 (page_break_before)")
                return False  # 已经有分页符
            
            for run in prev_para.runs:
                if hasattr(run, 'element'):
                    run_xml = str(run.element.xml)
                    if 'w:br' in run_xml and 'type="page"' in run_xml:
                        self._log_to_file(f"[修复] ✅ 摘要前一个段落的runs中已有分页符")
                        return False  # 已经有分页符
        
        # 3. 没有分页符，强制添加分页符
        self._log_to_file(f"[修复] ⚠️ 诚信承诺和摘要之间没有分页符，强制添加分页符")
        
        # 方法1：在摘要标题段落设置分页符
        abstract_para.paragraph_format.page_break_before = True
        self._log_to_file(f"[修复] 方法1：已在摘要标题段落设置分页符 (page_break_before)")
        
        # 方法3：在摘要标题的runs中添加分页符（最可靠的方法）
        # 如果摘要标题有runs，在第一个run前添加分页符
        if abstract_para.runs:
            # 获取第一个run
            first_run = abstract_para.runs[0]
            # 在第一个run前插入分页符（使用正确的XML格式）
            br_xml = '<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="page"/>'
            br = parse_xml(br_xml)
            first_run._element.getparent().insert(0, br)
            self._log_to_file(f"[修复] ✅ 已在摘要标题的第一个run前添加分页符")
        else:
            # 如果没有runs，创建一个run并添加分页符
            run = abstract_para.add_run()
            br_xml = '<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="page"/>'
            br = parse_xml(br_xml)
            run._element.getparent().insert(0, br)
            self._log_to_file(f"[修复] ✅ 已创建run并添加分页符")
        
        self._log_to_file(f"[修复] ✅ 已使用多种方法强制添加分页符，确保诚信承诺和摘要分开")
        return True

    def _ensure_abstract_separation(self, document: Document) -> bool:
        """
        确保中文摘要和英文摘要分开在不同页
        
        Returns:
            bool: 是否进行了修复
        """
        # 1. 查找中文摘要和英文摘要的位置
        section_ranges = self._find_section_ranges(document)
        
        if "abstract_zh" not in section_ranges:
            self._log_to_file(f"[修复] ❌ 未找到中文摘要，跳过分页修复")
            return False
        
        if "abstract_en" not in section_ranges:
            self._log_to_file(f"[修复] ❌ 未找到英文摘要，跳过分页修复")
            self._log_to_file(f"[修复] 已找到的section: {list(section_ranges.keys())}")
            # 尝试重新查找英文摘要（可能是大小写问题）
            for idx in range(0, len(document.paragraphs)):
                para_text = document.paragraphs[idx].text.strip() if document.paragraphs[idx].text else ""
                if re.match(r'^abstract', para_text, re.IGNORECASE):
                    self._log_to_file(f"[修复] 重新找到英文摘要，段落索引: {idx}, 文本: {para_text[:50]}")
                    # 手动设置英文摘要范围
                    section_ranges["abstract_en"] = (idx, len(document.paragraphs))
                    break
            if "abstract_en" not in section_ranges:
                return False
        
        abstract_zh_start, abstract_zh_end = section_ranges["abstract_zh"]
        abstract_en_start, _ = section_ranges["abstract_en"]
        
        self._log_to_file(f"[修复] 中文摘要范围: {abstract_zh_start} 到 {abstract_zh_end}")
        self._log_to_file(f"[修复] 英文摘要起始位置: {abstract_en_start}")
        
        # 2. 检查英文摘要标题前是否有分页符
        if abstract_en_start >= len(document.paragraphs):
            self._log_to_file(f"[修复] ⚠️ 英文摘要位置超出文档范围")
            return False
        
        abstract_en_para = document.paragraphs[abstract_en_start]
        
        # 检查英文摘要标题本身是否有分页符
        if abstract_en_para.paragraph_format.page_break_before:
            self._log_to_file(f"[修复] ✅ 英文摘要标题已有分页符 (page_break_before)")
            return False  # 已经有分页符
        
        # 检查英文摘要标题的runs中是否有分页符
        for run in abstract_en_para.runs:
            if hasattr(run, 'element'):
                run_xml = str(run.element.xml)
                if 'w:br' in run_xml and 'type="page"' in run_xml:
                    self._log_to_file(f"[修复] ✅ 英文摘要标题的runs中已有分页符")
                    return False  # 已经有分页符
        
        # 检查前一个段落是否有分页符
        if abstract_en_start > 0:
            prev_para = document.paragraphs[abstract_en_start - 1]
            if prev_para.paragraph_format.page_break_before:
                self._log_to_file(f"[修复] ✅ 英文摘要前一个段落已有分页符 (page_break_before)")
                return False  # 已经有分页符
            
            for run in prev_para.runs:
                if hasattr(run, 'element'):
                    run_xml = str(run.element.xml)
                    if 'w:br' in run_xml and 'type="page"' in run_xml:
                        self._log_to_file(f"[修复] ✅ 英文摘要前一个段落的runs中已有分页符")
                        return False  # 已经有分页符
        
        # 3. 没有分页符，强制添加分页符
        self._log_to_file(f"[修复] ⚠️ 中文摘要和英文摘要之间没有分页符，强制添加分页符")
        
        # 方法1：在英文摘要标题段落设置分页符
        abstract_en_para.paragraph_format.page_break_before = True
        
        # 方法2：在英文摘要标题的runs中添加分页符（最可靠的方法）
        if abstract_en_para.runs:
            # 获取第一个run
            first_run = abstract_en_para.runs[0]
            # 在第一个run前插入分页符
            br_xml = '<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="page"/>'
            br = parse_xml(br_xml)
            first_run._element.getparent().insert(0, br)
            self._log_to_file(f"[修复] ✅ 已在英文摘要标题的第一个run前添加分页符")
        else:
            # 如果没有runs，创建一个run并添加分页符
            run = abstract_en_para.add_run()
            br_xml = '<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="page"/>'
            br = parse_xml(br_xml)
            run._element.getparent().insert(0, br)
            self._log_to_file(f"[修复] ✅ 已创建run并添加分页符")
        
        self._log_to_file(f"[修复] ✅ 已使用多种方法强制添加分页符，确保中文摘要和英文摘要分开")
        return True

    def _check_and_remove_blank_pages(self, document: Document) -> list:
        """
        检测并删除整页空白页
        
        规则：
        - 在整个文档中检测整页空白（包括封面、诚信承诺、摘要、目录、正文等所有部分）
        - 封面、诚信承诺、摘要、目录等部分的空白行不删除，但整页空白要删除
        - 检测连续的空白段落，如果这些空白段落导致整页空白，则删除
        - 不允许整页空白页存在
        
        Returns:
            问题列表
        """
        issues = []
        
        # 检测整页空白页的方法：
        # 1. 查找连续的大量空白段落（可能是整页空白）
        # 2. 检查这些空白段落是否包含分页符
        # 3. 如果确认是整页空白，删除这些空白段落
        
        # 定义整页空白的阈值：
        # - 连续10个以上空白段落可能是整页空白
        # - 连续5个以上空白段落，且前后有分页符，可能是只有页眉的空白页
        BLANK_PAGE_THRESHOLD = 10
        BLANK_PAGE_WITH_HEADER_THRESHOLD = 5  # 只有页眉的空白页阈值
        
        def has_page_break(paragraph) -> bool:
            """检查段落是否包含分页符"""
            # 检查段落格式中的分页符
            if paragraph.paragraph_format.page_break_before:
                return True
            # 检查runs中的分页符
            for run in paragraph.runs:
                if hasattr(run, 'element'):
                    run_xml = str(run.element.xml)
                    if 'w:br' in run_xml and 'type="page"' in run_xml:
                        return True
            return False
        
        consecutive_blanks = 0
        blank_start_idx = None
        
        # 获取诚信承诺和摘要的范围，确保不删除它们之间的内容
        section_ranges = self._find_section_ranges(document)
        integrity_start = None
        integrity_end = None
        abstract_zh_start = None
        abstract_zh_end = None
        abstract_en_start = None
        abstract_en_end = None
        if "integrity" in section_ranges:
            integrity_start, integrity_end = section_ranges["integrity"]
        if "abstract_zh" in section_ranges:
            abstract_zh_start, abstract_zh_end = section_ranges["abstract_zh"]
        if "abstract_en" in section_ranges:
            abstract_en_start, abstract_en_end = section_ranges["abstract_en"]
            self._log_to_file(f"[空白页检测] 英文摘要范围: {abstract_en_start} 到 {abstract_en_end}")
        else:
            self._log_to_file(f"[空白页检测] ⚠️ 未找到英文摘要范围")
        
        # 在整个文档中检测整页空白
        # 使用while循环，因为删除段落后索引会变化
        idx = 0
        while idx < len(document.paragraphs):
            paragraph = document.paragraphs[idx]
            para_text = paragraph.text.strip() if paragraph.text else ""
            
            # 检查是否在诚信承诺和摘要之间
            # 允许删除空白段落，但不删除包含分页符的段落
            is_between_integrity_and_abstract = False
            if integrity_end is not None and abstract_zh_start is not None:
                if integrity_end <= idx < abstract_zh_start:
                    is_between_integrity_and_abstract = True
            
            # 检查是否在英文摘要之后（可能是空白页）
            is_after_abstract_en = False
            if abstract_en_end is not None:
                if idx >= abstract_en_end:
                    is_after_abstract_en = True
                    # 添加诊断日志（仅记录前几个段落，避免日志过多）
                    is_blank = len(para_text) == 0
                    if idx == abstract_en_end or (idx < abstract_en_end + 5 and is_blank):
                        self._log_to_file(f"[空白页检测] 段落 {idx} 在英文摘要之后（abstract_en_end={abstract_en_end}），文本: '{para_text[:30]}'，是否空白: {is_blank}")
            
            if is_between_integrity_and_abstract:
                # 在诚信承诺和摘要之间，只删除空白段落（不包含分页符的）
                # 这样可以删除多余的空白页，但保留分页符
                if is_blank and not has_page_break(paragraph):
                    # 空白段落且不包含分页符，可以删除
                    if consecutive_blanks == 0:
                        blank_start_idx = idx
                    consecutive_blanks += 1
                    # 如果连续空白段落较多，可能是空白页，标记为可删除
                    if consecutive_blanks >= BLANK_PAGE_WITH_HEADER_THRESHOLD:
                        # 检查前后是否有分页符
                        has_break_before = False
                        if blank_start_idx > 0 and blank_start_idx - 1 < len(document.paragraphs):
                            prev_para = document.paragraphs[blank_start_idx - 1]
                            if prev_para.paragraph_format.page_break_before or any('w:br' in str(run.element.xml) and 'type="page"' in str(run.element.xml) for run in prev_para.runs if hasattr(run, 'element')):
                                has_break_before = True
                        
                        has_break_after = False
                        if idx + 1 < len(document.paragraphs):
                            next_para = document.paragraphs[idx + 1]
                            if next_para.paragraph_format.page_break_before or any('w:br' in str(run.element.xml) and 'type="page"' in str(run.element.xml) for run in next_para.runs if hasattr(run, 'element')):
                                has_break_after = True
                        
                        # 如果前后有分页符，说明是空白页，删除这些空白段落
                        if has_break_before or has_break_after:
                            delete_end = min(blank_start_idx + consecutive_blanks - 1, len(document.paragraphs) - 1)
                            deleted_count = 0
                            for delete_idx in range(delete_end, blank_start_idx - 1, -1):
                                if delete_idx >= 0 and delete_idx < len(document.paragraphs):
                                    para_to_delete = document.paragraphs[delete_idx]
                                    if len(para_to_delete.text.strip()) == 0 and not has_page_break(para_to_delete):
                                        para_to_delete._element.getparent().remove(para_to_delete._element)
                                        deleted_count += 1
                                        if delete_idx < idx:
                                            idx -= 1
                            
                            if deleted_count > 0:
                                issues.append({
                                    "type": "blank_page_removed",
                                    "message": f"已删除诚信承诺和摘要之间的 {deleted_count} 个空白段落（空白页）",
                                    "suggestion": "已自动删除空白页",
                                    "blank_start": blank_start_idx,
                                    "blank_count": deleted_count,
                                })
                                consecutive_blanks = 0
                                blank_start_idx = None
                                continue
                else:
                    # 非空白段落或包含分页符，重置计数
                    consecutive_blanks = 0
                    blank_start_idx = None
                idx += 1
                continue
            
            # 检查是否在英文摘要之后，如果是空白段落，需要特别处理
            if is_after_abstract_en:
                is_blank = len(para_text) == 0
                if is_blank:
                    if consecutive_blanks == 0:
                        blank_start_idx = idx
                        self._log_to_file(f"[空白页检测] 在英文摘要后发现空白段落开始，段落索引: {idx}")
                    consecutive_blanks += 1
                    # 如果连续空白段落较多，可能是空白页，需要删除
                    if consecutive_blanks >= BLANK_PAGE_WITH_HEADER_THRESHOLD:
                        self._log_to_file(f"[空白页检测] 英文摘要后连续空白段落达到阈值: {consecutive_blanks}，开始检查是否为空白页")
                        # 检查前后是否有分页符
                        has_break_before = False
                        if blank_start_idx > 0 and blank_start_idx - 1 < len(document.paragraphs):
                            prev_para = document.paragraphs[blank_start_idx - 1]
                            if prev_para.paragraph_format.page_break_before or any('w:br' in str(run.element.xml) and 'type="page"' in str(run.element.xml) for run in prev_para.runs if hasattr(run, 'element')):
                                has_break_before = True
                        
                        has_break_after = False
                        if idx + 1 < len(document.paragraphs):
                            next_para = document.paragraphs[idx + 1]
                            if next_para.paragraph_format.page_break_before or any('w:br' in str(run.element.xml) and 'type="page"' in str(run.element.xml) for run in next_para.runs if hasattr(run, 'element')):
                                has_break_after = True
                        
                        # 如果前后有分页符，或者连续空白段落很多，说明是空白页，删除这些空白段落
                        if has_break_before or has_break_after or consecutive_blanks >= BLANK_PAGE_THRESHOLD:
                            self._log_to_file(f"[空白页检测] 确认英文摘要后有空白页，has_break_before={has_break_before}, has_break_after={has_break_after}, consecutive_blanks={consecutive_blanks}")
                            delete_end = min(blank_start_idx + consecutive_blanks - 1, len(document.paragraphs) - 1)
                            deleted_count = 0
                            for delete_idx in range(delete_end, blank_start_idx - 1, -1):
                                if delete_idx >= 0 and delete_idx < len(document.paragraphs):
                                    para_to_delete = document.paragraphs[delete_idx]
                                    if len(para_to_delete.text.strip()) == 0 and not has_page_break(para_to_delete):
                                        para_to_delete._element.getparent().remove(para_to_delete._element)
                                        deleted_count += 1
                                        if delete_idx < idx:
                                            idx -= 1
                            
                            if deleted_count > 0:
                                self._log_to_file(f"[空白页检测] ✅ 已删除英文摘要后的 {deleted_count} 个空白段落（空白页），从段落 {blank_start_idx} 到 {delete_end}")
                                issues.append({
                                    "type": "blank_page_removed",
                                    "message": f"已删除英文摘要后的 {deleted_count} 个空白段落（空白页）",
                                    "suggestion": "已自动删除空白页",
                                    "blank_start": blank_start_idx,
                                    "blank_count": deleted_count,
                                })
                                consecutive_blanks = 0
                                blank_start_idx = None
                                continue
                    idx += 1
                    continue
                else:
                    # 遇到非空白段落，检查之前是否有大量空白需要删除
                    if consecutive_blanks >= BLANK_PAGE_WITH_HEADER_THRESHOLD and blank_start_idx is not None:
                        # 检查空白段落前是否有分页符
                        has_break_before = False
                        if blank_start_idx > 0 and blank_start_idx - 1 < len(document.paragraphs):
                            prev_para = document.paragraphs[blank_start_idx - 1]
                            if prev_para.paragraph_format.page_break_before or any('w:br' in str(run.element.xml) and 'type="page"' in str(run.element.xml) for run in prev_para.runs if hasattr(run, 'element')):
                                has_break_before = True
                        
                        # 如果前面有分页符，或者连续空白段落很多，说明是空白页，删除这些空白段落
                        if has_break_before or consecutive_blanks >= BLANK_PAGE_THRESHOLD:
                            delete_end = min(blank_start_idx + consecutive_blanks - 1, len(document.paragraphs) - 1)
                            deleted_count = 0
                            for delete_idx in range(delete_end, blank_start_idx - 1, -1):
                                if delete_idx >= 0 and delete_idx < len(document.paragraphs):
                                    para_to_delete = document.paragraphs[delete_idx]
                                    if len(para_to_delete.text.strip()) == 0 and not has_page_break(para_to_delete):
                                        para_to_delete._element.getparent().remove(para_to_delete._element)
                                        deleted_count += 1
                                        if delete_idx < idx:
                                            idx -= 1
                            
                            if deleted_count > 0:
                                issues.append({
                                    "type": "blank_page_removed",
                                    "message": f"已删除英文摘要后的 {deleted_count} 个空白段落（空白页）",
                                    "suggestion": "已自动删除空白页",
                                    "blank_start": blank_start_idx,
                                    "blank_count": deleted_count,
                                })
                                consecutive_blanks = 0
                                blank_start_idx = None
                                continue
                    # 重置计数
                    consecutive_blanks = 0
                    blank_start_idx = None
            
            # 检查是否是空白段落
            is_blank = len(para_text) == 0
            
            if is_blank:
                if consecutive_blanks == 0:
                    blank_start_idx = idx
                consecutive_blanks += 1
                idx += 1
            else:
                # 遇到非空白段落
                # 如果之前有大量连续空白（可能是整页空白），检查是否需要删除
                # 或者有中等数量的空白且前后有分页符（可能是只有页眉的空白页）
                if (consecutive_blanks >= BLANK_PAGE_THRESHOLD or 
                    (consecutive_blanks >= BLANK_PAGE_WITH_HEADER_THRESHOLD and blank_start_idx is not None)) and blank_start_idx is not None:
                    # 检查这些空白段落前后是否有分页符，如果有，可能是整页空白
                    # 检查空白段落之前是否有分页符
                    has_break_before = False
                    if blank_start_idx > 0 and blank_start_idx - 1 < len(document.paragraphs):
                        prev_para = document.paragraphs[blank_start_idx - 1]
                        if prev_para.paragraph_format.page_break_before:
                            has_break_before = True
                        else:
                            for run in prev_para.runs:
                                if hasattr(run, 'element'):
                                    run_xml = str(run.element.xml)
                                    if 'w:br' in run_xml and 'type="page"' in run_xml:
                                        has_break_before = True
                                        break
                    
                    # 检查空白段落之后是否有分页符
                    has_break_after = False
                    if idx < len(document.paragraphs):
                        next_para = document.paragraphs[idx]
                        if next_para.paragraph_format.page_break_before:
                            has_break_after = True
                        else:
                            for run in next_para.runs:
                                if hasattr(run, 'element'):
                                    run_xml = str(run.element.xml)
                                    if 'w:br' in run_xml and 'type="page"' in run_xml:
                                        has_break_after = True
                                        break
                    
                    # 如果空白段落前后都有分页符，或者空白段落数量非常多，可能是整页空白
                    # 或者有中等数量的空白且前后有分页符（可能是只有页眉的空白页）
                    # 删除这些空白段落，但保留最后一个，避免导致新的整页空白
                    should_delete = False
                    if consecutive_blanks >= BLANK_PAGE_THRESHOLD * 2:
                        # 空白段落数量非常多，肯定是整页空白
                        should_delete = True
                    elif consecutive_blanks >= BLANK_PAGE_THRESHOLD:
                        # 空白段落数量多，可能是整页空白
                        should_delete = True
                    elif consecutive_blanks >= BLANK_PAGE_WITH_HEADER_THRESHOLD and (has_break_before or has_break_after):
                        # 中等数量的空白且前后有分页符，可能是只有页眉的空白页
                        should_delete = True
                    
                    if should_delete:
                        # 删除空白段落，但保留最后一个
                        deleted_count = 0
                        # 从后往前删除，保留最后一个空白段落
                        delete_end = blank_start_idx + consecutive_blanks - 1
                        # 确保索引在有效范围内
                        delete_end = min(delete_end, len(document.paragraphs) - 1)
                        for delete_idx in range(delete_end, blank_start_idx, -1):
                            if delete_idx >= 0 and delete_idx < len(document.paragraphs):
                                para_to_delete = document.paragraphs[delete_idx]
                                if len(para_to_delete.text.strip()) == 0:
                                    # 检查是否包含字段代码
                                    para_xml = para_to_delete._element.xml if hasattr(para_to_delete._element, 'xml') else ""
                                    if 'TOC' in para_xml or 'w:fldChar' in para_xml or 'w:instrText' in para_xml:
                                        continue
                                    # 检查是否包含分页符，如果包含则不删除（避免导致空白页）
                                    if has_page_break(para_to_delete):
                                        continue
                                    para_to_delete._element.getparent().remove(para_to_delete._element)
                                    deleted_count += 1
                                    # 如果删除的段落在当前索引之前，需要调整索引
                                    if delete_idx < idx:
                                        idx -= 1
                        
                        if deleted_count > 0:
                            issues.append({
                                "type": "blank_page_removed",
                                "message": f"已删除第 {blank_start_idx + 1} 段到第 {blank_start_idx + consecutive_blanks} 段之间的 {deleted_count} 个空白段落（疑似整页空白）",
                                "suggestion": "已自动删除整页空白页",
                                "blank_start": blank_start_idx,
                                "blank_count": deleted_count,
                            })
                            # 删除段落后，重新从删除位置开始检查
                            idx = blank_start_idx
                            consecutive_blanks = 0
                            blank_start_idx = None
                            continue
                
                consecutive_blanks = 0
                blank_start_idx = None
                idx += 1
        
        # 处理文档末尾的整页空白
        # 检查末尾是否有分页符，如果有，可能是只有页眉的空白页
        has_break_before_end = False
        if blank_start_idx is not None and blank_start_idx > 0 and blank_start_idx - 1 < len(document.paragraphs):
            prev_para = document.paragraphs[blank_start_idx - 1]
            if prev_para.paragraph_format.page_break_before:
                has_break_before_end = True
            else:
                for run in prev_para.runs:
                    if hasattr(run, 'element'):
                        run_xml = str(run.element.xml)
                        if 'w:br' in run_xml and 'type="page"' in run_xml:
                            has_break_before_end = True
                            break
        
        if ((consecutive_blanks >= BLANK_PAGE_THRESHOLD) or 
            (consecutive_blanks >= BLANK_PAGE_WITH_HEADER_THRESHOLD and has_break_before_end)) and blank_start_idx is not None:
            # 删除末尾的整页空白，但保留最后一个空白段落
            deleted_count = 0
            delete_end = min(blank_start_idx + consecutive_blanks - 1, len(document.paragraphs) - 1)
            for delete_idx in range(delete_end, blank_start_idx, -1):
                if delete_idx >= 0 and delete_idx < len(document.paragraphs):
                    para_to_delete = document.paragraphs[delete_idx]
                    if len(para_to_delete.text.strip()) == 0:
                        # 检查是否包含字段代码
                        para_xml = para_to_delete._element.xml if hasattr(para_to_delete._element, 'xml') else ""
                        if 'TOC' in para_xml or 'w:fldChar' in para_xml or 'w:instrText' in para_xml:
                            continue
                        # 检查是否包含分页符，如果包含则不删除（避免导致空白页）
                        if has_page_break(para_to_delete):
                            continue
                        para_to_delete._element.getparent().remove(para_to_delete._element)
                        deleted_count += 1
            
            if deleted_count > 0:
                issues.append({
                    "type": "blank_page_removed",
                    "message": f"已删除文档末尾第 {blank_start_idx + 1} 段到第 {blank_start_idx + consecutive_blanks} 段之间的 {deleted_count} 个空白段落（疑似整页空白）",
                    "suggestion": "已自动删除整页空白页",
                    "blank_start": blank_start_idx,
                    "blank_count": deleted_count,
                })
        
        return issues

    def _save_file_to_storage(self, key: str, content: bytes) -> bool:
        """
        保存文件到云存储
        
        Args:
            key: 存储键（路径）
            content: 文件内容（字节）
        
        Returns:
            是否成功
        """
        if not self.use_storage:
            return False
        try:
            file_obj = io.BytesIO(content)
            return self.storage.upload_file(key, file_obj)
        except Exception as e:
            print(f"[Storage] Failed to save {key}: {e}")
            return False

    def _load_file_from_storage(self, key: str) -> Optional[bytes]:
        """
        从云存储加载文件
        
        Args:
            key: 存储键（路径）
        
        Returns:
            文件内容（字节），如果不存在则返回 None
        """
        if not self.use_storage:
            return None
        try:
            return self.storage.download_file(key)
        except Exception as e:
            print(f"[Storage] Failed to load {key}: {e}")
            return None

    def _save_to_storage(self, document_id: str, files: Dict[str, Path]) -> None:
        """
        将文档文件保存到云存储
        
        Args:
            document_id: 文档ID
            files: 文件路径字典
        """
        if not self.use_storage:
            return
        
        prefix = f"documents/{document_id}"
        
        # 保存所有文件
        for file_type, file_path in files.items():
            if file_path.exists():
                # 对于PDF文件，确保使用正确的扩展名
                if file_type == "pdf":
                    key = f"{prefix}/pdf.pdf"
                else:
                    key = f"{prefix}/{file_type}.{file_path.suffix[1:]}"  # 去掉点号
                
                file_size = file_path.stat().st_size
                print(f"[Storage] 准备上传文件: {file_type} -> {key}, 大小: {file_size / 1024:.2f} KB")
                content = file_path.read_bytes()
                if self._save_file_to_storage(key, content):
                    print(f"[Storage] ✅ 成功上传: {key}")
                else:
                    print(f"[Storage] ❌ 上传失败: {key}")
            else:
                print(f"[Storage] ⚠️ 文件不存在，跳过上传: {file_type} -> {file_path}")

    def _get_file_from_storage_or_local(self, document_id: str, file_type: str, extension: str, local_path: Path) -> Optional[Path]:
        """
        从云存储或本地文件系统获取文件
        
        Args:
            document_id: 文档ID
            file_type: 文件类型（original, final, preview, html, report, pdf）
            extension: 文件扩展名（docx, html, json, pdf）
            local_path: 本地文件路径（用于回退）
        
        Returns:
            文件路径（如果找到），否则返回 None
        """
        # 优先从云存储读取
        if self.use_storage:
            key = f"documents/{document_id}/{file_type}.{extension}"
            print(f"[Storage] 查找文件: key={key}, file_type={file_type}, extension={extension}")
            if self.storage.file_exists(key):
                print(f"[Storage] 文件存在于云存储: {key}")
                content = self._load_file_from_storage(key)
                if content:
                    print(f"[Storage] 成功下载文件: {key}, 大小: {len(content) / 1024:.2f} KB")
                    # 确保本地目录存在
                    local_path.parent.mkdir(parents=True, exist_ok=True)
                    # 写入本地临时文件
                    local_path.write_bytes(content)
                    print(f"[Storage] 已保存到本地: {local_path}")
                    return local_path
                else:
                    print(f"[Storage] ⚠️ 文件存在但下载失败: {key}")
            else:
                print(f"[Storage] ⚠️ 文件不存在于云存储: {key}")
        
        # 回退到本地文件系统
        print(f"[Storage] 检查本地文件: {local_path}")
        if local_path.exists():
            local_size = local_path.stat().st_size
            print(f"[Storage] ✅ 找到本地文件: {local_path}, 大小: {local_size / 1024:.2f} KB")
            return local_path
        else:
            print(f"[Storage] ⚠️ 本地文件不存在: {local_path}")
        
        return None

    def _add_pdf_watermarks(self, pdf_path: Path, output_path: Path, watermark_text: str = "www.geshixiugai.cn", watermarks_per_page: int = 10) -> bool:
        """在PDF的每一页添加多个水印
        
        Args:
            pdf_path: 原始PDF文件路径
            output_path: 输出PDF文件路径
            watermark_text: 水印文本
            watermarks_per_page: 每页水印数量
        
        Returns:
            bool: 是否成功
        """
        try:
            from pypdf import PdfReader, PdfWriter
            from reportlab.pdfgen import canvas
            from reportlab.lib.colors import red
            import io
            import math
            
            print(f"[PDF水印] 开始为PDF添加水印: {pdf_path}")
            print(f"[PDF水印] 水印文本: {watermark_text}, 每页水印数: {watermarks_per_page}")
            
            # 读取原始PDF
            reader = PdfReader(str(pdf_path))
            writer = PdfWriter()
            
            # 获取PDF页数
            num_pages = len(reader.pages)
            print(f"[PDF水印] PDF总页数: {num_pages}")
            
            # 为每一页添加水印，同时检查并删除空白页
            pages_to_keep = []
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                
                # 获取页面尺寸
                page_box = page.mediabox
                page_width = float(page_box.width)
                page_height = float(page_box.height)
                
                # 检查页面是否是空白页
                # 提取页面文本内容
                try:
                    page_text = page.extract_text()
                    # 如果页面文本为空或只有空白字符，可能是空白页
                    # 但也要考虑页眉页脚，所以如果文本长度小于10个字符，认为是空白页
                    if page_text and len(page_text.strip()) > 10:
                        # 页面有内容，保留
                        pages_to_keep.append(page_num)
                        print(f"[PDF水印] 第 {page_num + 1} 页有内容，保留")
                    else:
                        # 页面可能是空白页，检查是否有图像或其他内容
                        # 如果页面有图像或其他对象，也保留
                        if '/XObject' in page.get('/Resources', {}):
                            pages_to_keep.append(page_num)
                            print(f"[PDF水印] 第 {page_num + 1} 页有图像，保留")
                        else:
                            print(f"[PDF水印] 第 {page_num + 1} 页是空白页，将删除")
                except Exception as e:
                    # 如果提取文本失败，保留页面（可能是扫描件或特殊格式）
                    pages_to_keep.append(page_num)
                    print(f"[PDF水印] 第 {page_num + 1} 页提取文本失败，保留: {e}")
                
                # 如果页面需要保留，添加水印
                if page_num in pages_to_keep:
                    print(f"[PDF水印] 处理第 {page_num + 1} 页, 尺寸: {page_width}x{page_height}")
                    
                    # 创建水印PDF（使用reportlab）
                    watermark_pdf = io.BytesIO()
                    c = canvas.Canvas(watermark_pdf, pagesize=(page_width, page_height))
                    
                    # 设置水印样式 - 浅红色、半透明、水平放置
                    # 使用浅红色（RGB: 255, 200, 200）并设置透明度
                    from reportlab.lib.colors import Color
                    light_red = Color(1.0, 0.78, 0.78, alpha=0.3)  # 浅红色，30%透明度
                    c.setFillColor(light_red)
                    
                    # 根据页面大小计算字体大小，适中即可
                    font_size = max(30, int(page_width / 20))
                    c.setFont("Helvetica-Bold", font_size)
                    
                    # 计算文本宽度（用于居中显示）
                    text_width = c.stringWidth(watermark_text, "Helvetica-Bold", font_size)
                    
                    # 每页水平放置3个水印，均匀分布在A4纸上
                    num_watermarks = 3
                    
                    # 计算每个水印的位置（水平均匀分布）
                    # 留出边距，确保水印不会太靠近边缘
                    margin_x = page_width / 10
                    margin_y = page_height / 10
                    usable_width = page_width - 2 * margin_x
                    usable_height = page_height - 2 * margin_y
                    
                    # 计算水平间距（3个水印，4个间隔，水平方向均匀分布）
                    x_spacing = usable_width / (num_watermarks + 1)
                    
                    # 垂直位置：在页面的上、中、下三个位置均匀分布
                    y_positions = [
                        margin_y + usable_height * 0.25,  # 上1/4位置
                        margin_y + usable_height * 0.5,   # 中间位置
                        margin_y + usable_height * 0.75    # 下3/4位置
                    ]
                    
                    # 添加3个水印，水平放置，均匀分布
                    for i in range(num_watermarks):
                        # 计算水平位置（均匀分布）
                        x = margin_x + (i + 1) * x_spacing
                        # 计算垂直位置（上、中、下均匀分布）
                        y = y_positions[i]
                        
                        # 绘制水印文本（水平放置，不旋转）
                        c.saveState()
                        c.translate(x, y)
                        # 不旋转，保持水平
                        # 使用浅红色，半透明
                        c.setFillColor(light_red)
                        # 居中显示文本
                        c.drawString(-text_width / 2, 0, watermark_text)
                        c.restoreState()
                    
                    watermark_count = num_watermarks
                    
                    c.save()
                    watermark_pdf.seek(0)
                    
                    # 读取水印PDF
                    watermark_reader = PdfReader(watermark_pdf)
                    watermark_page = watermark_reader.pages[0]
                    
                    # 合并水印到原页面
                    page.merge_page(watermark_page)
                    
                    # 添加到输出PDF
                    writer.add_page(page)
                    
                    print(f"[PDF水印] 第 {page_num + 1} 页水印添加完成（共 {watermark_count} 个水印）")
            
            # 保存输出PDF
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            # 统计删除的空白页
            deleted_pages = num_pages - len(pages_to_keep)
            if deleted_pages > 0:
                print(f"[PDF水印] ✅ 已删除 {deleted_pages} 个空白页")
            print(f"[PDF水印] ✅ 最终PDF页数: {len(pages_to_keep)} (原始: {num_pages})")
            
            output_size = output_path.stat().st_size
            print(f"[PDF水印] ✅ PDF水印添加成功: {output_path}, 大小: {output_size / 1024:.2f} KB")
            return True
            
        except ImportError as e:
            print(f"[PDF水印] ❌ 缺少必要的库: {e}")
            print(f"[PDF水印] 请安装: pip install reportlab")
            return False
        except Exception as e:
            print(f"[PDF水印] ❌ 添加水印失败: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _generate_watermarked_preview(self, final_path: Path, preview_path: Path) -> None:
        shutil.copy2(final_path, preview_path)
        document = Document(preview_path)
        watermark_text = "预览版 仅供查看"
        
        # 创建VML水印形状，设置为背景层，难以删除
        ns = (
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office"'
        )
        shape_template = (
            f'<w:pict {ns}>'
            '<v:shape id="watermark" o:spid="_x0000_s1025" type="#_x0000_t136" '
            'style="position:absolute;margin-left:0;margin-top:0;width:468pt;height:600pt;'
            'rotation:315;opacity:0.15;z-index:-251654144;mso-position-horizontal:center;'
            'mso-position-vertical:center;mso-wrap-style:none;mso-wrap-distance-left:0;'
            'mso-wrap-distance-right:0;mso-wrap-distance-top:0;mso-wrap-distance-bottom:0;">'
            '<v:fill opacity="0"/>'
            '<v:stroke color="#d10f0f"/>'
            f'<v:textpath style="font-family:微软雅黑;font-size:72pt;font-weight:bold" string="{watermark_text}"/>'
            '<o:lock v:ext="edit" rotation="t" text="t" aspectratio="t"/>'
            '</v:shape>'
            '</w:pict>'
        )
        
        # 方法1: 在页眉中添加水印（覆盖所有页面）
        for section in document.sections:
            header = section.header
            if header.is_linked_to_previous:
                header.is_linked_to_previous = False
            # 清空现有页眉内容
            for para in header.paragraphs:
                para.clear()
            paragraph = header.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            run._r.append(parse_xml(shape_template))
        
        # 方法2: 在正文的每个段落中嵌入水印（作为背景层）
        # 每隔几个段落插入一次，避免文档过大
        watermark_interval = max(1, len(document.paragraphs) // 20)  # 大约20个水印
        for i, paragraph in enumerate(document.paragraphs):
            # 跳过空段落和标题段落
            if not paragraph.text.strip() or len(paragraph.text.strip()) < 3:
                continue
            # 每隔一定间隔插入水印
            if i % watermark_interval == 0:
                # 在段落开头插入水印形状
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                # 创建独立的水印形状，位置相对于段落
                para_shape = (
                    f'<w:pict {ns}>'
                    '<v:shape id="watermark_para" o:spid="_x0000_s1026" type="#_x0000_t136" '
                    'style="position:absolute;margin-left:0;margin-top:0;width:400pt;height:400pt;'
                    'rotation:315;opacity:0.12;z-index:-251654144;mso-position-horizontal:center;'
                    'mso-position-vertical:center;mso-wrap-style:none;">'
                    '<v:fill opacity="0"/>'
                    '<v:stroke color="#d10f0f"/>'
                    f'<v:textpath style="font-family:微软雅黑;font-size:60pt;font-weight:bold" string="{watermark_text}"/>'
                    '<o:lock v:ext="edit" rotation="t" text="t" aspectratio="t"/>'
                    '</v:shape>'
                    '</w:pict>'
                )
                run._r.append(parse_xml(para_shape))
        
        document.save(preview_path)
    
    def _generate_html_preview(self, docx_path: Path, html_path: Path, stats: Dict) -> None:
        """将Word文档转换为HTML预览，尽量保持与原文档一致"""
        print(f"[HTML预览] 开始生成HTML预览，输入文件: {docx_path}")
        print(f"[HTML预览] 输出文件: {html_path}")
        
        # 优先尝试使用LibreOffice转换（保留格式最好）
        print("[HTML预览] 尝试使用LibreOffice转换...")
        if self._try_libreoffice_conversion(docx_path, html_path, stats):
            print("[HTML预览] 使用LibreOffice转换成功")
            return
        
        # 回退到自定义HTML生成
        print("[HTML预览] LibreOffice不可用，使用自定义HTML生成")
        print(f"[HTML预览] 正在读取Word文档: {docx_path}")
        document = Document(docx_path)
        print(f"[HTML预览] Word文档读取成功，总段落数: {len(document.paragraphs)}")
        
        # 不再在预览文档中添加检测结果，保持文档干净
        # 检测结果只在首页（报告）中显示
        
        html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>文档预览 - 预览版</title>
    <style>
        body {{
            font-family: "SimSun", "宋体", "Times New Roman", serif;
            max-width: 210mm;
            margin: 20px auto;
            padding: 20px;
            background: #f5f5f5;
            line-height: 1.6;
        }}
        .document-container {{
            background: white;
            padding: 40px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            min-height: 297mm;
        }}
        .watermark {{
            position: fixed;
            top: 50%;
            left: 50%;
            /* 移除transform，避免WeasyPrint内部错误 */
            margin-top: -36px;
            margin-left: -200px;
            font-size: 72px;
            color: rgba(209, 15, 15, 0.15);
            font-weight: bold;
            pointer-events: none;
            z-index: 1;
            white-space: nowrap;
        }}
        p {{
            margin: 0;
            padding: 0;
            position: relative;
            z-index: 2;
            /* 默认格式会被内联样式覆盖 */
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin: 20px 0 10px 0;
            position: relative;
            z-index: 2;
        }}
        h1 {{ font-size: 18pt; font-weight: bold; text-align: center; }}
        h2 {{ font-size: 16pt; font-weight: bold; }}
        h3 {{ font-size: 14pt; font-weight: bold; }}
        .center {{ text-align: center; }}
        .bold {{ font-weight: bold; }}
        .no-indent {{ text-indent: 0; }}
        .warning {{
            background: #fff3cd;
            border: 1px solid #ffc107;
            padding: 15px;
            margin: 20px 0;
            border-radius: 5px;
            text-align: center;
            font-weight: bold;
            color: #856404;
        }}
        @media print {{
            body {{ background: white; }}
            .document-container {{ box-shadow: none; }}
        }}
    </style>
</head>
<body>
    <div class="watermark">预览版 仅供查看</div>
    <div class="document-container">
"""
        
        paragraph_count = 0
        total_text_length = 0
        chinese_char_count = 0
        print(f"[HTML预览] 开始处理 {len(document.paragraphs)} 个段落...")
        for idx, paragraph in enumerate(document.paragraphs):
            # 每处理100个段落输出一次进度
            if idx > 0 and idx % 100 == 0:
                print(f"[HTML预览] 已处理 {idx}/{len(document.paragraphs)} 个段落...")
            # 改进文字提取：优先使用 paragraph.text，如果为空则从 runs 中提取
            text = paragraph.text.strip()
            if not text:
                # 如果 paragraph.text 为空，尝试从 runs 中提取所有文字
                text = "".join([run.text for run in paragraph.runs if run.text]).strip()
            
            # 统计文字长度和中文字符数量
            if text:
                total_text_length += len(text)
                # 统计中文字符（Unicode范围：\u4e00-\u9fff）
                chinese_chars = [c for c in text if '\u4e00' <= c <= '\u9fff']
                chinese_char_count += len(chinese_chars)
                # 调试：记录前几个包含中文的段落
                if chinese_chars and idx < 5:
                    print(f"[HTML预览] 段落 {idx} 包含中文: {text[:50]}... (中文字符数: {len(chinese_chars)})")
            
            # 检查段落格式中是否有分页符
            # python-docx中，分页符通常通过paragraph_format.page_break_before或runs中的break元素表示
            page_break_before = False
            if paragraph.paragraph_format.page_break_before:
                page_break_before = True
                print(f"[HTML预览] 检测到分页符（段落 {idx}）")
            
            # 检查runs中是否有分页符
            for run in paragraph.runs:
                if hasattr(run, 'element'):
                    run_xml = str(run.element.xml)
                    if 'w:br' in run_xml and 'type="page"' in run_xml:
                        page_break_before = True
                        print(f"[HTML预览] 检测到run中的分页符（段落 {idx}）")
                        break
            
            # 如果检测到分页符，添加分页标记
            # 注意：在浏览器预览中，page-break-before: always 会显示为空白页
            # 因此只在打印或PDF生成时使用分页符，浏览器预览中不添加空白页
            # 如果需要，可以添加一个轻微的分隔线来标识分页位置（但不占用空间）
            if page_break_before:
                # 使用CSS类，在浏览器中不显示空白，只在打印/PDF时生效
                html_content += '<div class="page-break" style="display: none;"></div>\n'
            
            # 检查段落是否包含图片
            has_image = self._paragraph_has_image_or_equation(paragraph)
            images_html = ""
            
            if has_image:
                # 提取段落中的图片
                if idx < 5 or idx % 50 == 0:  # 只记录前5个或每50个
                    print(f"[HTML预览] 段落 {idx} 包含图片，正在提取...")
                images_html = self._extract_images_from_paragraph(paragraph, document)
                if idx < 5 or idx % 50 == 0:
                    print(f"[HTML预览] 段落 {idx} 图片提取完成，HTML长度: {len(images_html)} 字符")
                # 如果检测到图片但提取失败，记录警告
                if not images_html:
                    print(f"[HTML预览] ⚠️ 警告：段落 {idx} 检测到图片但提取失败！")
                # 如果提取到图片但HTML为空，记录警告
                if not images_html:
                    print(f"[HTML预览] ⚠️ 警告：段落 {idx} 检测到图片但提取失败！")
            
            # 判断段落样式（提前定义，避免作用域错误）
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # 如果既没有文本也没有图片，跳过（但保留空段落以维持格式）
            if not text and not images_html:
                html_content += "<p>&nbsp;</p>\n"
                continue
            
            # 调试：记录段落信息
            if idx < 10 or (text and len(text) > 0):  # 只记录前10个段落或有文字的段落
                print(f"[HTML预览] 段落 {idx}: 文字长度={len(text)}, 有图片={bool(images_html)}, 样式={style_name}")
            alignment = paragraph.alignment
            
            # 构建样式
            style_attrs = []
            classes = []
            
            # 判断是否是标题（用于确定字体）
            is_heading_para = False
            if "Heading" in style_name or "标题" in style_name:
                is_heading_para = True
            elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and len(text) <= 20:
                # 居中对齐的短文本可能是标题
                is_heading_para = True
            elif text and text[0].isdigit() and len(text) <= 20:
                # 以数字开头的短文本可能是标题
                if re.match(r'^(\d+\.\d+\.\d+|\d+\.\d+|\d+)([，,。.：:；;]?)$', text):
                    is_heading_para = True
            
            if is_heading_para:
                level = 1
                if "1" in style_name or "一" in style_name:
                    level = 1
                elif "2" in style_name or "二" in style_name:
                    level = 2
                elif "3" in style_name or "三" in style_name:
                    level = 3
                else:
                    level = 2
                # 转义标题文字（只转义特殊字符，保留中文）
                if text:
                    escaped_title = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                else:
                    escaped_title = ""
                html_content += f"<h{level}>{escaped_title}</h{level}>\n"
                if images_html:
                    html_content += f"<div style='text-align: center; margin: 10px 0;'>{images_html}</div>\n"
            else:
                # 普通段落 - 从Word文档中提取完整的格式信息
                # 提取段落格式
                para_format = docx_format_utils.extract_paragraph_format(paragraph)
                
                # 应用对齐方式
                if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    style_attrs.append("text-align: center;")
                elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                    style_attrs.append("text-align: right;")
                elif alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                    style_attrs.append("text-align: justify;")
                else:
                    style_attrs.append("text-align: left;")
                
                # 应用字体和字号（从runs中提取）
                font_name = para_format.get("font_name")
                font_size = para_format.get("font_size")
                
                # 如果从runs中提取不到字体，尝试从段落样式中获取
                if not font_name and paragraph.style:
                    try:
                        # 尝试从段落样式的字体设置中获取
                        style_font = paragraph.style.font
                        if style_font and style_font.name:
                            font_name = style_font.name
                            if idx < 10:
                                print(f"[HTML预览] 段落 {idx} 从样式获取字体: {font_name}")
                    except:
                        pass
                
                # 根据提取到的字体信息，确定最终使用的字体
                # 支持：黑体、宋体、楷体、Times New Roman
                font_family_css = None
                
                if font_name:
                    font_name_lower = font_name.lower()
                    # 黑体识别
                    if "黑" in font_name or "simhei" in font_name_lower or "hei" in font_name_lower or "heiti" in font_name_lower:
                        font_family_css = '"SimHei", "黑体", "STHeiti", "WenQuanYi Micro Hei", "WenQuanYi Zen Hei", sans-serif'
                        if idx < 10:
                            print(f"[HTML预览] 段落 {idx} 字体: {font_name} -> 黑体")
                    # 宋体识别
                    elif "宋" in font_name or "simsun" in font_name_lower or "song" in font_name_lower or "songti" in font_name_lower:
                        font_family_css = '"SimSun", "宋体", "STSong", "STSongti-SC-Regular", "WenQuanYi Micro Hei", "WenQuanYi Zen Hei", serif'
                        if idx < 10:
                            print(f"[HTML预览] 段落 {idx} 字体: {font_name} -> 宋体")
                    # 楷体识别
                    elif "楷" in font_name or "kaiti" in font_name_lower or "kai" in font_name_lower:
                        font_family_css = '"KaiTi", "楷体", "STKaiti", "WenQuanYi Micro Hei", "WenQuanYi Zen Hei", serif'
                        if idx < 10:
                            print(f"[HTML预览] 段落 {idx} 字体: {font_name} -> 楷体")
                    # Times New Roman 识别
                    elif "times" in font_name_lower or "new roman" in font_name_lower or "tnr" in font_name_lower:
                        font_family_css = '"Times New Roman", "Times", "Liberation Serif", "DejaVu Serif", serif'
                        if idx < 10:
                            print(f"[HTML预览] 段落 {idx} 字体: {font_name} -> Times New Roman")
                
                # 如果没有识别到字体，根据段落类型使用默认字体
                if not font_family_css:
                    if is_heading_para:
                        # 标题默认使用黑体
                        font_family_css = '"SimHei", "黑体", "STHeiti", "WenQuanYi Micro Hei", "WenQuanYi Zen Hei", sans-serif'
                        if idx < 10:
                            print(f"[HTML预览] 段落 {idx} 标题使用默认字体：黑体（原字体: {font_name or '未提取'}）")
                    else:
                        # 正文默认使用宋体
                        font_family_css = '"SimSun", "宋体", "STSong", "STSongti-SC-Regular", "WenQuanYi Micro Hei", "WenQuanYi Zen Hei", serif'
                        if idx < 10:
                            print(f"[HTML预览] 段落 {idx} 正文使用默认字体：宋体（原字体: {font_name or '未提取'}）")
                
                style_attrs.append(f'font-family: {font_family_css};')
                
                if font_size:
                    style_attrs.append(f"font-size: {font_size}pt;")
                else:
                    # 如果没有字号，使用默认字号
                    if is_heading_para:
                        style_attrs.append("font-size: 16pt;")
                    else:
                        style_attrs.append("font-size: 12pt;")
                
                # 应用加粗
                is_bold = para_format.get("bold") or any(run.bold for run in paragraph.runs if run.bold)
                if is_bold:
                    style_attrs.append("font-weight: bold;")
                
                # 应用行距
                line_spacing = para_format.get("line_spacing")
                if line_spacing:
                    if isinstance(line_spacing, (int, float)):
                        # 固定行距（磅）
                        style_attrs.append(f"line-height: {line_spacing}pt;")
                    elif line_spacing == "single":
                        style_attrs.append("line-height: 1.0;")
                    elif line_spacing == "double":
                        style_attrs.append("line-height: 2.0;")
                    elif line_spacing == "1.5":
                        style_attrs.append("line-height: 1.5;")
                
                # 应用首行缩进
                first_line_indent = para_format.get("first_line_indent")
                if first_line_indent and first_line_indent > 0:
                    style_attrs.append(f"text-indent: {first_line_indent}pt;")
                else:
                    # 默认首行缩进2字符（24pt）
                    style_attrs.append("text-indent: 24pt;")
                
                # 应用段前段后间距
                space_before = para_format.get("space_before")
                space_after = para_format.get("space_after")
                if space_before and space_before > 0:
                    style_attrs.append(f"margin-top: {space_before}pt;")
                if space_after and space_after > 0:
                    style_attrs.append(f"margin-bottom: {space_after}pt;")
                
                # 应用左右缩进
                left_indent = para_format.get("left_indent")
                right_indent = para_format.get("right_indent")
                if left_indent and left_indent > 0:
                    style_attrs.append(f"margin-left: {left_indent}pt;")
                if right_indent and right_indent > 0:
                    style_attrs.append(f"margin-right: {right_indent}pt;")
                
                class_attr = f' class="{" ".join(classes)}"' if classes else ""
                style_attr = f' style="{" ".join(style_attrs)}"' if style_attrs else ""
                
                # 处理文本中的特殊字符
                # 注意：xml.sax.saxutils.escape 只转义 & < >，不会影响中文字符
                # 但为了安全，我们只转义必要的字符，保留中文字符
                if text:
                    escaped_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                else:
                    escaped_text = ""
                
                # 如果有图片，先显示图片，再显示文本
                if images_html:
                    # 图片段落通常居中显示，确保图片能正确显示
                    html_content += f'<div style="text-align: center; margin: 10px 0; page-break-inside: avoid;">{images_html}</div>\n'
                    if idx < 5:  # 只记录前5个段落的详细信息
                        print(f"[HTML预览] 段落 {idx} 添加图片HTML: {len(images_html)} 字符")
                if text:
                    # 确保文字被添加到HTML中
                    html_content += f'<p{class_attr}{style_attr}>{escaped_text}</p>\n'
        
        html_content += """    </div>
    <div class="warning">
        ⚠️ 这是预览版本，仅供查看。如需下载正式版，请完成支付。
    </div>
</body>
</html>"""
        
        print(f"[HTML预览] 段落处理完成，开始写入HTML文件...")
        # 记录统计信息
        print(f"[HTML预览] HTML生成完成，总段落数: {paragraph_count}, 总文字长度: {total_text_length} 字符")
        print(f"[HTML预览] 中文字符数: {chinese_char_count} 字符")
        print(f"[HTML预览] HTML内容大小: {len(html_content) / 1024:.2f} KB")
        # 检查HTML中是否包含中文字符
        html_chinese_count = len([c for c in html_content if '\u4e00' <= c <= '\u9fff'])
        print(f"[HTML预览] HTML中的中文字符数: {html_chinese_count} 字符")
        if chinese_char_count > 0 and html_chinese_count == 0:
            print(f"[HTML预览] ⚠️ 警告：提取到 {chinese_char_count} 个中文字符，但HTML中只有 {html_chinese_count} 个！")
        
        html_path.write_text(html_content, encoding="utf-8")
    
    def _extract_images_from_paragraph(self, paragraph, document: Document) -> str:
        """从段落中提取图片并转换为HTML img标签"""
        import zipfile
        
        images_html = ""
        image_count = 0
        
        try:
            # 获取文档的zip文件路径（docx是zip格式）
            docx_path = document.part.package
            
            # 方法1: 从runs中提取图片
            for run in paragraph.runs:
                if not hasattr(run, 'element'):
                    continue
                
                try:
                    run_xml = str(run.element.xml)
                    # 排除水印
                    if 'v:shape' in run_xml.lower() and 'textpath' in run_xml.lower():
                        continue
                    
                    # 查找图片关系ID（支持多种格式）
                    image_id = None
                    # 尝试多种方式查找图片ID
                    if 'r:embed' in run_xml:
                        # 内嵌图片
                        match = re.search(r'r:embed="([^"]+)"', run_xml)
                        if match:
                            image_id = match.group(1)
                    elif 'r:link' in run_xml:
                        # 链接图片
                        match = re.search(r'r:link="([^"]+)"', run_xml)
                        if match:
                            image_id = match.group(1)
                    # 也尝试查找a:blip中的embed属性
                    if not image_id and 'a:blip' in run_xml:
                        match = re.search(r'r:embed="([^"]+)"', run_xml)
                        if match:
                            image_id = match.group(1)
                    
                    if image_id:
                        # 从文档中提取图片数据
                        try:
                            # 尝试从多个位置获取图片
                            image_part = None
                            
                            # 方法1: 从主文档部分获取
                            if hasattr(document.part, 'related_parts') and image_id in document.part.related_parts:
                                image_part = document.part.related_parts[image_id]
                                print(f"[HTML预览] 从主文档部分找到图片: {image_id}")
                            
                            # 方法2: 从run的part获取
                            if not image_part and hasattr(run, 'part') and hasattr(run.part, 'related_parts'):
                                if image_id in run.part.related_parts:
                                    image_part = run.part.related_parts[image_id]
                                    print(f"[HTML预览] 从run.part找到图片: {image_id}")
                            
                            # 方法3: 从文档的所有部分查找
                            if not image_part:
                                # 尝试从文档的所有相关部分查找
                                for rel in document.part.rels.values():
                                    if rel.rId == image_id:
                                        image_part = rel.target_part
                                        print(f"[HTML预览] 从文档关系中找到图片: {image_id}")
                                        break
                            
                            if not image_part:
                                print(f"[HTML预览] 警告: 未找到图片关系ID: {image_id}")
                                continue
                                
                            image_data = image_part.blob
                            if not image_data:
                                print(f"[HTML预览] 警告: 图片数据为空: {image_id}")
                                continue
                            
                            # 确定图片格式
                            content_type = image_part.content_type if hasattr(image_part, 'content_type') else ''
                            
                            # 检查是否为不支持的格式（WMF、EMF等）
                            if 'wmf' in content_type.lower() or 'emf' in content_type.lower() or 'x-wmf' in content_type.lower():
                                print(f"[HTML预览] ⚠️ 跳过不支持的图片格式: {content_type} (WeasyPrint不支持WMF/EMF格式)")
                                continue
                            
                            if 'jpeg' in content_type or 'jpg' in content_type:
                                img_format = 'jpeg'
                            elif 'png' in content_type:
                                img_format = 'png'
                            elif 'gif' in content_type:
                                img_format = 'gif'
                            elif 'bmp' in content_type:
                                img_format = 'bmp'
                            elif 'webp' in content_type:
                                img_format = 'webp'
                            else:
                                # 检查文件头来确定格式
                                if image_data.startswith(b'\x89PNG'):
                                    img_format = 'png'
                                elif image_data.startswith(b'\xff\xd8\xff'):
                                    img_format = 'jpeg'
                                elif image_data.startswith(b'GIF'):
                                    img_format = 'gif'
                                elif image_data.startswith(b'BM'):
                                    img_format = 'bmp'
                                elif image_data.startswith(b'RIFF') and b'WEBP' in image_data[:12]:
                                    img_format = 'webp'
                                else:
                                    # 尝试从文件扩展名判断
                                    if hasattr(image_part, 'partname'):
                                        partname = str(image_part.partname)
                                        if '.jpg' in partname or '.jpeg' in partname:
                                            img_format = 'jpeg'
                                        elif '.png' in partname:
                                            img_format = 'png'
                                        elif '.gif' in partname:
                                            img_format = 'gif'
                                        else:
                                            print(f"[HTML预览] ⚠️ 未知图片格式: {content_type}，跳过")
                                            continue
                                    else:
                                        print(f"[HTML预览] ⚠️ 未知图片格式: {content_type}，跳过")
                                        continue
                                    
                                # 如果到这里还没有设置 img_format，使用默认值
                                if 'img_format' not in locals():
                                    img_format = 'png'  # 默认
                            
                            # 转换为base64
                            base64_data = base64.b64encode(image_data).decode('utf-8')
                            data_uri = f"data:image/{img_format};base64,{base64_data}"
                            
                            # 创建img标签
                            images_html += f'<img src="{data_uri}" style="max-width: 100%; height: auto; margin: 10px 0;" alt="图片 {image_count + 1}" />'
                            image_count += 1
                            print(f"[HTML预览] 成功提取图片 {image_count}，格式: {img_format}，大小: {len(image_data)} 字节")
                            
                        except Exception as e:
                            print(f"[HTML预览] 提取图片失败: {e}")
                            import traceback
                            print(f"[HTML预览] 错误堆栈: {traceback.format_exc()}")
                            continue
                            
                except Exception as e:
                    print(f"[HTML预览] 处理run时出错: {e}")
                    continue
            
            # 方法2: 从段落的内联形状中提取图片（即使方法1已经找到图片，也继续查找，因为一个段落可能有多个图片）
            if hasattr(paragraph, '_element'):
                try:
                    # 查找drawing元素（使用findall配合qn，而不是xpath with namespaces）
                    drawings = paragraph._element.findall('.//' + qn('w:drawing'))
                    
                    for drawing in drawings:
                        # 查找图片关系ID
                        blip_elements = drawing.findall('.//' + qn('a:blip'))
                        
                        for blip in blip_elements:
                            embed_attr = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            link_attr = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link')
                            
                            image_id = embed_attr or link_attr
                            if image_id:
                                try:
                                    # 检查是否已经处理过这个图片（避免重复）
                                    # 这里简化处理，允许重复（因为可能有不同的引用方式）
                                    
                                    # 尝试从多个位置获取图片
                                    image_part = None
                                    
                                    # 方法1: 从主文档部分获取
                                    if hasattr(document.part, 'related_parts') and image_id in document.part.related_parts:
                                        image_part = document.part.related_parts[image_id]
                                        print(f"[HTML预览] 从drawing找到图片（主文档）: {image_id}")
                                    
                                    # 方法2: 从文档的所有关系查找
                                    if not image_part:
                                        for rel in document.part.rels.values():
                                            if rel.rId == image_id:
                                                image_part = rel.target_part
                                                print(f"[HTML预览] 从drawing找到图片（关系）: {image_id}")
                                                break
                                    
                                    if not image_part:
                                        print(f"[HTML预览] 警告: 从drawing未找到图片关系ID: {image_id}")
                                        continue
                                        
                                    image_data = image_part.blob
                                    if not image_data:
                                        print(f"[HTML预览] 警告: 从drawing获取的图片数据为空: {image_id}")
                                        continue
                                    
                                    # 确定图片格式
                                    content_type = image_part.content_type if hasattr(image_part, 'content_type') else ''
                                    
                                    # 检查是否为不支持的格式（WMF、EMF等）
                                    if 'wmf' in content_type.lower() or 'emf' in content_type.lower() or 'x-wmf' in content_type.lower():
                                        print(f"[HTML预览] ⚠️ 跳过不支持的图片格式: {content_type} (WeasyPrint不支持WMF/EMF格式)")
                                        # 可以添加一个占位符图片
                                        images_html += f'<div style="border: 1px dashed #ccc; padding: 20px; text-align: center; color: #999; margin: 10px 0;">[图片格式不支持: {content_type}]</div>'
                                        continue
                                    
                                    if 'jpeg' in content_type or 'jpg' in content_type:
                                        img_format = 'jpeg'
                                    elif 'png' in content_type:
                                        img_format = 'png'
                                    elif 'gif' in content_type:
                                        img_format = 'gif'
                                    elif 'bmp' in content_type:
                                        img_format = 'bmp'
                                    elif 'webp' in content_type:
                                        img_format = 'webp'
                                    else:
                                        # 检查文件头来确定格式
                                        if image_data.startswith(b'\x89PNG'):
                                            img_format = 'png'
                                        elif image_data.startswith(b'\xff\xd8\xff'):
                                            img_format = 'jpeg'
                                        elif image_data.startswith(b'GIF'):
                                            img_format = 'gif'
                                        elif image_data.startswith(b'BM'):
                                            img_format = 'bmp'
                                        elif image_data.startswith(b'RIFF') and b'WEBP' in image_data[:12]:
                                            img_format = 'webp'
                                        else:
                                            print(f"[HTML预览] ⚠️ 未知图片格式: {content_type}，跳过")
                                            images_html += f'<div style="border: 1px dashed #ccc; padding: 20px; text-align: center; color: #999; margin: 10px 0;">[图片格式未知: {content_type}]</div>'
                                            continue
                                    
                                    # 转换为base64
                                    base64_data = base64.b64encode(image_data).decode('utf-8')
                                    data_uri = f"data:image/{img_format};base64,{base64_data}"
                                    
                                    # 创建img标签
                                    images_html += f'<img src="{data_uri}" style="max-width: 100%; height: auto; margin: 10px 0;" alt="图片 {image_count + 1}" />'
                                    image_count += 1
                                    print(f"[HTML预览] 从drawing成功提取图片 {image_count}，格式: {img_format}，大小: {len(image_data)} 字节")
                                    
                                except Exception as e:
                                    print(f"[HTML预览] 从drawing提取图片失败: {e}")
                                    import traceback
                                    print(f"[HTML预览] 错误堆栈: {traceback.format_exc()}")
                                    continue
                                    
                except Exception as e:
                    print(f"[HTML预览] 处理drawing时出错: {e}")
                    import traceback
                    print(f"[HTML预览] 错误堆栈: {traceback.format_exc()}")
                    pass
            
            # 方法3: 如果前两种方法都没找到图片，尝试直接从zip文件中提取
            # 这适用于某些特殊格式的图片或关系ID查找失败的情况
            if not images_html and hasattr(document, 'part') and hasattr(document.part, 'package'):
                try:
                    # 获取docx文件的路径
                    docx_file_path = None
                    if hasattr(document.part.package, 'name'):
                        docx_file_path = document.part.package.name
                    elif hasattr(document.part.package, '__file__'):
                        docx_file_path = document.part.package.__file__
                    
                    if docx_file_path:
                        import zipfile
                        from pathlib import Path
                        
                        docx_path = Path(docx_file_path)
                        if docx_path.exists() and docx_path.suffix.lower() == '.docx':
                            print(f"[HTML预览] 尝试从zip文件直接提取图片: {docx_path}")
                            
                            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                                # 查找所有图片文件（通常在word/media/目录下）
                                image_files = [f for f in zip_ref.namelist() 
                                             if f.startswith('word/media/') and 
                                             any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'])]
                                
                                print(f"[HTML预览] 在zip文件中找到 {len(image_files)} 个图片文件")
                                
                                # 尝试从段落XML中查找引用的图片文件名
                                para_xml = str(paragraph._element.xml) if hasattr(paragraph, '_element') else ''
                                
                                for img_file in image_files:
                                    # 检查这个图片是否可能属于当前段落
                                    # 通过检查图片文件名是否在段落XML中被引用
                                    img_filename = Path(img_file).name
                                    
                                    # 如果段落包含drawing或图片相关元素，尝试匹配
                                    if ('drawing' in para_xml.lower() or 'pic:pic' in para_xml.lower() or 'a:blip' in para_xml.lower()):
                                        try:
                                            # 读取图片数据
                                            image_data = zip_ref.read(img_file)
                                            
                                            # 确定图片格式
                                            img_format = 'png'  # 默认
                                            if img_file.lower().endswith('.jpg') or img_file.lower().endswith('.jpeg'):
                                                img_format = 'jpeg'
                                            elif img_file.lower().endswith('.png'):
                                                img_format = 'png'
                                            elif img_file.lower().endswith('.gif'):
                                                img_format = 'gif'
                                            elif img_file.lower().endswith('.bmp'):
                                                img_format = 'bmp'
                                            elif img_file.lower().endswith('.webp'):
                                                img_format = 'webp'
                                            
                                            # 转换为base64
                                            base64_data = base64.b64encode(image_data).decode('utf-8')
                                            data_uri = f"data:image/{img_format};base64,{base64_data}"
                                            
                                            # 创建img标签（只添加一次，避免重复）
                                            if img_filename not in images_html:
                                                images_html += f'<img src="{data_uri}" style="max-width: 100%; height: auto; margin: 10px 0;" alt="图片 {image_count + 1}" />'
                                                image_count += 1
                                                print(f"[HTML预览] 从zip文件成功提取图片 {image_count}: {img_filename}，格式: {img_format}，大小: {len(image_data)} 字节")
                                                
                                                # 如果已经找到一个图片，就停止（避免一个段落显示多个图片）
                                                # 如果需要显示多个图片，可以移除这个break
                                                if image_count >= 1:
                                                    break
                                                    
                                        except Exception as e:
                                            print(f"[HTML预览] 从zip文件读取图片失败 {img_file}: {e}")
                                            continue
                                            
                except Exception as e:
                    print(f"[HTML预览] 从zip文件提取图片时出错: {e}")
                    import traceback
                    print(f"[HTML预览] 错误堆栈: {traceback.format_exc()}")
                    pass
        
        except Exception as e:
            print(f"[HTML预览] 提取图片时发生错误: {e}")
            import traceback
            print(f"[HTML预览] 错误堆栈: {traceback.format_exc()}")
        
        if images_html:
            print(f"[HTML预览] 段落图片提取完成，共提取 {image_count} 张图片")
        else:
            # 如果没找到图片，但段落包含drawing元素，记录警告
            if hasattr(paragraph, '_element'):
                para_xml = str(paragraph._element.xml)
                if 'drawing' in para_xml.lower() or 'pic:pic' in para_xml.lower():
                    print(f"[HTML预览] 警告: 段落包含drawing元素但未提取到图片，XML片段: {para_xml[:200]}")
        
        return images_html
    
    def _generate_pdf_preview(self, docx_path: Path, pdf_path: Path, stats: Dict) -> bool:
        """将Word文档转换为PDF预览（使用WeasyPrint从HTML转PDF）
        
        注意：LibreOffice 在某些环境下无法正常工作，因此直接使用 WeasyPrint
        """
        print(f"[PDF预览] 开始生成PDF预览，输入文件: {docx_path}, 输出文件: {pdf_path}")
        try:
            from weasyprint import HTML, CSS
            # 不再导入FontConfiguration，避免transform错误
            print("[PDF预览] WeasyPrint导入成功")
        except ImportError as e:
            print(f"[PDF预览] ❌ weasyprint未安装，跳过PDF生成: {e}")
            import traceback
            traceback.print_exc()
            return False
        
        try:
            # 先生成HTML（用于PDF转换）
            html_path = pdf_path.with_suffix('.html')
            print(f"[PDF预览] 开始生成HTML预览: {html_path}")
            self._generate_html_preview(docx_path, html_path, stats)
            
            # 检查HTML文件是否生成成功
            if not html_path.exists():
                print(f"[PDF预览] 错误: HTML文件未生成: {html_path}")
                return False
            
            # 读取Word文档的页面设置
            from docx import Document
            doc = Document(docx_path)
            page_settings = self._extract_page_settings(doc)
            
            # 读取HTML内容
            html_content = html_path.read_text(encoding='utf-8')
            
            # 生成PDF专用样式（使用Word文档的页面设置）
            pdf_css = self._generate_pdf_css(page_settings)
            
            # 在HTML的head中添加CSS和meta标签（确保UTF-8编码）
            if '</head>' in html_content:
                # 检查是否已有charset meta标签
                if 'charset' not in html_content.lower():
                    html_content = html_content.replace('</head>', '<meta charset="UTF-8">\n</head>')
                html_content = html_content.replace('</head>', f'<style>{pdf_css}</style></head>')
            else:
                # 如果没有head标签，添加一个
                if '<html' in html_content:
                    html_content = html_content.replace('<html', '<html><head><meta charset="UTF-8"><style>' + pdf_css + '</style></head>')
            
            print(f"[PDF预览] 开始转换HTML到PDF，HTML大小: {len(html_content) / 1024:.2f} KB")
            
            # 统计HTML中的图片数量（用于调试）
            import re
            img_count = len(re.findall(r'<img[^>]+>', html_content, re.IGNORECASE))
            data_uri_count = len(re.findall(r'data:image/[^;]+;base64,', html_content, re.IGNORECASE))
            print(f"[PDF预览] HTML中包含 {img_count} 个img标签，其中 {data_uri_count} 个使用data URI")
            
            # 检查HTML中的中文字符数量
            html_chinese_count = len([c for c in html_content if '\u4e00' <= c <= '\u9fff'])
            print(f"[PDF预览] HTML中的中文字符数: {html_chinese_count} 字符")
            if html_chinese_count > 0:
                print(f"[PDF预览] ✅ HTML中包含中文字符，如果PDF中看不到中文，可能是服务器缺少中文字体")
                # 显示前几个中文字符作为示例
                chinese_chars_in_html = [c for c in html_content if '\u4e00' <= c <= '\u9fff'][:10]
                if chinese_chars_in_html:
                    print(f"[PDF预览] HTML中的中文字符示例: {''.join(chinese_chars_in_html)}")
            else:
                print(f"[PDF预览] ⚠️ 警告：HTML中没有中文字符！可能是文字提取或转义时丢失了")
            
            # 使用weasyprint转换
            # 设置base_url为HTML文件所在目录，帮助weasyprint解析相对路径和data URI
            # 注意：不使用FontConfiguration()，因为它可能导致transform错误
            html_doc = HTML(
                string=html_content,
                base_url=str(html_path.parent)  # 设置base_url，帮助解析图片
            )
            
            # 检查系统可用字体（用于诊断）
            try:
                import subprocess
                result = subprocess.run(['fc-list', ':lang=zh'], capture_output=True, text=True, timeout=5)
                if result.returncode == 0:
                    fonts = result.stdout.strip().split('\n')
                    chinese_fonts = [f for f in fonts if any(keyword in f.lower() for keyword in ['song', 'simsun', '宋', 'hei', 'simhei', '黑', 'wqy', 'wenquanyi'])]
                    print(f"[PDF预览] 系统检测到 {len(chinese_fonts)} 个中文字体:")
                    for font in chinese_fonts[:5]:  # 只显示前5个
                        print(f"[PDF预览]   - {font[:100]}")
                    if len(chinese_fonts) == 0:
                        print(f"[PDF预览] ⚠️ 警告：系统未检测到中文字体！PDF可能无法正确显示中文")
                else:
                    print(f"[PDF预览] ⚠️ 无法检测系统字体（fc-list命令失败）")
            except Exception as e:
                print(f"[PDF预览] ⚠️ 字体检测失败: {e}")
            
            print(f"[PDF预览] 开始生成PDF文件...")
            # 生成PDF（不使用font_config，避免transform错误）
            # 根据WeasyPrint文档，font_config是可选的，不使用也能正常工作
            html_doc.write_pdf(
                pdf_path,
                optimize_images=False,  # 禁用图片优化，避免某些内部错误
            )
            
            pdf_size = pdf_path.stat().st_size
            print(f"[PDF预览] PDF生成成功，大小: {pdf_size / 1024:.2f} KB")
            
            # 检查PDF中的字体（使用pypdf）
            try:
                from pypdf import PdfReader
                from pypdf.generic import IndirectObject
                reader = PdfReader(str(pdf_path))
                if len(reader.pages) > 0:
                    page = reader.pages[0]
                    # 获取Resources对象，可能是IndirectObject，需要先获取实际对象
                    resources = page.get('/Resources', {})
                    if isinstance(resources, IndirectObject):
                        resources = resources.get_object()
                    
                    if resources and '/Font' in resources:
                        fonts_used = resources['/Font']
                        # 如果fonts_used是IndirectObject，也需要获取实际对象
                        if isinstance(fonts_used, IndirectObject):
                            fonts_used = fonts_used.get_object()
                        
                        if fonts_used:
                            print(f"[PDF预览] PDF中使用的字体:")
                            font_embedded_count = 0
                            font_referenced_count = 0
                            for font_name, font_obj in fonts_used.items():
                                try:
                                    # 确保font_obj是实际对象
                                    if isinstance(font_obj, IndirectObject):
                                        font_info = font_obj.get_object()
                                    else:
                                        font_info = font_obj
                                    
                                    base_font = font_info.get('/BaseFont', 'Unknown')
                                    
                                    # 检查字体是否嵌入（如果有/FontDescriptor和/FontFile，说明字体已嵌入）
                                    is_embedded = False
                                    if '/FontDescriptor' in font_info:
                                        font_desc = font_info['/FontDescriptor']
                                        # 如果font_desc是IndirectObject，需要获取实际对象
                                        if isinstance(font_desc, IndirectObject):
                                            font_desc = font_desc.get_object()
                                        if isinstance(font_desc, dict):
                                            # 检查是否有字体文件（嵌入的字体）
                                            if any(key in font_desc for key in ['/FontFile', '/FontFile2', '/FontFile3']):
                                                is_embedded = True
                                                font_embedded_count += 1
                                            else:
                                                font_referenced_count += 1
                                    
                                    font_status = "已嵌入" if is_embedded else "仅引用（未嵌入）"
                                    print(f"[PDF预览]   - {font_name}: {base_font} ({font_status})")
                                    
                                    # 检查是否是中文字体
                                    if any(keyword in str(base_font).lower() for keyword in ['song', 'simsun', '宋', 'hei', 'simhei', '黑', 'wqy', 'wenquanyi']):
                                        if not is_embedded:
                                            print(f"[PDF预览]     ⚠️ 警告：中文字体未嵌入，在不同系统上可能显示不同字体")
                                except Exception as e:
                                    print(f"[PDF预览]   - {font_name}: (无法读取字体信息: {e})")
                            
                            print(f"[PDF预览] 字体统计: {font_embedded_count} 个已嵌入, {font_referenced_count} 个仅引用")
                            if font_referenced_count > 0:
                                print(f"[PDF预览] ⚠️ 注意：有 {font_referenced_count} 个字体未嵌入，在不同系统（如Mac）上可能显示不同字体")
                        else:
                            print(f"[PDF预览] ⚠️ PDF中未找到字体信息（fonts_used为空）")
                    else:
                        print(f"[PDF预览] ⚠️ PDF中未找到字体信息")
            except Exception as e:
                print(f"[PDF预览] ⚠️ 无法读取PDF字体信息: {e}")
                import traceback
                print(f"[PDF预览] 错误详情: {traceback.format_exc()}")
            
            # 检查PDF中是否包含中文字符（通过读取PDF文本内容）
            if html_chinese_count > 0:
                try:
                    # 尝试读取PDF的文本内容（使用pypdf）
                    from pypdf import PdfReader
                    reader = PdfReader(str(pdf_path))
                    pdf_text = ""
                    for page in reader.pages[:3]:  # 只检查前3页
                        pdf_text += page.extract_text() or ""
                    pdf_chinese_count = len([c for c in pdf_text if '\u4e00' <= c <= '\u9fff'])
                    print(f"[PDF预览] PDF中的中文字符数: {pdf_chinese_count} 字符")
                    if pdf_chinese_count == 0:
                        print(f"[PDF预览] ❌ 错误：HTML中有 {html_chinese_count} 个中文字符，但PDF中只有 {pdf_chinese_count} 个！")
                        print(f"[PDF预览] 💡 解决方案：服务器需要安装中文字体包")
                        print(f"[PDF预览] 💡 安装命令（CentOS/RHEL）: sudo yum install -y wqy-microhei-fonts wqy-zenhei-fonts")
                        print(f"[PDF预览] 💡 安装命令（Ubuntu/Debian）: sudo apt-get install -y fonts-wqy-microhei fonts-wqy-zenhei")
                    else:
                        print(f"[PDF预览] ✅ PDF中包含中文字符，字体支持正常")
                except Exception as e:
                    print(f"[PDF预览] 无法检查PDF中的中文字符: {e}")
            
            # 验证PDF文件是否有效（至少应该有一定大小）
            if pdf_size < 1024:  # 小于1KB可能有问题
                print(f"[PDF预览] 警告: PDF文件大小异常小 ({pdf_size} 字节)，可能生成失败")
                return False
            
            return True
            
        except Exception as e:
            print(f"[PDF预览] 生成PDF失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _extract_page_settings(self, document: Document) -> Dict:
        """从Word文档中提取页面设置"""
        settings = {
            "paper_size": "A4",  # 默认A4
            "margins": {
                "top": 2.54,      # 默认1英寸 = 2.54cm
                "bottom": 2.54,
                "left": 3.18,     # 默认1.25英寸 = 3.18cm
                "right": 3.18,
            },
            "orientation": "portrait"  # 默认纵向
        }
        
        try:
            # 获取第一个section的页面设置（通常所有section使用相同设置）
            if document.sections:
                section = document.sections[0]
                page_width = section.page_width
                page_height = section.page_height
                
                # 判断页面方向
                if page_width > page_height:
                    settings["orientation"] = "landscape"
                else:
                    settings["orientation"] = "portrait"
                
                # 判断纸张大小（转换为厘米）
                width_cm = page_width / 360000  # Word内部单位转换为厘米
                height_cm = page_height / 360000
                
                # 常见纸张大小判断
                if abs(width_cm - 21.0) < 0.5 and abs(height_cm - 29.7) < 0.5:
                    settings["paper_size"] = "A4"
                elif abs(width_cm - 21.59) < 0.5 and abs(height_cm - 27.94) < 0.5:
                    settings["paper_size"] = "Letter"
                elif abs(width_cm - 21.0) < 0.5 and abs(height_cm - 29.7) < 0.5:
                    settings["paper_size"] = "A4"
                else:
                    # 自定义大小，使用实际尺寸
                    settings["paper_size"] = f"{width_cm}cm {height_cm}cm"
                
                # 提取页边距（转换为厘米）
                # Word内部单位：1英寸 = 914400 EMU，1厘米 = 360000 EMU
                # 所以除以360000得到厘米
                top_margin_cm = section.top_margin / 360000
                bottom_margin_cm = section.bottom_margin / 360000
                left_margin_cm = section.left_margin / 360000
                right_margin_cm = section.right_margin / 360000
                
                settings["margins"]["top"] = round(top_margin_cm, 2)
                settings["margins"]["bottom"] = round(bottom_margin_cm, 2)
                settings["margins"]["left"] = round(left_margin_cm, 2)
                settings["margins"]["right"] = round(right_margin_cm, 2)
                
                print(f"[PDF预览] 提取页面设置: {settings['paper_size']}, 方向: {settings['orientation']}")
                print(f"[PDF预览] 提取页边距: 上={settings['margins']['top']}cm, 下={settings['margins']['bottom']}cm, 左={settings['margins']['left']}cm, 右={settings['margins']['right']}cm")
                print(f"[PDF预览] Word原始页边距(EMU): 上={section.top_margin}, 下={section.bottom_margin}, 左={section.left_margin}, 右={section.right_margin}")
        except Exception as e:
            print(f"[PDF预览] 提取页面设置失败，使用默认值: {e}")
        
        return settings
    
    def _generate_pdf_css(self, page_settings: Dict) -> str:
        """根据页面设置生成PDF CSS"""
        paper_size = page_settings.get("paper_size", "A4")
        orientation = page_settings.get("orientation", "portrait")
        margins = page_settings.get("margins", {})
        
        # 构建@page规则
        # 使用提取到的实际页边距，如果没有则使用Word文档的默认值（2.54cm = 1英寸）
        # 注意：不要使用硬编码的默认值，应该从Word文档中提取
        margin_top = f"{margins.get('top', 2.54):.2f}cm" if margins.get('top') else "2.54cm"
        margin_bottom = f"{margins.get('bottom', 2.54):.2f}cm" if margins.get('bottom') else "2.54cm"
        margin_left = f"{margins.get('left', 3.18):.2f}cm" if margins.get('left') else "3.18cm"
        margin_right = f"{margins.get('right', 2.54):.2f}cm" if margins.get('right') else "2.54cm"
        
        print(f"[PDF预览] 应用页边距: 上={margin_top}, 下={margin_bottom}, 左={margin_left}, 右={margin_right}")
        
        # 如果纸张大小是自定义的，直接使用
        if "cm" in str(paper_size) and " " in str(paper_size):
            size_value = paper_size
        else:
            # 标准纸张大小
            size_value = paper_size
            if orientation == "landscape":
                size_value = f"{paper_size} landscape"
        
        css = f"""
            @page {{
                size: {size_value};
                margin-top: {margin_top};
                margin-bottom: {margin_bottom};
                margin-left: {margin_left};
                margin-right: {margin_right};
                /* 使用背景色和边框让分页更明显 */
                background: #ffffff;
            }}
            @page:first {{
                /* 第一页特殊处理 */
            }}
            /* 在每页底部添加分页线 - 使用更简单的方法 */
            @page {{
                @bottom-center {{
                    content: "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ content: "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ content: "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ HTML生成时，在每页底部添加可见的分页线。让我改用更简单的方法：在HTML内容中直接插入HR标签或带边框的div。
</think>
改用更直接的方法：在 HTML 内容中直接插入可见的分页线元素：
<｜tool▁calls▁begin｜><｜tool▁call▁begin｜>
read_file
            body {{
                font-family: "SimSun", "宋体", "STSong", "STSongti-SC-Regular", "WenQuanYi Micro Hei", "WenQuanYi Zen Hei", serif;
                padding: 0;
                margin: 0;
                background: #ffffff;
                /* 确保中文字符正确显示 */
                -webkit-font-smoothing: antialiased;
            }}
            /* 默认段落字体：宋体（内联样式会覆盖此设置） */
            /* 注意：WeasyPrint需要系统安装对应的字体文件才能正确渲染 */
            p {{
                font-family: "SimSun", "宋体", "STSong", "STSongti-SC-Regular", "WenQuanYi Micro Hei", "WenQuanYi Zen Hei", serif;
            }}
            /* 默认标题字体：黑体（内联样式会覆盖此设置） */
            h1, h2, h3, h4, h5, h6 {{
                font-family: "SimHei", "黑体", "STHeiti", "WenQuanYi Micro Hei", "WenQuanYi Zen Hei", sans-serif;
            }}
            /* 支持多种字体：黑体、宋体、楷体、Times New Roman */
            /* 这些字体栈会在内联样式中被使用 */
            /* 确保图片正确显示 */
            img {{
                max-width: 100%;
                height: auto;
                display: block;
                margin: 10px auto;
                page-break-inside: avoid;
            }}
            /* 分页控制 - 只在打印或PDF生成时生效，浏览器预览中不显示空白页 */
            .page-break {{
                /* 在浏览器中隐藏，不占用空间 */
                display: none;
                /* 只在打印或PDF生成时生效 */
                page-break-before: always;
            }}
            /* 打印时显示分页符 */
            @media print {{
                .page-break {{
                    display: block;
                    page-break-before: always;
                    height: 0;
                    margin: 0;
                    padding: 0;
                }}
            }}
            /* 在每个段落后添加轻微的分隔（帮助识别分页） */
            p {{
                orphans: 3;
                widows: 3;
                margin-bottom: 0.5em;
            }}
            /* 在标题后添加更多间距，帮助识别分页 */
            h1, h2, h3, h4, h5, h6 {{
                page-break-after: avoid;
                margin-top: 1em;
                margin-bottom: 0.5em;
            }}
            /* 图片和表格分页控制 */
            img, table {{
                page-break-inside: avoid;
            }}
            /* 文档容器样式 - 添加边框让分页更明显 */
            .document-container {{
                border: 1px solid #e0e0e0;
                padding: 20px;
                margin: 0;
                background: #ffffff;
                /* 每页都有独立的容器边框 */
                box-shadow: 0 0 0 1px #d0d0d0;
            }}
            /* 在每页底部添加分页标记 */
            .page-end {{
                border-bottom: 2px solid #cccccc;
                margin-bottom: 20px;
                padding-bottom: 10px;
            }}
            .watermark {{
                position: fixed;
                top: 50%;
                left: 50%;
                /* 完全移除transform和复杂的CSS属性，避免WeasyPrint内部错误 */
                margin-top: -36px;  /* 字体大小的一半 */
                margin-left: -200px;  /* 大约文本宽度的一半 */
                font-size: 72px;
                color: rgba(209, 15, 15, 0.15);
                font-weight: bold;
                pointer-events: none;
                z-index: 1;
                /* 移除writing-mode，避免可能的兼容性问题 */
            }}
            """
        return css
    
    def _verify_format_changes(self, original_path: Path, final_path: Path, rules: Dict) -> Dict:
        """验证格式修改是否正确：对比原始文档和修改后的文档"""
        from docx import Document
        # docx_format_utils 已经在文件顶部导入，直接使用即可
        
        verification = {
            "summary": {},
            "errors": [],
            "warnings": []
        }
        
        try:
            original_doc = Document(original_path)
            final_doc = Document(final_path)
            
            # 确保两个文档的段落数量一致
            if len(original_doc.paragraphs) != len(final_doc.paragraphs):
                verification["warnings"].append(
                    f"段落数量不一致：原始文档 {len(original_doc.paragraphs)} 段，修改后 {len(final_doc.paragraphs)} 段"
                )
            
            # 对比每个段落的格式
            total_paragraphs = min(len(original_doc.paragraphs), len(final_doc.paragraphs))
            format_changes_count = 0
            font_correct_count = 0
            line_spacing_correct_count = 0
            format_errors = []
            
            for idx in range(total_paragraphs):
                orig_para = original_doc.paragraphs[idx]
                final_para = final_doc.paragraphs[idx]
                
                orig_format = docx_format_utils.extract_paragraph_format(orig_para)
                final_format = docx_format_utils.extract_paragraph_format(final_para)
                
                # 检查格式是否有变化
                has_changes = False
                for key in orig_format:
                    if orig_format.get(key) != final_format.get(key):
                        has_changes = True
                        break
                
                if has_changes:
                    format_changes_count += 1
                    
                    # 验证字体是否正确（正文应该是宋体，标题应该是黑体）
                    text = final_para.text.strip()
                    if text:
                        # 判断段落类型
                        is_title = any(keyword in final_para.style.name.lower() for keyword in ["heading", "标题", "title"])
                        
                        if is_title:
                            # 标题应该是黑体
                            if final_format.get("font_name") and "黑" in final_format.get("font_name", ""):
                                font_correct_count += 1
                            elif final_format.get("font_name"):
                                format_errors.append(
                                    f"段落 {idx}（标题）字体应为黑体，实际为：{final_format.get('font_name')}"
                                )
                        else:
                            # 正文应该是宋体
                            if final_format.get("font_name") and "宋" in final_format.get("font_name", ""):
                                font_correct_count += 1
                        
                        # 验证行距（正文应该是20磅）
                        if not is_title and final_format.get("line_spacing"):
                            line_spacing = final_format.get("line_spacing")
                            if isinstance(line_spacing, (int, float)) and abs(line_spacing - 20) < 1:
                                line_spacing_correct_count += 1
                            elif isinstance(line_spacing, (int, float)) and line_spacing != 20:
                                format_errors.append(
                                    f"段落 {idx}（正文）行距应为20磅，实际为：{line_spacing}"
                                )
            
            verification["summary"] = {
                "总段落数": total_paragraphs,
                "格式修改段落数": format_changes_count,
                "字体正确段落数": font_correct_count,
                "行距正确段落数": line_spacing_correct_count
            }
            
            # 如果格式错误太多，记录为错误
            if format_errors:
                verification["errors"].extend(format_errors[:20])  # 最多记录20个错误
            
            # 如果格式修改的段落太少，可能是格式没有正确应用
            if format_changes_count == 0 and total_paragraphs > 0:
                verification["warnings"].append("未检测到任何格式修改，可能格式规则未正确应用")
            
        except Exception as e:
            verification["errors"].append(f"格式验证过程出错: {e}")
            import traceback
            print(f"[格式验证] 错误堆栈: {traceback.format_exc()}")
        
        return verification
    
    def _try_libreoffice_conversion(self, docx_path: Path, html_path: Path, stats: Dict) -> bool:
        """尝试使用LibreOffice将Word文档转换为HTML（保留格式最好）"""
        import subprocess
        import shutil
        
        print("[HTML预览] 检查LibreOffice是否可用...")
        # 检查LibreOffice是否可用
        libreoffice_cmd = None
        for cmd in ['libreoffice', 'soffice']:
            if shutil.which(cmd):
                libreoffice_cmd = cmd
                print(f"[HTML预览] 找到LibreOffice命令: {cmd}")
                break
        
        if not libreoffice_cmd:
            print("[HTML预览] LibreOffice未安装，使用自定义HTML生成")
            return False
        
        try:
            # 创建临时目录用于输出
            temp_dir = html_path.parent / "temp_html"
            temp_dir.mkdir(exist_ok=True)
            
            # 使用LibreOffice转换
            # --headless: 无界面模式
            # --convert-to html: 转换为HTML
            # --outdir: 输出目录
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'html',
                '--outdir', str(temp_dir),
                str(docx_path)
            ]
            
            print(f"[HTML预览] 执行LibreOffice转换命令: {' '.join(cmd)}")
            
            # 准备环境变量，确保包含必要的 PATH
            import os
            env = os.environ.copy()
            # 确保 PATH 包含 /usr/bin 和 /bin（uname 等命令需要）
            current_path = env.get('PATH', '')
            if '/usr/bin' not in current_path:
                env['PATH'] = f"/usr/bin:/bin:{current_path}"
            if '/bin' not in env['PATH']:
                env['PATH'] = f"/bin:{env['PATH']}"
            
            print(f"[HTML预览] 开始执行LibreOffice转换，超时时间: 60秒...")
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,  # 60秒超时
                env=env  # 使用包含完整 PATH 的环境变量
            )
            print(f"[HTML预览] LibreOffice转换完成，返回码: {result.returncode}")
            
            if result.returncode != 0:
                print(f"[HTML预览] LibreOffice转换失败: {result.stderr}")
                return False
            
            # 查找生成的HTML文件
            html_file_name = docx_path.stem + '.html'
            generated_html = temp_dir / html_file_name
            
            if not generated_html.exists():
                print(f"[HTML预览] LibreOffice生成的HTML文件不存在: {generated_html}")
                return False
            
            # 读取生成的HTML内容
            html_content = generated_html.read_text(encoding='utf-8', errors='ignore')
            
            # 清理临时文件
            try:
                generated_html.unlink()
                temp_dir.rmdir()
            except:
                pass
            
            # 不再在预览文档中添加检测结果，保持文档干净
            # 检测结果只在首页（报告）中显示
            
            # 添加水印和警告样式
            watermark_style = """
            <style>
                .preview-watermark {
                    position: fixed;
                    top: 50%;
                    left: 50%;
                    /* 移除transform，避免WeasyPrint内部错误 */
                    margin-top: -36px;
                    margin-left: -200px;
                    font-size: 72px;
                    color: rgba(209, 15, 15, 0.15);
                    font-weight: bold;
                    pointer-events: none;
                    z-index: 9999;
                    white-space: nowrap;
                }
                .preview-warning {
                    background: #fff3cd;
                    border: 1px solid #ffc107;
                    padding: 15px;
                    margin: 20px 0;
                    border-radius: 5px;
                    text-align: center;
                    font-weight: bold;
                    color: #856404;
                }
            </style>
            """
            
            # 在head标签中插入样式
            if '</head>' in html_content:
                html_content = html_content.replace('</head>', watermark_style + '</head>')
            
            # 在body标签后插入水印（不插入检测结果）
            if '<body' in html_content:
                # 找到body标签结束位置
                body_end = html_content.find('>', html_content.find('<body'))
                if body_end != -1:
                    insert_pos = body_end + 1
                    insert_content = '<div class="preview-watermark">预览版 仅供查看</div>'
                    html_content = html_content[:insert_pos] + insert_content + html_content[insert_pos:]
            
            # 在文档末尾添加警告
            if '</body>' in html_content:
                warning_html = '<div class="preview-warning">⚠️ 这是预览版本，仅供查看。如需下载正式版，请完成支付。</div>'
                html_content = html_content.replace('</body>', warning_html + '</body>')
            
            # 保存HTML文件
            html_path.write_text(html_content, encoding='utf-8')
            
            print(f"[HTML预览] LibreOffice转换成功，HTML大小: {len(html_content) / 1024:.2f} KB")
            return True
            
        except subprocess.TimeoutExpired:
            print("[HTML预览] LibreOffice转换超时")
            return False
        except Exception as e:
            print(f"[HTML预览] LibreOffice转换出错: {e}")
            return False
    
    def _try_libreoffice_pdf_conversion(self, docx_path: Path, pdf_path: Path) -> bool:
        """尝试使用LibreOffice将Word文档直接转换为PDF（最接近Word效果）"""
        import subprocess
        import shutil
        import os
        import sys
        
        # 立即输出日志
        log_msg = f"[PDF预览] ========== 开始 LibreOffice PDF 转换 =========="
        print(log_msg, file=sys.stderr, flush=True)
        try:
            with open("/var/log/geshixiugai/error.log", "a") as f:
                f.write(f"{log_msg}\n")
        except Exception:
            pass
        
        # 检查LibreOffice是否可用
        libreoffice_cmd = None
        
        # 方法1: 直接检查常见路径（最快最可靠）
        direct_paths = [
            '/bin/libreoffice',      # 阿里云 Linux
            '/bin/soffice',          # 阿里云 Linux
            '/usr/bin/libreoffice',  # 标准路径
            '/usr/bin/soffice',      # 标准路径
            '/usr/local/bin/libreoffice',
            '/usr/local/bin/soffice',
        ]
        
        for path in direct_paths:
            if os.path.exists(path) and os.access(path, os.X_OK):
                # 验证可以执行
                try:
                    result = subprocess.run(
                        [path, '--version'],
                        capture_output=True,
                        text=True,
                        timeout=3
                    )
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        log_msg = f"[PDF预览] 在路径找到LibreOffice: {path}"
                        print(log_msg, file=sys.stderr, flush=True)
                        try:
                            with open("/var/log/geshixiugai/error.log", "a") as f:
                                f.write(f"{log_msg}\n")
                        except Exception:
                            pass
                        log_msg = f"[PDF预览] LibreOffice版本: {result.stdout.strip()}"
                        print(log_msg, file=sys.stderr, flush=True)
                        try:
                            with open("/var/log/geshixiugai/error.log", "a") as f:
                                f.write(f"{log_msg}\n")
                        except Exception:
                            pass
                        break
                except Exception as e:
                    print(f"[PDF预览] 路径 {path} 存在但无法执行: {e}")
                    continue
        
        # 方法2: 使用 which 查找（如果直接路径找不到）
        if not libreoffice_cmd:
            for cmd in ['libreoffice', 'soffice']:
                cmd_path = shutil.which(cmd)
                if cmd_path:
                    libreoffice_cmd = cmd_path
                    print(f"[PDF预览] 通过 which 找到LibreOffice命令: {cmd_path}")
                    break
        
        # 方法3: 尝试直接执行命令（可能在 PATH 中但 which 检测不到）
        if not libreoffice_cmd:
            for cmd in ['libreoffice', 'soffice']:
                try:
                    # 尝试执行 --version 命令来验证
                    result = subprocess.run(
                        [cmd, '--version'],
                        capture_output=True,
                        text=True,
                        timeout=5,
                        env=os.environ.copy()  # 使用当前环境变量
                    )
                    if result.returncode == 0:
                        libreoffice_cmd = cmd
                        print(f"[PDF预览] 通过执行验证找到LibreOffice命令: {cmd}")
                        print(f"[PDF预览] LibreOffice版本: {result.stdout.strip()}")
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired) as e:
                    print(f"[PDF预览] 命令 {cmd} 执行失败: {e}")
                    continue
        
        # 方法3: 尝试常见安装路径（包括 macOS）
        if not libreoffice_cmd:
            import platform
            is_macos = platform.system() == 'Darwin'
            
            common_paths = [
                '/bin/libreoffice',  # 阿里云 Linux 常见路径
                '/bin/soffice',       # 阿里云 Linux 常见路径（soffice）
                '/usr/bin/libreoffice',
                '/usr/bin/soffice',
                '/usr/local/bin/libreoffice',
                '/usr/local/bin/soffice',
                '/opt/libreoffice*/program/soffice',
            ]
            
            # macOS 特有路径
            if is_macos:
                # Homebrew 安装路径（Intel 和 Apple Silicon）
                common_paths.extend([
                    '/opt/homebrew/bin/libreoffice',  # Apple Silicon (M1/M2)
                    '/usr/local/bin/libreoffice',     # Intel
                    '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # 手动安装
                    '/Applications/LibreOffice.app/Contents/MacOS/libreoffice',  # 手动安装（如果存在）
                ])
            
            for path_pattern in common_paths:
                if '*' in path_pattern:
                    # 处理通配符路径
                    import glob
                    matches = glob.glob(path_pattern)
                    if matches:
                        libreoffice_cmd = matches[0]
                        print(f"[PDF预览] 在常见路径找到LibreOffice: {libreoffice_cmd}")
                        break
                else:
                    if os.path.exists(path_pattern) and os.access(path_pattern, os.X_OK):
                        libreoffice_cmd = path_pattern
                        print(f"[PDF预览] 在常见路径找到LibreOffice: {libreoffice_cmd}")
                        break
        
        if not libreoffice_cmd:
            log_msg = "[PDF预览] LibreOffice未找到，无法使用LibreOffice转换PDF"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            log_msg = "[PDF预览] 诊断信息:"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            log_msg = f"[PDF预览]   - PATH环境变量: {os.environ.get('PATH', '未设置')}"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            log_msg = f"[PDF预览]   - 检查的路径: {direct_paths}"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            # 检查哪些路径存在
            existing_paths = [p for p in direct_paths if os.path.exists(p)]
            if existing_paths:
                print(f"[PDF预览]   - 存在的路径: {existing_paths}")
                for p in existing_paths:
                    stat_info = os.stat(p)
                    print(f"[PDF预览]     {p}: 权限={oct(stat_info.st_mode)}, 所有者UID={stat_info.st_uid}")
            else:
                print(f"[PDF预览]   - 没有找到任何LibreOffice路径")
            print("[PDF预览] 提示: 请确保LibreOffice已安装并在PATH中，或使用 'which libreoffice' 检查")
            return False
        
        try:
            # 准备输出目录
            output_dir = pdf_path.parent
            
            # 使用LibreOffice转换
            # --headless: 无界面模式
            # --convert-to pdf: 转换为PDF，并嵌入字体以确保跨平台一致性
            # --outdir: 输出目录
            # 使用标准PDF转换命令
            # 注意：LibreOffice默认会嵌入可用字体，确保跨平台一致性
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                str(docx_path)
            ]
            
            log_msg = f"[PDF预览] 执行LibreOffice PDF转换命令: {' '.join(cmd)}"
            print(log_msg, flush=True)
            
            # 准备环境变量，确保包含必要的 PATH 和 HOME
            env = os.environ.copy()
            # 确保 PATH 包含 /usr/bin 和 /bin（uname 等命令需要）
            current_path = env.get('PATH', '')
            if '/usr/bin' not in current_path:
                env['PATH'] = f"/usr/bin:/bin:{current_path}"
            if '/bin' not in env['PATH']:
                env['PATH'] = f"/bin:{env['PATH']}"
            # 设置 HOME 环境变量（LibreOffice 可能需要）
            if 'HOME' not in env:
                env['HOME'] = str(abs_output_dir.parent)
            # 设置用户目录（LibreOffice 可能需要）
            if 'USER' not in env:
                import pwd
                try:
                    current_uid = os.getuid()
                    user_info = pwd.getpwuid(current_uid)
                    env['USER'] = user_info.pw_name
                    env['HOME'] = user_info.pw_dir
                except Exception:
                    pass
            
            # 执行转换（超时60秒）
            # 注意：使用绝对路径，避免相对路径问题
            abs_docx_path = docx_path.resolve()
            abs_output_dir = output_dir.resolve()
            
            # 确保输出目录存在且权限正确
            abs_output_dir.mkdir(parents=True, exist_ok=True)
            os.chmod(abs_output_dir, 0o755)
            
            # 确保输入文件可读
            os.chmod(abs_docx_path, 0o644)
            
            # 处理中文文件名问题：总是使用临时文件，避免中文文件名导致的问题
            import tempfile
            import shutil
            import uuid
            temp_input = None
            temp_output_name = None
            use_temp_file = True  # 总是使用临时文件，避免中文文件名问题
            
            # 生成唯一的临时文件名
            temp_id = str(uuid.uuid4())[:8]
            temp_input = abs_output_dir / f"temp_input_{temp_id}{abs_docx_path.suffix}"
            temp_output_name = f"temp_input_{temp_id}.pdf"
            
            # 复制文件到临时文件
            shutil.copy2(abs_docx_path, temp_input)
            os.chmod(temp_input, 0o644)
            log_msg = f"[PDF预览] 使用临时文件避免中文文件名问题: {temp_input} -> {temp_output_name}"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            
            # 使用临时文件
            input_file = temp_input
            expected_pdf_name = temp_output_name
            
            # 添加环境变量，确保LibreOffice使用正确的字体设置
            # 禁用OpenCL，避免渲染问题
            env['SAL_DISABLE_OPENCL'] = '1'
            # 设置LibreOffice用户配置目录（如果不存在则创建）
            lo_user_dir = Path.home() / '.config' / 'libreoffice'
            if not lo_user_dir.exists():
                try:
                    lo_user_dir.mkdir(parents=True, exist_ok=True)
                except Exception:
                    pass
            
            # 重要：LibreOffice在转换PDF时会保留原始字体
            # 但如果系统缺少字体，会自动替换为可用字体
            # 解决方案：
            # 1. 安装所需字体（见 安装中文字体.sh）
            # 2. 配置字体替换规则（见 配置LibreOffice字体替换.sh）
            # 3. 确保PDF导出时嵌入字体（LibreOffice默认会嵌入可用字体）
            
            log_msg = f"[PDF预览] LibreOffice配置目录: {lo_user_dir}"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            
            # 检查文件权限，如果文件属于其他用户，可能需要使用sudo
            # 但首先尝试直接执行
            # 注意：使用PDF导出过滤器，确保保留原始字体
            # LibreOffice默认会保留字体，但如果系统缺少字体，可能会替换
            # 尝试方法1：使用标准PDF转换命令
            # 注意：LibreOffice 7.1.8.1版本可能不支持复杂的JSON格式参数
            # 先使用简单的PDF转换，LibreOffice默认会嵌入可用字体
            cmd_abs = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(abs_output_dir),
                str(input_file)
            ]
            
            log_msg = f"[PDF预览] 尝试方法1：使用PDF转换命令，保留原始字体"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            
            log_msg = f"[PDF预览] 使用绝对路径执行命令: {' '.join(cmd_abs)}"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            log_msg = f"[PDF预览] 输入文件: {abs_docx_path}, 存在: {abs_docx_path.exists()}"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            log_msg = f"[PDF预览] 输出目录: {abs_output_dir}, 存在: {abs_output_dir.exists()}"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            
            # 检查文件权限
            try:
                file_stat = abs_docx_path.stat()
                log_msg = f"[PDF预览] 文件权限: {oct(file_stat.st_mode)}, 所有者UID: {file_stat.st_uid}, GID: {file_stat.st_gid}"
                print(log_msg, file=sys.stderr, flush=True)
                try:
                    with open("/var/log/geshixiugai/error.log", "a") as f:
                        f.write(f"{log_msg}\n")
                except Exception:
                    pass
                # 检查当前用户
                current_uid = os.getuid()
                log_msg = f"[PDF预览] 当前用户UID: {current_uid}"
                print(log_msg, file=sys.stderr, flush=True)
                try:
                    with open("/var/log/geshixiugai/error.log", "a") as f:
                        f.write(f"{log_msg}\n")
                except Exception:
                    pass
            except Exception as e:
                log_msg = f"[PDF预览] 无法获取文件权限信息: {e}"
                print(log_msg, file=sys.stderr, flush=True)
                try:
                    with open("/var/log/geshixiugai/error.log", "a") as f:
                        f.write(f"{log_msg}\n")
                except Exception:
                    pass
            
            result = subprocess.run(
                cmd_abs,
                capture_output=True,
                text=True,
                timeout=60,
                env=env  # 使用包含完整 PATH 的环境变量
            )
            
            log_msg = f"[PDF预览] LibreOffice执行完成，返回码: {result.returncode}"
            print(log_msg, file=sys.stderr, flush=True)
            try:
                with open("/var/log/geshixiugai/error.log", "a") as f:
                    f.write(f"{log_msg}\n")
            except Exception:
                pass
            if result.stdout:
                log_msg = f"[PDF预览] LibreOffice标准输出: {result.stdout}"
                print(log_msg, file=sys.stderr, flush=True)
                try:
                    with open("/var/log/geshixiugai/error.log", "a") as f:
                        f.write(f"{log_msg}\n")
                except Exception:
                    pass
            if result.stderr:
                log_msg = f"[PDF预览] LibreOffice错误输出: {result.stderr}"
                print(log_msg, file=sys.stderr, flush=True)
                try:
                    with open("/var/log/geshixiugai/error.log", "a") as f:
                        f.write(f"{log_msg}\n")
                except Exception:
                    pass
            
            # 如果出现 "no export filter" 错误，尝试使用不同的方法
            if result.stderr and "no export filter" in result.stderr.lower():
                log_msg = f"[PDF预览] 检测到 'no export filter' 错误，尝试使用 soffice 命令"
                print(log_msg, file=sys.stderr, flush=True)
                try:
                    with open("/var/log/geshixiugai/error.log", "a") as f:
                        f.write(f"{log_msg}\n")
                except Exception:
                    pass
                
                # 尝试使用 soffice 命令
                soffice_cmd = None
                for path in ['/bin/soffice', '/usr/bin/soffice']:
                    if os.path.exists(path) and os.access(path, os.X_OK):
                        soffice_cmd = path
                        break
                
                if soffice_cmd:
                    # 使用标准PDF转换命令
                    cmd_abs2 = [
                        soffice_cmd,
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', str(abs_output_dir),
                        str(input_file)
                    ]
                    
                    log_msg = f"[PDF预览] 尝试使用 soffice 命令: {' '.join(cmd_abs2)}"
                    print(log_msg, file=sys.stderr, flush=True)
                    try:
                        with open("/var/log/geshixiugai/error.log", "a") as f:
                            f.write(f"{log_msg}\n")
                    except Exception:
                        pass
                    
                    result = subprocess.run(
                        cmd_abs2,
                        capture_output=True,
                        text=True,
                        timeout=60,
                        env=env
                    )
                    
                    log_msg = f"[PDF预览] soffice 执行完成，返回码: {result.returncode}"
                    print(log_msg, file=sys.stderr, flush=True)
                    try:
                        with open("/var/log/geshixiugai/error.log", "a") as f:
                            f.write(f"{log_msg}\n")
                    except Exception:
                        pass
                    if result.stderr:
                        log_msg = f"[PDF预览] soffice 错误输出: {result.stderr}"
                        print(log_msg, file=sys.stderr, flush=True)
                        try:
                            with open("/var/log/geshixiugai/error.log", "a") as f:
                                f.write(f"{log_msg}\n")
                        except Exception:
                            pass
            
            # LibreOffice 会在输出目录生成与输入文件同名的PDF
            # 例如：preview.docx -> preview.pdf
            # 使用绝对路径查找
            generated_pdf = abs_output_dir / expected_pdf_name
            
            # 等待文件生成（LibreOffice 可能需要一点时间）
            import time
            max_wait = 10  # 增加到10秒，给LibreOffice更多时间
            wait_interval = 0.5  # 每0.5秒检查一次
            waited = 0
            while not generated_pdf.exists() and waited < max_wait:
                time.sleep(wait_interval)
                waited += wait_interval
                log_msg = f"[PDF预览] 等待PDF文件生成... ({waited:.1f}s)"
                print(log_msg, file=sys.stderr, flush=True)
                try:
                    with open("/var/log/geshixiugai/error.log", "a") as f:
                        f.write(f"{log_msg}\n")
                except Exception:
                    pass
            
            # 检查是否成功生成PDF文件（即使返回码非零，也可能成功转换）
            if generated_pdf.exists():
                # 验证文件大小（应该大于0）
                pdf_size = generated_pdf.stat().st_size
                if pdf_size > 0:
                    log_msg = f"[PDF预览] ✅ PDF文件已生成: {generated_pdf}, 大小: {pdf_size / 1024:.2f} KB"
                    print(log_msg, file=sys.stderr, flush=True)
                    try:
                        with open("/var/log/geshixiugai/error.log", "a") as f:
                            f.write(f"{log_msg}\n")
                    except Exception:
                        pass
                    # 即使返回码非零，只要文件生成了就认为成功
                    # 将生成的PDF移动到目标路径
                    shutil.move(str(generated_pdf), str(pdf_path))
                    log_msg = f"[PDF预览] LibreOffice PDF转换成功，文件已移动到: {pdf_path}"
                    print(log_msg, file=sys.stderr, flush=True)
                    try:
                        with open("/var/log/geshixiugai/error.log", "a") as f:
                            f.write(f"{log_msg}\n")
                    except Exception:
                        pass
                    # 清理临时文件
                    if temp_input and temp_input.exists():
                        temp_input.unlink()
                    return True
                else:
                    log_msg = f"[PDF预览] ⚠️ PDF文件存在但大小为0: {generated_pdf}"
                    print(log_msg, file=sys.stderr, flush=True)
                    try:
                        with open("/var/log/geshixiugai/error.log", "a") as f:
                            f.write(f"{log_msg}\n")
                    except Exception:
                        pass
            
            # 如果找不到预期的PDF文件，尝试列出输出目录中的所有文件
            if not generated_pdf.exists():
                print(f"[PDF预览] 预期PDF文件不存在: {generated_pdf}")
                print(f"[PDF预览] 检查输出目录中的所有文件:")
                try:
                    all_files = list(output_dir.glob("*"))
                    for f in all_files:
                        print(f"[PDF预览]   - {f.name} ({f.stat().st_size} bytes)")
                    # 尝试查找任何 PDF 文件
                    pdf_files = list(output_dir.glob("*.pdf"))
                    if pdf_files:
                        print(f"[PDF预览] 找到 {len(pdf_files)} 个PDF文件，使用第一个: {pdf_files[0]}")
                        generated_pdf = pdf_files[0]
                    else:
                        log_msg = f"[PDF预览] 输出目录中没有找到任何PDF文件"
                        print(log_msg, flush=True)
                        if result.returncode != 0:
                            log_msg = f"[PDF预览] LibreOffice返回错误码: {result.returncode}"
                            print(log_msg, flush=True)
                            if result.stdout:
                                log_msg = f"[PDF预览] LibreOffice标准输出: {result.stdout}"
                                print(log_msg, flush=True)
                            if result.stderr:
                                log_msg = f"[PDF预览] LibreOffice错误输出: {result.stderr}"
                                print(log_msg, flush=True)
                        # 检查输出目录权限
                        try:
                            dir_stat = abs_output_dir.stat()
                            log_msg = f"[PDF预览] 输出目录权限: {oct(dir_stat.st_mode)}, 所有者UID: {dir_stat.st_uid}"
                            print(log_msg, flush=True)
                        except Exception as e:
                            log_msg = f"[PDF预览] 无法获取目录权限: {e}"
                            print(log_msg, flush=True)
                        # 输出详细的诊断信息
                        log_msg = f"[PDF预览] 转换失败诊断: 输入文件={abs_docx_path}, 输出目录={abs_output_dir}, LibreOffice命令={libreoffice_cmd}"
                        print(log_msg, flush=True)
                        return False
                except Exception as e:
                    print(f"[PDF预览] 列出文件时出错: {e}")
                    return False
            
            # 如果找到了PDF文件，移动到目标位置
            if generated_pdf.exists():
                pdf_size = generated_pdf.stat().st_size
                if pdf_size > 0:
                    # 移动到目标位置
                    if generated_pdf != pdf_path:
                        shutil.move(str(generated_pdf), str(pdf_path))
                    else:
                        # 如果已经是目标路径，验证文件大小
                        pdf_size = pdf_path.stat().st_size
                    
                    print(f"[PDF预览] ✅ LibreOffice PDF转换成功，大小: {pdf_size / 1024:.2f} KB")
                    
                    # 验证PDF文件是否有效（至少应该有一定大小）
                    if pdf_size < 1024:  # 小于1KB可能有问题
                        print(f"[PDF预览] ⚠️ 警告: PDF文件大小异常小 ({pdf_size} 字节)，可能生成失败")
                        return False
                    
                    return True
                else:
                    print(f"[PDF预览] ⚠️ PDF文件存在但大小为0: {generated_pdf}")
                    return False
            else:
                print(f"[PDF预览] ❌ 未找到生成的PDF文件")
                if result.returncode != 0:
                    print(f"[PDF预览] LibreOffice返回错误码: {result.returncode}")
                return False
            
        except subprocess.TimeoutExpired:
            print("[PDF预览] LibreOffice PDF转换超时（超过60秒）")
            return False
        except Exception as e:
            print(f"[PDF预览] LibreOffice PDF转换出错: {e}")
            import traceback
            traceback.print_exc()
            return False

